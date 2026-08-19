"""Convention detection from a reference DOCX.

The strongest check available: render a document in a known style, feed it back
through the extractor, and assert it recovers that style's conventions — columns,
heading numbering, caption position/wording, table rules, reference numbering.
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from app.docos.parser import parse_docx_bytes
from app.paper.renderer import render_paper
from app.paper.schema import PaperSpec
from app.paper.styles import get_stylesheet
from app.paper.styles.conventions import detect_conventions
from app.paper.styles.extract import derive_stylesheet_from_docx

SPEC = {
    "meta": {
        "title": "Round Trip",
        "authors": [{"name": "A. Author", "affiliation": "Lab"}],
        "abstract": "An abstract of sufficient length to be recognised.",
        "keywords": ["alpha", "beta"],
    },
    "blocks": [
        {"type": "heading", "level": 1, "text": "Introduction"},
        {"type": "paragraph", "text": "Body prose in the introduction section here."},
        {"type": "heading", "level": 2, "text": "Background"},
        {"type": "paragraph", "text": "More body prose supporting the background."},
        {"type": "heading", "level": 1, "text": "Methodology"},
        {"type": "paragraph", "text": "Prose describing the method in detail."},
        {"type": "table", "caption": "Performance results",
         "columns": ["Model", "Acc"], "rows": [["CNN", "0.94"], ["SVM", "0.85"]]},
        {"type": "figure", "caption": "Accuracy by model",
         "chart": {"kind": "bar", "title": "Accuracy",
                   "labels": ["CNN", "SVM"], "values": [0.94, 0.85]}},
        {"type": "heading", "level": 1, "text": "Conclusion"},
        {"type": "paragraph", "text": "Concluding prose for the round-trip sample."},
    ],
    "references": [
        "A. Author, \"A title,\" Journal, 2023.",
        "B. Writer, \"Another title,\" Proc. Conf., 2024.",
    ],
}


def _rendered(style: str, tmp_path) -> bytes:
    out = tmp_path / f"{style}.docx"
    render_paper(PaperSpec.model_validate(SPEC), out, style=style)
    return out.read_bytes()


def _detect(data: bytes) -> dict:
    graph = parse_docx_bytes(data, title="t")
    return detect_conventions(Document(io.BytesIO(data)), graph)


# ── round trip: IEEE ────────────────────────────────────────────────────────

def test_detects_ieee_conventions(tmp_path):
    found = _detect(_rendered("ieee", tmp_path))

    assert found["columns"] == 2
    assert found["heading_scheme"] == "roman_alpha"
    assert found["table_caption_position"] == "above"
    assert found["table_number_style"] == "roman"
    assert found["figure_caption_position"] == "below"
    assert found["table_borders"] == "horizontal"
    assert found["number_references"] is True
    assert found["references_title"] == "References"
    # IEEE: "TABLE I" with the title on the next line; "Fig. 1. " inline
    assert found["table_caption_prefix"] == "TABLE {num}"
    assert found["table_caption_separator"] == "\n"
    assert found["figure_caption_prefix"] == "Fig. {num}. "
    assert found["abstract_lead"].startswith("Abstract")
    assert found["abstract_as_heading"] is False


# ── round trip: assignment (decimal, single column, gridded) ────────────────

def test_detects_assignment_conventions(tmp_path):
    found = _detect(_rendered("assignment", tmp_path))

    assert found["columns"] == 1
    assert found["heading_scheme"] == "decimal"
    assert found["table_borders"] == "grid"
    assert found["table_header_fill"] == "EDEFF2"
    assert found["table_caption_position"] == "above"
    assert found["table_number_style"] == "arabic"
    assert found["figure_caption_position"] == "below"
    assert found["table_caption_prefix"] == "Table {num}. "
    assert found["abstract_as_heading"] is True


# ── round trip: IEEE in one column ──────────────────────────────────────────

def test_detects_ieee_1col_conventions(tmp_path):
    """Same conventions as IEEE, one column — the column count must be read from
    the sample rather than assumed from the rest of the style."""
    found = _detect(_rendered("ieee_1col", tmp_path))
    assert found["columns"] == 1
    assert found["heading_scheme"] == "roman_alpha"
    assert found["table_number_style"] == "roman"
    assert found["figure_caption_position"] == "below"


# ── the full derivation now carries conventions, not just fonts ─────────────

def test_derived_style_recovers_ieee_from_a_different_base(tmp_path):
    """Derive from an IEEE sample while basing on the assignment sheet — the
    conventions must come from the sample, not the base."""
    derived = derive_stylesheet_from_docx(
        _rendered("ieee", tmp_path), name="Learned IEEE",
        base=get_stylesheet("assignment"), source_filename="ieee_sample.docx",
    )
    base = get_stylesheet("assignment")
    assert base.page.columns == 1 and base.heading_scheme == "decimal"  # base differs

    assert derived.page.columns == 2                    # learned, not inherited
    assert derived.heading_scheme == "roman_alpha"
    assert derived.table_number_style == "roman"
    assert derived.table_borders == "horizontal"
    assert derived.figure_caption_position == "below"
    assert derived.table_caption_prefix == "TABLE {num}"
    assert "columns" in derived.detected
    assert "heading_scheme" in derived.detected


def test_derived_ieee_style_renders_like_ieee(tmp_path):
    """A style learned from an IEEE sample should itself produce IEEE-shaped output."""
    derived = derive_stylesheet_from_docx(
        _rendered("ieee", tmp_path), name="Learned IEEE",
        base=get_stylesheet("assignment"),
    )
    out = tmp_path / "again.docx"
    render_paper(PaperSpec.model_validate(SPEC), out, style=derived)

    again = _detect(out.read_bytes())
    assert again["columns"] == 2
    assert again["heading_scheme"] == "roman_alpha"
    assert again["table_caption_position"] == "above"
    assert again["figure_caption_position"] == "below"


# ── conservatism: say nothing when the sample says nothing ──────────────────

def test_empty_document_detects_nothing_significant():
    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    found = _detect(buf.getvalue())

    assert "heading_scheme" not in found          # no headings to judge
    assert "table_borders" not in found           # no tables
    assert "table_caption_position" not in found
    assert "number_references" not in found


def test_single_column_default_is_not_evidence():
    """Every blank Word document reports one column — that is the default, not a
    choice. A near-empty sample must not overwrite a two-column base."""
    doc = Document()
    doc.add_paragraph("One lonely paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    assert "columns" not in _detect(buf.getvalue())


def test_multi_column_is_always_trusted(tmp_path):
    """Two columns can only happen deliberately, so it needs no corroboration."""
    assert _detect(_rendered("ieee", tmp_path))["columns"] == 2


def test_unknown_conventions_keep_the_base(tmp_path):
    doc = Document()
    doc.add_paragraph("Just some prose with no structure to speak of.")
    buf = io.BytesIO()
    doc.save(buf)

    base = get_stylesheet("ieee")
    derived = derive_stylesheet_from_docx(buf.getvalue(), name="Sparse", base=base)
    # nothing detectable → base conventions preserved and still renderable
    assert derived.heading_scheme == base.heading_scheme
    assert derived.table_borders == base.table_borders
    assert derived.table_caption_prefix == base.table_caption_prefix
