"""Rendering an imported document the way the file actually specifies it.

A document imported in the viewer's default font, at the viewer's default line
spacing, does not look like the document — and the wrong metrics need more room
than the page has, which is how text ended up cut off at the foot of a page.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.docos.graph import NodeType  # noqa: E402
from app.docos.parser.docx_parser import parse_docx_bytes  # noqa: E402

_NS = ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
       'xmlns:v="urn:schemas-microsoft-com:vml" '
       'xmlns:o="urn:schemas-microsoft-com:office:office"')


def _render(doc):
    buf = io.BytesIO()
    doc.save(buf)
    return parse_docx_bytes(buf.getvalue(), title="t")


# ── typeface ────────────────────────────────────────────────────────────────

def test_the_documents_own_font_is_carried():
    """python-docx's template names its font by theme, as Word does."""
    graph = _render(Document())
    page = graph.root.metadata["page"]
    assert page["default_font"], "a theme font must be resolved to a real typeface"
    assert page["default_size_pt"] > 0


def test_a_font_set_on_the_style_is_found_without_run_formatting():
    """Most documents set their typeface once, on a style, never on a run."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Garamond"
    normal.font.size = Pt(13)
    doc.add_paragraph("styled by the style, not the run")

    body = next(n for n in _render(doc).root.children if n.type == NodeType.BODY)
    assert body.style.font_family == "Garamond"
    assert body.style.font_size == 13


def test_run_formatting_still_wins_over_the_style():
    doc = Document()
    doc.styles["Normal"].font.name = "Garamond"
    p = doc.add_paragraph()
    run = p.add_run("directly formatted")
    run.font.name = "Consolas"

    body = next(n for n in _render(doc).root.children if n.type == NodeType.BODY)
    assert body.style.font_family == "Consolas"


# ── spacing ─────────────────────────────────────────────────────────────────

def test_line_spacing_is_carried():
    doc = Document()
    p = doc.add_paragraph("double spaced")
    p.paragraph_format.line_spacing = 2.0

    body = next(n for n in _render(doc).root.children if n.type == NodeType.BODY)
    assert body.metadata["line_spacing"] == 2.0
    assert body.metadata["line_spacing_exact"] is False


def test_exact_line_spacing_is_marked_as_exact():
    doc = Document()
    p = doc.add_paragraph("exactly 18pt")
    p.paragraph_format.line_spacing = Pt(18)

    body = next(n for n in _render(doc).root.children if n.type == NodeType.BODY)
    assert body.metadata["line_spacing_exact"] is True
    assert body.metadata["line_spacing"] == 18.0


def test_paragraph_gaps_are_carried():
    doc = Document()
    p = doc.add_paragraph("spaced")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    body = next(n for n in _render(doc).root.children if n.type == NodeType.BODY)
    assert body.metadata["space_before_pt"] == 12.0
    assert body.metadata["space_after_pt"] == 6.0


# ── page boundaries around a rule ───────────────────────────────────────────

def test_a_rule_keeps_the_page_marker_it_sits_on():
    """The rule branch used to drop it, which moved every later page boundary
    and truncated the pages that held one."""
    doc = Document()
    doc.add_paragraph("before")
    p = doc.add_paragraph()
    p._p.append(parse_xml(
        f'<w:r {_NS}><w:lastRenderedPageBreak/>'
        '<w:pict><v:rect style="width:0;height:1.5pt" o:hr="t"/></w:pict></w:r>'))
    doc.add_paragraph("after")

    nodes = _render(doc).root.children
    marked = [n for n in nodes if n.metadata.get("page_break_before")]
    assert marked, "the page boundary must survive whatever the paragraph became"


# ── emphasis that lives on the style, not the runs ──────────────────────────

def _emphasis_docx() -> bytes:
    """Four paragraphs, each a different way of saying (or not saying) italic."""
    import io
    from docx import Document
    from docx.shared import Pt

    d = Document()
    p = d.add_paragraph(); r = p.add_run("Every run italic. " * 6); r.italic = True

    style = d.styles.add_style("AbstractStyle", 1)
    style.font.italic = True
    style.font.size = Pt(9)
    d.add_paragraph("Italic from the paragraph style. " * 6, style="AbstractStyle")

    p = d.add_paragraph()
    lead = p.add_run("A long italic lead. " * 6); lead.italic = True
    tail = p.add_run("Short roman tail."); tail.italic = False

    p = d.add_paragraph()
    label = p.add_run("Index Terms— "); label.italic = True
    body = p.add_run("A much longer upright body of text. " * 6); body.italic = False

    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


def test_italic_is_read_from_the_paragraph_style_too():
    """An IEEE abstract is italic because its *style* is, and no run says so."""
    graph = parse_docx_bytes(_emphasis_docx(), title="E")
    paras = [n for n in graph.nodes() if n.content and "Italic from the paragraph style" in n.content]
    assert paras and paras[0].style.italic is True
    assert paras[0].style.font_size == 9.0


def test_mixed_runs_are_weighted_by_how_much_text_carries_them():
    graph = parse_docx_bytes(_emphasis_docx(), title="E")
    by_text = {n.content[:20]: n.style for n in graph.nodes() if n.content}

    # A long italic lead outweighs a short upright tail...
    lead = next(s for t, s in by_text.items() if t.startswith("A long italic lead"))
    assert lead.italic is True
    # ...and a three-word italic label does not make a whole paragraph italic.
    label = next(s for t, s in by_text.items() if t.startswith("Index Terms—"))
    assert label.italic is False


def test_every_run_italic_still_reads_as_italic():
    graph = parse_docx_bytes(_emphasis_docx(), title="E")
    para = next(n for n in graph.nodes() if n.content and n.content.startswith("Every run italic"))
    assert para.style.italic is True


# ── inline runs ─────────────────────────────────────────────────────────────

def _mixed_runs_docx() -> bytes:
    """A paragraph whose words are not all formatted alike."""
    import io
    from docx import Document
    from docx.shared import Pt

    d = Document()
    p = d.add_paragraph()
    p.add_run("Recent work ")
    bold = p.add_run("substantially outperforms"); bold.bold = True
    p.add_run(" the earlier baseline ")
    cite = p.add_run("12"); cite.font.superscript = True
    p.add_run(" on every subject.")

    # Word loves to split a sentence into identically formatted pieces.
    p2 = d.add_paragraph()
    for piece in ("One ", "sentence ", "in ", "four ", "runs."):
        p2.add_run(piece)

    d.save(io := __import__("io").BytesIO()); return io.getvalue()


def test_inline_formatting_survives_the_paragraph():
    graph = parse_docx_bytes(_mixed_runs_docx(), title="M")
    para = next(n for n in graph.nodes() if n.content.startswith("Recent work"))

    assert para.content == "Recent work substantially outperforms the earlier baseline 12 on every subject."
    formatting = [(r.text, r.style.bold, r.style.vertical_align) for r in para.inline_runs()]
    assert ("substantially outperforms", True, None) in formatting
    assert ("12", None, "superscript") in formatting


def test_a_uniform_paragraph_carries_no_runs():
    """Runs describe variation; a paragraph without any needs none of them."""
    graph = parse_docx_bytes(_mixed_runs_docx(), title="M")
    para = next(n for n in graph.nodes() if n.content.startswith("One sentence"))
    assert para.runs == []
    assert [r.text for r in para.inline_runs()] == ["One sentence in four runs."]


def test_inline_formatting_survives_a_round_trip_through_docx():
    from app.docos.export import graph_to_docx_bytes

    before = parse_docx_bytes(_mixed_runs_docx(), title="M")
    after = parse_docx_bytes(graph_to_docx_bytes(before), title="M")

    def shape(graph):
        return [[(r.text, r.style.bold, r.style.italic, r.style.vertical_align)
                 for r in n.inline_runs()]
                for n in graph.nodes() if n.content]

    assert shape(after) == shape(before)


def test_rewriting_the_words_drops_formatting_that_described_the_old_ones():
    graph = parse_docx_bytes(_mixed_runs_docx(), title="M")
    para = next(n for n in graph.nodes() if n.content.startswith("Recent work"))
    assert len(para.inline_runs()) > 1

    para.set_text("A completely different sentence.")
    assert para.runs == []
    assert [r.text for r in para.inline_runs()] == ["A completely different sentence."]

    # Even a caller that forgets and assigns content directly stays safe.
    para2 = next(n for n in graph.nodes() if n.content.startswith("One sentence"))
    para2.runs = [__import__("app.docos.graph", fromlist=["Run"]).Run(text="stale")]
    para2.content = "new words entirely"
    assert [r.text for r in para2.inline_runs()] == ["new words entirely"]


# ── spacing a document states somewhere other than the paragraph ────────────

def _spaced_docx() -> bytes:
    """python-docx's template states 10pt after and 1.15 line spacing in
    docDefaults, and 24pt before a Heading 1 in the style — and nothing at all
    on any paragraph."""
    import io
    from docx import Document
    from docx.shared import Pt

    d = Document()
    d.add_heading("I. Introduction", 1)
    d.add_paragraph("A body paragraph that states no spacing of its own.")
    tight = d.add_paragraph("This one sets its own, and its own must win.")
    tight.paragraph_format.space_after = Pt(0)
    tight.paragraph_format.line_spacing = 1.0

    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


def test_spacing_is_inherited_from_the_document_defaults():
    graph = parse_docx_bytes(_spaced_docx(), title="S")
    body = next(n for n in graph.nodes() if n.content.startswith("A body paragraph"))

    assert body.metadata["space_after_pt"] == 10.0, "w:after=200 twips is 10pt"
    assert body.metadata["line_spacing"] == 1.15, "w:line=276 auto is 1.15"


def test_spacing_is_inherited_from_the_paragraph_style():
    graph = parse_docx_bytes(_spaced_docx(), title="S")
    heading = next(n for n in graph.nodes() if n.content.startswith("I. Introduction"))
    assert heading.metadata["space_before_pt"] == 24.0


def test_a_paragraph_that_states_its_own_spacing_still_wins():
    graph = parse_docx_bytes(_spaced_docx(), title="S")
    tight = next(n for n in graph.nodes() if n.content.startswith("This one sets"))
    assert tight.metadata["space_after_pt"] == 0.0
    assert tight.metadata["line_spacing"] == 1.0


# ── headings a document never declared ──────────────────────────────────────

def _hand_formatted_docx() -> bytes:
    """A paper written the way most are: Normal style, headings set by hand."""
    import io
    from docx import Document
    from docx.shared import Pt

    d = Document()
    d.styles["Normal"].font.size = Pt(10)

    def para(text, bold=False, size=10):
        p = d.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    para("I. INTRODUCTION", bold=True, size=12)
    para("Falls are a major cause of injury among older adults, motivating research.")
    para("A. Dataset", bold=True, size=11)
    para("MobiAct contains recordings from 67 subjects performing twelve activities.")
    para("II. RELATED WORK", bold=True, size=12)
    para("Most existing approaches formulate fall detection as a joint problem.")

    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


def test_headings_are_recognised_by_shape_when_no_style_says_so():
    graph = parse_docx_bytes(_hand_formatted_docx(), title="H")
    found = [(n.type.value, n.content) for n in graph.nodes()
             if n.type.value in ("heading", "subheading")]

    assert ("heading", "I. INTRODUCTION") in found, "a roman numeral opens a section"
    assert ("heading", "II. RELATED WORK") in found
    assert ("subheading", "A. Dataset") in found, "a letter numbers a subsection"


def test_a_bold_sentence_is_not_a_heading():
    """Calling a paragraph a heading is worse than missing one, so everything
    that merely looks emphatic stays where it is."""
    import io
    from docx import Document
    from docx.shared import Pt

    d = Document()
    d.styles["Normal"].font.size = Pt(10)
    for text, bold, size in [
        ("This sentence is bold for emphasis and should stay a paragraph.", True, 10),
        ("Figure 1. Distribution of per-subject accuracy across folds.", True, 9),
        ("Table II. Per-class precision and recall.", True, 9),
        ("1. Load the dataset from disk into memory before windowing begins.", False, 10),
        ("Note:", True, 10),
        ("Zhang, Y., and Li, Q. (2020). Fall detection. IEEE Access.", False, 10),
    ]:
        p = d.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    buf = io.BytesIO(); d.save(buf)

    graph = parse_docx_bytes(buf.getvalue(), title="F")
    assert not [n for n in graph.nodes() if n.type.value in ("heading", "subheading")]


def test_a_declared_heading_style_still_wins():
    """Inference only fills a silence; it never argues with the document."""
    import io
    from docx import Document

    d = Document()
    d.add_heading("Introduction", level=1)
    d.add_heading("Dataset", level=2)
    buf = io.BytesIO(); d.save(buf)

    graph = parse_docx_bytes(buf.getvalue(), title="D")
    kinds = [n.type.value for n in graph.nodes() if n.content]
    assert kinds == ["heading", "subheading"]


# ── Word's own equations ────────────────────────────────────────────────────

_MATH_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'


def _omml(inner: str) -> str:
    return f"<m:oMath {_MATH_NS}>{inner}</m:oMath>"


def _equation_docx(inline: bool = False) -> bytes:
    """A document with a real Word equation, the way Word stores one."""
    import io
    from docx import Document
    from docx.oxml import parse_xml

    fraction = ("<m:f><m:num><m:r><m:t>a</m:t></m:r></m:num>"
                "<m:den><m:r><m:t>b</m:t></m:r></m:den></m:f>")
    d = Document()
    if inline:
        p = d.add_paragraph()
        p.add_run("We minimise ")
        p._p.append(parse_xml(_omml(fraction)))
        p.add_run(" over the training set.")
    else:
        d.add_paragraph("The objective is defined below.")
        p = d.add_paragraph()
        p._p.append(parse_xml(_omml(fraction)))
        d.add_paragraph("It is minimised over the set.")

    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


def test_a_word_equation_is_read_rather_than_dropped():
    """A paragraph holding only an equation used to look empty and be skipped,
    so the equation left the document entirely."""
    graph = parse_docx_bytes(_equation_docx(), title="E")
    equation = next((n for n in graph.nodes() if n.metadata.get("equations")), None)

    assert equation is not None, "the equation's paragraph survived"
    assert equation.content == r"$\frac{a}{b}$"


def test_an_inline_equation_keeps_its_place_in_the_sentence():
    """Reading the runs and then the equations would put every equation at the
    end of its own paragraph."""
    graph = parse_docx_bytes(_equation_docx(inline=True), title="E")
    para = next(n for n in graph.nodes() if n.content)

    assert para.content == r"We minimise $\frac{a}{b}$ over the training set."


def test_an_untouched_equation_is_written_back_as_word_wrote_it():
    from docx.oxml.ns import qn
    from app.docos.export import graph_to_docx_bytes

    original = _equation_docx()
    graph = parse_docx_bytes(original, title="E")
    exported = graph_to_docx_bytes(graph)

    def equations(data: bytes) -> int:
        import io
        from docx import Document
        return len(Document(io.BytesIO(data)).element.body.findall(".//" + qn("m:oMath")))

    assert equations(original) == 1
    assert equations(exported) == 1, "the equation came back as OMML, not as text"


def test_a_rewritten_equation_is_written_as_the_words_it_became():
    """Asked to turn the equations into something readable, the assistant
    replaces the LaTeX — and then there is no equation left to restore."""
    import io
    from docx import Document
    from docx.oxml.ns import qn
    from app.docos.export import graph_to_docx_bytes

    graph = parse_docx_bytes(_equation_docx(), title="E")
    node = next(n for n in graph.nodes() if n.metadata.get("equations"))
    node.set_text("a over b")

    exported = Document(io.BytesIO(graph_to_docx_bytes(graph)))
    assert not exported.element.body.findall(".//" + qn("m:oMath"))
    assert any("a over b" in p.text for p in exported.paragraphs)


def test_the_converter_handles_the_shapes_maths_actually_takes():
    from docx.oxml import parse_xml
    from app.docos.parser.omml import omml_to_latex

    cases = {
        "<m:f><m:num><m:f><m:num><m:r><m:t>a</m:t></m:r></m:num>"
        "<m:den><m:r><m:t>b</m:t></m:r></m:den></m:f></m:num>"
        "<m:den><m:r><m:t>c</m:t></m:r></m:den></m:f>": r"\frac{\frac{a}{b}}{c}",

        '<m:nary><m:naryPr><m:chr m:val="∑"/></m:naryPr>'
        "<m:sub><m:r><m:t>i=1</m:t></m:r></m:sub><m:sup><m:r><m:t>N</m:t></m:r></m:sup>"
        "<m:e><m:r><m:t>x</m:t></m:r></m:e></m:nary>": r"\sum_{i=1}^N x",

        "<m:rad><m:deg/><m:e><m:r><m:t>x</m:t></m:r></m:e></m:rad>": r"\sqrt{x}",

        '<m:acc><m:accPr><m:chr m:val="̂"/></m:accPr>'
        "<m:e><m:r><m:t>y</m:t></m:r></m:e></m:acc>": r"\hat{y}",

        "<m:sSubSup><m:e><m:r><m:t>x</m:t></m:r></m:e>"
        "<m:sub><m:r><m:t>i</m:t></m:r></m:sub>"
        "<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSubSup>": "x_i^2",

        # Something the converter does not know: keep the words, lose the shape.
        "<m:groupChr><m:e><m:r><m:t>abc</m:t></m:r></m:e></m:groupChr>": "abc",
    }
    for inner, expected in cases.items():
        assert omml_to_latex(parse_xml(_omml(inner))) == expected, inner[:40]
