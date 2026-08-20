"""Importing a DOCX into the DocOS graph.

The failures these cover both changed what the reader saw: a horizontal rule
arriving as a figure, and a real figure arriving with no picture in it.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.docos.graph import NodeType  # noqa: E402
from app.docos.parser.docx_parser import parse_docx_bytes  # noqa: E402

_VML_RULE = (
    '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    '     xmlns:v="urn:schemas-microsoft-com:vml"'
    '     xmlns:o="urn:schemas-microsoft-com:office:office">'
    '  <w:pict><v:rect style="width:0;height:1.5pt" o:hralign="center"'
    '                  o:hrstd="t" o:hr="t"/></w:pict>'
    '</w:r>'
)


def _png(size=(240, 120), colour="#3366cc") -> io.BytesIO:
    buf = io.BytesIO()
    Image.new("RGB", size, colour).save(buf, format="PNG")
    buf.seek(0)
    return buf


def _add_rule(doc) -> None:
    doc.add_paragraph()._p.append(parse_xml(_VML_RULE))


def _render(doc) -> list:
    out = io.BytesIO()
    doc.save(out)
    return parse_docx_bytes(out.getvalue(), title="t").root.children


def _types(nodes) -> list[str]:
    return [n.type.value for n in nodes]


# ── horizontal rules ────────────────────────────────────────────────────────

def test_a_word_rule_is_a_rule_not_a_figure():
    """Word writes a rule as a VML rect inside w:pict — drawing markup, no picture."""
    doc = Document()
    doc.add_paragraph("before")
    _add_rule(doc)
    doc.add_paragraph("after")

    types = _types(_render(doc))
    assert NodeType.HORIZONTAL_RULE.value in types
    assert NodeType.FIGURE.value not in types


def test_a_bordered_empty_paragraph_is_still_a_rule():
    doc = Document()
    p = doc.add_paragraph()
    p._p.append(parse_xml(
        '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:pBdr><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>'
        '</w:pPr>'))
    assert NodeType.HORIZONTAL_RULE.value in _types(_render(doc))


# ── figures ─────────────────────────────────────────────────────────────────

def test_a_picture_arrives_with_its_bytes():
    doc = Document()
    doc.add_picture(_png(), width=Inches(2))

    figures = [n for n in _render(doc) if n.type == NodeType.FIGURE]
    assert len(figures) == 1
    image = figures[0].children[0]
    assert image.type == NodeType.IMAGE
    src = image.metadata.get("src", "")
    assert src.startswith("data:image/"), "the picture must carry its own bytes"
    assert image.metadata["bytes"] > 0


def test_several_pictures_in_one_paragraph_each_get_a_node():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run().add_picture(_png(), width=Inches(1))
    p.add_run().add_picture(_png(colour="#cc3333"), width=Inches(1))

    figures = [n for n in _render(doc) if n.type == NodeType.FIGURE]
    assert len(figures) == 1
    assert len(figures[0].children) == 2
    assert all(c.metadata.get("src", "").startswith("data:image/")
               for c in figures[0].children)


def test_rules_and_figures_survive_the_same_document_in_order():
    doc = Document()
    doc.add_paragraph("intro")
    _add_rule(doc)
    doc.add_picture(_png(), width=Inches(2))
    doc.add_paragraph("outro")

    types = _types(_render(doc))
    assert types.index(NodeType.HORIZONTAL_RULE.value) < types.index(NodeType.FIGURE.value)
    assert types.count(NodeType.FIGURE.value) == 1
    assert types.count(NodeType.HORIZONTAL_RULE.value) == 1


def test_an_oversized_picture_keeps_its_place_without_the_bytes(monkeypatch):
    """The document's shape must survive an image we decline to inline."""
    import app.docos.parser.docx_parser as parser
    monkeypatch.setattr(parser, "_MAX_INLINE_IMAGE_BYTES", 10)

    doc = Document()
    doc.add_picture(_png(), width=Inches(2))

    figures = [n for n in _render(doc) if n.type == NodeType.FIGURE]
    image = figures[0].children[0]
    assert "src" not in image.metadata
    assert image.metadata["too_large"] > 10


# ── things that must not be mistaken for pictures ───────────────────────────

def test_an_empty_paragraph_is_still_dropped():
    doc = Document()
    doc.add_paragraph("a")
    doc.add_paragraph("   ")
    doc.add_paragraph("b")
    assert _types(_render(doc)).count(NodeType.BODY.value) == 2


@pytest.mark.parametrize("text", ["a caption", "Figure 1. Something"])
def test_text_next_to_a_picture_becomes_its_caption(text):
    doc = Document()
    p = doc.add_paragraph(text)
    p.add_run().add_picture(_png(), width=Inches(1))

    figures = [n for n in _render(doc) if n.type == NodeType.FIGURE]
    assert figures[0].children[0].content == text


# ── pictures laid out in a table ────────────────────────────────────────────

def test_a_picture_in_a_table_cell_is_not_dropped():
    """Putting screenshots in a table is a common layout; the cell text alone
    was all that survived, so the pictures vanished."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Screenshot 1"
    table.rows[0].cells[1].paragraphs[0].add_run().add_picture(_png(), width=Inches(1))

    graph_nodes = _render(doc)
    images = [n for node in graph_nodes for n in node.walk() if n.type == NodeType.IMAGE]
    assert len(images) == 1
    assert images[0].metadata.get("src", "").startswith("data:image/")


def test_cell_text_still_survives_alongside_its_picture():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = "Figure 1"
    cell.paragraphs[0].add_run().add_picture(_png(), width=Inches(1))

    table_node = next(n for n in _render(doc) if n.type == NodeType.TABLE)
    cell_node = table_node.children[0].children[0]
    assert "Figure 1" in cell_node.content
    assert len(cell_node.children) == 1


# ── linked vs embedded pictures ─────────────────────────────────────────────

def _link_image(doc, target: str) -> None:
    """A picture the document points at rather than stores — what LibreOffice
    writes converting HTML, and Word writes for "Link to File"."""
    rid = doc.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
        is_external=True,
    )
    p = doc.add_paragraph()
    p._p.append(parse_xml(
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        '     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        '     xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        '     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        '     xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<w:drawing><wp:inline><a:graphic><a:graphicData>'
        '<pic:pic><pic:blipFill>'
        f'<a:blip r:link="{rid}"/>'
        '</pic:blipFill></pic:pic>'
        '</a:graphicData></a:graphic></wp:inline></w:drawing></w:r>'))


def _first_image(nodes):
    return next(n for node in nodes for n in node.walk() if n.type == NodeType.IMAGE)


def test_a_locally_linked_picture_is_reported_as_linked_not_unreadable():
    """The bytes are genuinely absent; saying "could not be read" blamed the
    importer for something the document never contained."""
    doc = Document()
    _link_image(doc, "file:///C:/pictures/diagram.png")

    image = _first_image(_render(doc))
    assert image.metadata["external"] is True
    assert image.metadata["linked_to"] == "file:///C:/pictures/diagram.png"
    assert "src" not in image.metadata
    assert "unreadable" not in image.metadata


def test_a_local_path_is_never_read_from_disk():
    """An imported document must not be able to make us open arbitrary files."""
    doc = Document()
    _link_image(doc, "file:///C:/Windows/win.ini")

    image = _first_image(_render(doc))
    assert "src" not in image.metadata, "a named local file must not be inlined"


def test_an_http_linked_picture_is_passed_through():
    doc = Document()
    _link_image(doc, "https://example.com/figure.png")

    image = _first_image(_render(doc))
    assert image.metadata["src"] == "https://example.com/figure.png"
    assert image.metadata["external"] is True


def test_an_embedded_picture_is_still_inlined():
    doc = Document()
    doc.add_picture(_png(), width=Inches(1))

    image = _first_image(_render(doc))
    assert image.metadata["src"].startswith("data:image/")
    assert "external" not in image.metadata
