"""Custom, user-defined stylesheets: JSON round-trip, storage, resolution,
rendering, and derivation from a reference DOCX."""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.shared import Pt

from app.paper.renderer import render_paper
from app.paper.schema import PaperSpec, Style
from app.paper.styles import get_stylesheet, list_styles, resolve_style
from app.paper.styles.base import StyleSheet
from app.paper.styles.store import StyleStore

SPEC = {
    "meta": {"title": "Custom Styled Doc", "abstract": "An abstract.", "keywords": ["a"]},
    "blocks": [
        {"type": "heading", "level": 1, "text": "Introduction"},
        {"type": "paragraph", "text": "Body text for the custom style test."},
        {"type": "table", "caption": "Data", "columns": ["A"], "rows": [["1"]]},
    ],
    "references": ["A. Author, Title, 2024."],
}


@pytest.fixture()
def store(tmp_path):
    return StyleStore(db_path=tmp_path / "styles.db")


# ── serialisability (the property that makes custom styles possible) ────────

def test_stylesheet_json_round_trip():
    sheet = get_stylesheet("ieee")
    data = sheet.model_dump(mode="json")
    restored = StyleSheet.model_validate(data)
    assert restored.id == "ieee"
    assert restored.page.columns == 2
    assert restored.body.size_pt == sheet.body.size_pt
    assert restored.heading1.small_caps is True


# ── storage + resolution ────────────────────────────────────────────────────

def test_save_and_resolve_custom_style(store, monkeypatch):
    import app.paper.styles.store as store_mod
    monkeypatch.setattr(store_mod, "get_style_store", lambda: store)

    custom = get_stylesheet("assignment").model_copy(deep=True)
    custom.name = "House Style"
    custom.body = custom.body.merged(Style(font="Georgia", size_pt=13))
    saved = store.save("usr_1", custom)

    assert saved.builtin is False
    assert saved.id.startswith("style_")

    # resolvable by id and by name, scoped to the owner
    assert resolve_style(saved.id, "usr_1").body.font == "Georgia"
    assert resolve_style("House Style", "usr_1").body.size_pt == 13
    # another user cannot see it → falls back to the default style
    assert resolve_style(saved.id, "usr_2").id == "ieee"


def test_a_custom_style_beats_a_builtin_alias(store, monkeypatch):
    """"Report" is one of the loose names pointing at the assignment sheet. A
    user who names their own style "Report" means their own style — an alias
    must not shadow it. A built-in's own id still wins, since it is ours."""
    import app.paper.styles.store as store_mod
    monkeypatch.setattr(store_mod, "get_style_store", lambda: store)

    mine = get_stylesheet("ieee").model_copy(deep=True)
    mine.name = "Report"
    mine.body = mine.body.merged(Style(font="Georgia"))
    store.save("usr_1", mine)

    assert resolve_style("Report", "usr_1").body.font == "Georgia"
    assert resolve_style("Report", "usr_2").id == "assignment"   # alias, for everyone else
    assert resolve_style("assignment", "usr_1").id == "assignment"  # an id cannot be claimed


def test_custom_style_appears_in_listing(store, monkeypatch):
    import app.paper.styles.store as store_mod
    monkeypatch.setattr(store_mod, "get_style_store", lambda: store)

    sheet = get_stylesheet("ieee").model_copy(deep=True)
    sheet.name = "My IEEE Variant"
    store.save("usr_1", sheet)

    listed = list_styles(owner_id="usr_1")
    ids = {s["name"] for s in listed}
    assert "My IEEE Variant" in ids
    assert "IEEE Conference" in ids          # built-ins still present
    assert len(list_styles()) == 3           # built-ins only without an owner


def test_delete_custom_style(store):
    sheet = get_stylesheet("report").model_copy(deep=True)
    sheet.name = "Temp"
    saved = store.save("usr_1", sheet)
    assert store.delete(saved.id, "usr_1") is True
    assert store.get(saved.id, "usr_1") is None
    assert store.delete(saved.id, "usr_1") is False


# ── rendering with a custom sheet ───────────────────────────────────────────

def test_render_with_custom_stylesheet_object(tmp_path):
    custom = get_stylesheet("assignment").model_copy(deep=True)
    custom.id = "house"
    custom.name = "House"
    custom.builtin = False
    custom.body = custom.body.merged(Style(font="Georgia", size_pt=13))
    custom.heading_scheme = "roman_alpha"

    out = tmp_path / "custom.docx"
    render_paper(PaperSpec.model_validate(SPEC), out, style=custom)

    doc = Document(str(out))
    paras = [p for p in doc.paragraphs if p.text.strip()]
    body = next(p for p in paras if p.text.startswith("Body text"))
    assert body.runs[0].font.name == "Georgia"
    assert body.runs[0].font.size == Pt(13)
    # the custom sheet's numbering scheme took effect
    assert any(p.text.startswith("I. Introduction") for p in paras)



def _reference_docx() -> bytes:
    doc = Document()
    h = doc.add_heading("A Heading", level=1)
    for run in h.runs:
        run.font.name = "Garamond"
        run.font.size = Pt(20)

    for _ in range(3):  # majority-vote needs a few samples
        p = doc.add_paragraph()
        run = p.add_run("Body paragraph of the reference document, long enough to matter.")
        run.font.name = "Garamond"
        run.font.size = Pt(13)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
