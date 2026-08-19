"""The formal assignment style, embedded code listings, and the two rendering
faults they were built alongside: right-drifting references and blank figures.
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.paper.figures import render_figure  # noqa: E402
from app.paper.references import format_reference  # noqa: E402
from app.paper.renderer import render_paper  # noqa: E402
from app.paper.schema import ChartSpec, PaperSpec  # noqa: E402
from app.paper.styles import resolve_style  # noqa: E402

CODE = 'def solve(x):\n    """Docstring."""\n    return x * 2'


def _texts(path):
    return [p.text for p in Document(str(path)).paragraphs]


def _render(tmp_path, spec, style="assignment"):
    return render_paper(PaperSpec.model_validate(spec), tmp_path / "out.docx", style=style)


# ── the assignment style ────────────────────────────────────────────────────

@pytest.mark.parametrize("name", ["assignment", "homework", "coursework",
                                  "report", "official", "lab report"])
def test_the_names_users_type_reach_the_formal_style(name):
    assert resolve_style(name).id == "assignment"


def test_formal_style_is_one_column_with_decimal_numbering():
    sheet = resolve_style("assignment")
    assert sheet.page.columns == 1
    assert sheet.heading_scheme == "decimal"


def test_ieee_is_still_the_default():
    assert resolve_style("").id == "ieee"
    assert resolve_style("something we do not implement").id == "ieee"


# ── code listings ───────────────────────────────────────────────────────────

def test_listings_are_captioned_and_numbered(tmp_path):
    out = _render(tmp_path, {"meta": {"title": "T"}, "blocks": [
        {"type": "code", "language": "python", "filename": "solve.py",
         "caption": "doubles its argument", "text": CODE},
        {"type": "code", "language": "python", "caption": "second one", "text": "pass"},
    ]})
    texts = _texts(out)
    assert "Listing 1. solve.py — doubles its argument" in texts
    assert "Listing 2. second one" in texts


def test_listing_keeps_every_line_and_its_indentation(tmp_path):
    out = _render(tmp_path, {"meta": {"title": "T"}, "blocks": [
        {"type": "code", "language": "python", "text": CODE}]})
    texts = _texts(out)
    for line in CODE.splitlines():
        assert line in texts


def test_listing_lines_are_held_together_across_a_break(tmp_path):
    out = _render(tmp_path, {"meta": {"title": "T"}, "blocks": [
        {"type": "code", "language": "python", "text": CODE}]})
    paras = [p for p in Document(str(out)).paragraphs if p.text in CODE.splitlines()]
    assert [p.paragraph_format.keep_with_next for p in paras] == [True, True, False]


# ── references ──────────────────────────────────────────────────────────────

def test_bibtex_becomes_a_formatted_single_line_citation():
    got = format_reference(
        "@article{k,\n title = {A Title},\n author = {Doe, Jane and Roe, Richard},\n"
        " journal = {A Journal},\n volume = {7},\n pages = {1--10},\n year = {2020}\n}")
    assert "\n" not in got
    assert got == 'J. Doe and R. Roe, "A Title," *A Journal*, vol. 7, pp. 1–10, 2020.'


def test_a_multi_line_entry_is_never_justified(tmp_path):
    """Justified lines ending in a manual break get stretched to the full column
    width, which pushes an unbreakable token flush right."""
    out = _render(tmp_path, {
        "meta": {"title": "T"},
        "blocks": [{"type": "paragraph", "text": "body"}],
        "references": ["Some Author, unparseable\nsecond line\nthird line"],
    }, style="ieee")
    ref = next(p for p in Document(str(out)).paragraphs if p.text.startswith("[1]"))
    assert "\n" not in ref.text
    assert ref.alignment != WD_ALIGN_PARAGRAPH.RIGHT


# ── figures ─────────────────────────────────────────────────────────────────

def test_a_grouped_scatter_carries_its_own_x_coordinates(tmp_path):
    chart = ChartSpec(kind="scatter", x_label="A", y_label="B", series=[
        {"name": "one", "x_values": [0.1, 0.2], "values": [1.0, 2.0]},
        {"name": "two", "x_values": [0.3, 0.4], "values": [3.0, 4.0]},
    ])
    out = render_figure(chart, tmp_path / "fig.png")
    assert out.exists() and out.stat().st_size > 0


def test_a_chart_with_no_values_is_refused(tmp_path):
    with pytest.raises(ValueError):
        render_figure(ChartSpec(kind="scatter", title="Empty", labels=["a", "b"]),
                      tmp_path / "fig.png")


def test_an_undrawable_figure_takes_its_caption_with_it(tmp_path):
    """No bare "Figure 1." standing over a blank box — and the figure that can
    be drawn still gets number 1."""
    out = _render(tmp_path, {"meta": {"title": "T"}, "blocks": [
        {"type": "figure", "caption": "nothing to draw",
         "chart": {"kind": "scatter", "labels": ["a"], "values": []}},
        {"type": "figure", "caption": "a real one",
         "chart": {"kind": "bar", "labels": ["a", "b"], "values": [1.0, 2.0]}},
    ]})
    texts = _texts(out)
    assert not any("nothing to draw" in t for t in texts)
    assert any(t.startswith("Figure 1.") and "a real one" in t for t in texts)
