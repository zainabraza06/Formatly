"""Typeset equations and code listings, neither of which needs LaTeX or an IDE."""
from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.paper.codeshot import THEMES, render_code_image  # noqa: E402
from app.paper.equations import looks_like_math, render_equation_png  # noqa: E402
from app.paper.renderer import render_paper  # noqa: E402
from app.paper.schema import PaperSpec  # noqa: E402

_CODE = 'import numpy as np\n\n\ndef f(x):\n    """Doc."""\n    return np.sqrt(x)\n'


def _images(path) -> int:
    doc = Document(str(path))
    return sum(1 for r in doc.part.rels.values() if "image" in r.reltype)


def _texts(path) -> list[str]:
    return [p.text for p in Document(str(path)).paragraphs]


# ── equations ───────────────────────────────────────────────────────────────

@pytest.mark.parametrize("expr", [
    r"\frac{a}{b}",
    r"\sum_{i=0}^{n} x_i",
    r"\sqrt{s_1^2 + s_2^2}",
    r"\alpha + \beta",
    "x_1",
])
def test_math_markup_is_recognised(expr):
    assert looks_like_math(expr)


@pytest.mark.parametrize("expr", ["y = mx + c", "The result is significant", ""])
def test_plain_expressions_are_left_as_text(expr):
    assert not looks_like_math(expr)


def test_renders_a_fraction_without_latex(tmp_path):
    out, width_in = render_equation_png(r"\frac{a}{b}", tmp_path / "eq.png", size_pt=10)
    assert out.exists()
    assert 0 < width_in < 4, f"implausible width {width_in}in"


@pytest.mark.parametrize("wrapper", ["${}$", "$${}$$", r"\[{}\]", r"\({}\)", "{}"])
def test_accepts_the_delimiters_models_actually_emit(wrapper, tmp_path):
    out, _ = render_equation_png(wrapper.format(r"\frac{a}{b}"), tmp_path / "eq.png")
    assert out.exists()


def test_unparseable_markup_raises_rather_than_writing_a_broken_image(tmp_path):
    with pytest.raises(ValueError):
        render_equation_png(r"y = \frac{a}{", tmp_path / "eq.png")


def test_document_typesets_maths_and_keeps_plain_equations_as_text(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "ieee"},
        "blocks": [
            {"type": "equation", "text": r"d = \frac{x_1 - x_2}{\sqrt{s}}"},
            {"type": "equation", "text": "y = mx + c"},
        ],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    assert _images(out) == 1, "the fraction should be an image, the plain one text"
    assert any("y = mx + c" in t for t in _texts(out))


def test_a_broken_equation_costs_its_typography_not_the_document(tmp_path):
    """The markup still has to reach the reader, even unparsed."""
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "ieee"},
        "blocks": [{"type": "equation", "text": r"z = \frac{a}{"}],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    assert _images(out) == 0
    assert any(r"\frac{a}{" in t for t in _texts(out))


# ── code listings ───────────────────────────────────────────────────────────

@pytest.mark.parametrize("theme", sorted(THEMES))
def test_renders_an_editor_screenshot(theme, tmp_path):
    out = render_code_image(_CODE, tmp_path / f"{theme}.png", language="python",
                            filename="f.py", theme=theme)
    assert out.exists() and out.stat().st_size > 1000


def test_unknown_language_still_renders(tmp_path):
    out = render_code_image("SELECT 1;", tmp_path / "c.png", language="not-a-language")
    assert out.exists()


def test_empty_code_is_refused(tmp_path):
    with pytest.raises(ValueError):
        render_code_image("   ", tmp_path / "c.png")


def test_code_renders_as_text_by_default_and_as_an_image_on_request(tmp_path):
    def build(render):
        spec = PaperSpec.model_validate({
            "meta": {"title": "T", "style": "ieee_1col"},
            "blocks": [{"type": "code", "language": "python", "text": _CODE,
                        "render": render, "filename": "f.py"}],
        })
        return render_paper(spec, tmp_path / f"{render}.docx")

    as_text = build("text")
    assert _images(as_text) == 0
    assert any("import numpy as np" in t for t in _texts(as_text))

    as_image = build("image")
    assert _images(as_image) == 1
    assert not any("import numpy as np" in t for t in _texts(as_image))


def test_listing_caption_carries_the_filename(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "ieee_1col"},
        "blocks": [{"type": "code", "language": "python", "text": _CODE,
                    "filename": "regime.py", "caption": "the scalars"}],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    caption = next(t for t in _texts(out) if t.startswith("Listing"))
    assert "regime.py" in caption and "the scalars" in caption


# ── terminal (program output) screenshots ───────────────────────────────────

from app.paper.codeshot import TERMINAL_THEMES, render_terminal_image  # noqa: E402
from app.paper.renderer import _is_terminal  # noqa: E402
from app.paper.schema import Code  # noqa: E402

_SESSION = "Enter your guess: 50\nToo high!\nEnter your guess: 37\nCorrect!\n"


@pytest.mark.parametrize("theme", sorted(TERMINAL_THEMES))
def test_renders_a_console_window(theme, tmp_path):
    out = render_terminal_image(_SESSION, tmp_path / f"{theme}.png",
                                title="Command Prompt", theme=theme)
    assert out.exists() and out.stat().st_size > 500


def test_empty_output_is_refused(tmp_path):
    with pytest.raises(ValueError):
        render_terminal_image("   ", tmp_path / "t.png")


def test_terminal_is_chosen_when_asked_for():
    assert _is_terminal(Code(text="x", window="terminal", language="cpp"))


def test_source_listings_stay_in_the_editor():
    assert not _is_terminal(Code(text="x", language="cpp", caption="the game loop"))
    assert not _is_terminal(Code(text="x", language="python", caption="sample output"))


def test_a_plain_text_output_listing_is_inferred_as_a_console():
    """Models label output as language "text"; a caption saying so is enough."""
    assert _is_terminal(Code(text="x", language="text", caption="Sample output: startup"))
    assert _is_terminal(Code(text="x", language="", caption="Program execution"))
    assert not _is_terminal(Code(text="x", language="text", caption="Pseudocode"))


def test_both_screenshot_kinds_render_into_one_document(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment"},
        "blocks": [
            {"type": "code", "language": "cpp", "text": "int main(){}",
             "render": "image", "window": "editor", "filename": "g.cpp"},
            {"type": "code", "language": "text", "text": _SESSION,
             "render": "image", "window": "terminal", "filename": "Command Prompt"},
        ],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    assert _images(out) == 2
    # neither listing should also appear as selectable text
    assert not any("Enter your guess" in t for t in _texts(out))
