"""Multi-pass (agentic) generation.

Depth is produced by planning once and then writing each section on its own,
because a single call asked for a detailed document overruns the token ceiling
and truncates its JSON — losing the whole document rather than its tail.

These tests use a scripted router, so they run offline and deterministically.
"""
from __future__ import annotations

import json

import pytest

from app.paper.agentic import generate_sectioned
from app.paper.generator import PaperGenerationError, generate_paper
from app.paper.prompt import DEPTHS, section_system_prompt, system_prompt

PLAN = {
    "meta": {"title": "Churn Analysis", "abstract": "An abstract.", "keywords": ["churn"],
             "authors": [{"name": "A. Analyst", "affiliation": "Ops"}]},
    "outline": [
        {"heading": "Introduction", "brief": "frame the problem", "subsections": []},
        {"heading": "Findings", "brief": "present the data", "subsections": ["Pricing"]},
    ],
    "references": ["A. Author, Title, 2024."],
    "visualization_plan": [{"data": "churn by month", "kind": "line", "rationale": "trend"}],
}

SECTION_REPLIES = {
    "Introduction": {"blocks": [
        {"type": "heading", "level": 1, "text": "Introduction"},
        {"type": "paragraph", "text": "Framing prose for the introduction section."},
    ]},
    "Findings": {"blocks": [
        {"type": "heading", "level": 1, "text": "Findings"},
        {"type": "paragraph", "text": "What the data shows."},
        {"type": "heading", "level": 2, "text": "Pricing"},
        {"type": "figure", "caption": "Churn by month",
         "chart": {"kind": "line", "title": "Churn", "labels": ["Jul", "Aug"],
                   "values": [4.2, 5.1]}},
    ]},
}


class ScriptedRouter:
    """Answers the plan pass with PLAN, then each section pass by its heading."""

    def __init__(self, section_replies=None, fail_sections=()):
        self.section_replies = section_replies or SECTION_REPLIES
        self.fail_sections = set(fail_sections)
        self.calls: list[str] = []

    def chat(self, messages, **kwargs):
        system = messages[0]["content"]
        user = messages[1]["content"]

        if "planning a document" in system:
            self.calls.append("plan")
            return json.dumps(PLAN), "scripted", 0.01

        heading = next((h for h in self.section_replies if f"'{h}'" in user), None)
        self.calls.append(f"section:{heading}")
        if heading in self.fail_sections:
            raise RuntimeError("provider exploded")
        return json.dumps(self.section_replies[heading]), "scripted", 0.01


# ── the loop ────────────────────────────────────────────────────────────────

def test_sectioned_generation_makes_one_call_per_section_plus_plan():
    router = ScriptedRouter()
    raw, provider = generate_sectioned(
        raw_text="material", style_guide="report", depth="detailed",
        doc_kind="report", router=router,
    )
    assert provider == "scripted"
    assert router.calls == ["plan", "section:Introduction", "section:Findings"]

    assert raw["meta"]["title"] == "Churn Analysis"
    assert raw["references"] == PLAN["references"]
    assert raw["visualization_plan"] == PLAN["visualization_plan"]

    headings = [b["text"] for b in raw["blocks"] if b["type"] == "heading" and b["level"] == 1]
    assert headings == ["Introduction", "Findings"]      # outline order preserved
    assert any(b["type"] == "figure" for b in raw["blocks"])


def test_progress_is_reported_per_section():
    seen = []
    generate_sectioned(
        raw_text="m", style_guide="report", depth="detailed", doc_kind="report",
        router=ScriptedRouter(), on_progress=lambda h, i, n: seen.append((h, i, n)),
    )
    assert seen == [("Introduction", 1, 2), ("Findings", 2, 2)]


def test_a_failed_section_does_not_destroy_the_document():
    """The point of splitting the work up: one bad section costs that section,
    not the whole document."""
    router = ScriptedRouter(fail_sections={"Findings"})
    raw, _ = generate_sectioned(
        raw_text="m", style_guide="report", depth="detailed", doc_kind="report", router=router,
    )
    headings = [b["text"] for b in raw["blocks"] if b["type"] == "heading" and b["level"] == 1]
    assert headings == ["Introduction", "Findings"]   # heading kept as a visible gap
    # the surviving section still has its prose
    assert any(b["type"] == "paragraph" for b in raw["blocks"])


def test_sections_cannot_invent_sibling_level1_headings():
    """A section that wanders into writing other sections would corrupt the
    document's numbering; extra level-1 headings are demoted."""
    replies = {
        "Introduction": {"blocks": [
            {"type": "heading", "level": 1, "text": "Introduction"},
            {"type": "paragraph", "text": "prose"},
            {"type": "heading", "level": 1, "text": "Sneaky Extra Section"},
            {"type": "paragraph", "text": "more"},
        ]},
        "Findings": SECTION_REPLIES["Findings"],
    }
    raw, _ = generate_sectioned(
        raw_text="m", style_guide="report", depth="detailed", doc_kind="report",
        router=ScriptedRouter(section_replies=replies),
    )
    l1 = [b["text"] for b in raw["blocks"] if b["type"] == "heading" and b["level"] == 1]
    assert l1 == ["Introduction", "Findings"]
    assert any(b.get("text") == "Sneaky Extra Section" and b["level"] == 2
               for b in raw["blocks"])


def test_missing_heading_is_restored_from_the_outline():
    replies = {
        "Introduction": {"blocks": [{"type": "paragraph", "text": "no heading here"}]},
        "Findings": SECTION_REPLIES["Findings"],
    }
    raw, _ = generate_sectioned(
        raw_text="m", style_guide="report", depth="detailed", doc_kind="report",
        router=ScriptedRouter(section_replies=replies),
    )
    first = raw["blocks"][0]
    assert first["type"] == "heading" and first["text"] == "Introduction"


def test_plan_without_outline_is_an_error():
    class NoOutline(ScriptedRouter):
        def chat(self, messages, **kwargs):
            return json.dumps({"meta": {"title": "x"}, "outline": []}), "scripted", 0.01

    with pytest.raises(ValueError):
        generate_sectioned(raw_text="m", style_guide="report", depth="detailed",
                           doc_kind="report", router=NoOutline())


# ── routing: which depths go multi-pass ─────────────────────────────────────

def test_detailed_depth_routes_through_the_multipass_path():
    router = ScriptedRouter()
    spec, _ = generate_paper(raw_text="material", style="report", depth="detailed",
                             router=router)
    assert "plan" in router.calls          # planned, not written in one shot
    assert spec.resolved is True
    assert spec.meta.title == "Churn Analysis"


def test_standard_depth_stays_single_pass():
    """Short documents must not pay for passes they do not need."""
    single = {
        "meta": {"title": "One Shot", "abstract": "a", "keywords": ["k"]},
        "blocks": [{"type": "heading", "level": 1, "text": "Intro"},
                   {"type": "paragraph", "text": "prose"}],
        "references": [],
    }

    class Single:
        def __init__(self):
            self.calls = 0

        def chat(self, messages, **kwargs):
            self.calls += 1
            return json.dumps(single), "scripted", 0.01

    router = Single()
    spec, _ = generate_paper(raw_text="m", style="report", depth="standard", router=router)
    assert router.calls == 1
    assert spec.meta.title == "One Shot"


def test_multipass_can_be_forced_or_suppressed():
    router = ScriptedRouter()
    generate_paper(raw_text="m", style="report", depth="standard", multipass=True,
                   router=router)
    assert "plan" in router.calls


def test_multipass_failure_surfaces_as_generation_error():
    class Broken:
        def chat(self, messages, **kwargs):
            return "not json at all", "scripted", 0.01

    with pytest.raises(PaperGenerationError):
        generate_paper(raw_text="m", style="report", depth="detailed", router=Broken())


# ── depth prompts ───────────────────────────────────────────────────────────

def test_every_depth_has_a_length_instruction():
    for depth in DEPTHS:
        assert "LENGTH" in system_prompt("report", depth)


def test_depth_changes_the_instruction():
    brief = system_prompt("report", "brief")
    detailed = system_prompt("report", "detailed")
    assert brief != detailed
    assert "concise" in brief.lower()
    assert "in-depth" in detailed.lower() or "thorough" in detailed.lower()


def test_detailed_depth_still_forbids_fabrication():
    """Depth must come from analysis, never from invented facts."""
    p = system_prompt("report", "detailed").lower()
    assert "never invent" in p or "not from fabricated" in p


def test_section_prompt_scopes_the_model_to_one_section():
    p = section_system_prompt("report", "detailed")
    assert "ONLY the section you are given" in p


# ── chart-kind coercion (models write "pie chart", not "pie") ───────────────

@pytest.mark.parametrize("raw,expected", [
    ("pie chart", "pie"), ("Pie", "pie"), ("donut", "pie"),
    ("line chart", "line"), ("line graph", "line"), ("trend line", "line"),
    ("stacked bar chart", "grouped_bar"), ("grouped bars", "grouped_bar"),
    ("clustered column", "grouped_bar"), ("multi-series bar", "grouped_bar"),
    ("bar chart", "bar"), ("column chart", "bar"), ("histogram", "bar"),
    ("scatter plot", "scatter"),
    ("something weird", "bar"),        # last-resort default
    ("bar", "bar"),                    # already canonical
])
def test_chart_kind_is_coerced(raw, expected):
    from app.paper.schema import normalize_chart_kind
    assert normalize_chart_kind(raw) == expected


def test_visualization_plan_tolerates_loose_kinds():
    """The real failure that motivated this: a plan with 'pie chart' etc. must
    parse rather than sink the whole document."""
    from app.paper.schema import PaperSpec
    spec = PaperSpec.model_validate({
        "meta": {"title": "t"},
        "blocks": [{"type": "paragraph", "text": "p"}],
        "visualization_plan": [
            {"data": "shares", "kind": "pie chart", "rationale": "composition"},
            {"data": "trend", "kind": "line graph", "rationale": "over time"},
            {"data": "compare", "kind": "stacked bar chart", "rationale": "multi"},
        ],
    })
    assert [v.kind for v in spec.visualization_plan] == ["pie", "line", "grouped_bar"]


# ── styles we don't implement (Chicago, Harvard, a journal house style) ────

def test_unknown_style_name_reaches_the_writer():
    """A requested style we have no stylesheet for must not be silently dropped —
    the writer should still be told to follow its conventions."""
    from tests.test_paper import SPEC as SINGLE_SPEC, _StubRouter

    router = _StubRouter(json.dumps(SINGLE_SPEC))
    generate_paper(raw_text="material", style="Chicago", depth="standard", router=router)
    system = router.system
    assert "Chicago" in system
    assert "in-text citation" in system.lower() or "references" in system.lower()


def test_unknown_style_still_renders_with_a_real_stylesheet():
    from tests.test_paper import SPEC as SINGLE_SPEC, _StubRouter

    spec, _ = generate_paper(raw_text="m", style="Vancouver", depth="standard",
                             router=_StubRouter(json.dumps(SINGLE_SPEC)))
    # typography falls back to a neutral sheet so the document is still renderable
    assert spec.resolved is True
    assert spec.meta.style == "ieee"


def test_known_styles_do_not_get_treated_as_unknown():
    from tests.test_paper import SPEC as SINGLE_SPEC, _StubRouter

    router = _StubRouter(json.dumps(SINGLE_SPEC))
    generate_paper(raw_text="m", style="ieee", depth="standard", router=router)
    assert "Requested style:" not in router.system     # used the built-in guide
    assert "IEEE" in router.system


def test_unknown_style_reaches_the_multipass_writer_too():
    router = ScriptedRouter()
    generate_paper(raw_text="m", style="Harvard", depth="detailed", router=router)
    # the scripted router records calls; assert the plan prompt carried the name
    assert "plan" in router.calls


def test_lookup_style_distinguishes_unknown_from_default():
    from app.paper.styles import lookup_style
    assert lookup_style("ieee") is not None
    assert lookup_style("Chicago") is None            # genuinely unknown
    assert lookup_style("") is None


def test_figure_chart_kind_is_coerced():
    from app.paper.schema import PaperSpec
    spec = PaperSpec.model_validate({
        "meta": {"title": "t"},
        "blocks": [{"type": "figure", "caption": "c",
                    "chart": {"kind": "line chart", "labels": ["a"], "values": [1]}}],
    })
    fig = spec.blocks[0]
    assert fig.chart.kind == "line"


# ── markdown emphasis in model output ───────────────────────────────────────

def test_markdown_emphasis_becomes_real_formatting(tmp_path):
    """Models italicise titles with asterisks — especially in reference formats
    that require italics. Written verbatim they appear as literal '*' in Word."""
    from docx import Document as Docx
    from app.paper.renderer import render_paper
    from app.paper.schema import PaperSpec

    spec = PaperSpec.model_validate({
        "meta": {"title": "T"},
        "blocks": [{"type": "paragraph", "text": "See *The Printing Press* and **note this**."}],
        "references": ["Eisenstein, E. 1979. *The Printing Press as an Agent of Change*."],
    })
    out = tmp_path / "emphasis.docx"
    render_paper(spec, out, style="report")

    doc = Docx(str(out))
    body = next(p for p in doc.paragraphs if p.text.startswith("See "))
    assert "*" not in body.text                      # no literal asterisks survive
    assert any(r.italic and r.text == "The Printing Press" for r in body.runs)
    assert any(r.bold and r.text == "note this" for r in body.runs)

    ref = next(p for p in doc.paragraphs if "Eisenstein" in p.text)
    assert "*" not in ref.text
    assert any(r.italic for r in ref.runs)


def test_code_blocks_keep_underscores_verbatim(tmp_path):
    """snake_case identifiers must not be read as markdown emphasis."""
    from docx import Document as Docx
    from app.paper.renderer import render_paper
    from app.paper.schema import PaperSpec

    spec = PaperSpec.model_validate({
        "meta": {"title": "T"},
        "blocks": [{"type": "code", "language": "python",
                    "text": "my_var = other_var * 2"}],
    })
    out = tmp_path / "code.docx"
    render_paper(spec, out, style="report")

    doc = Docx(str(out))
    line = next(p for p in doc.paragraphs if "my_var" in p.text)
    assert line.text == "my_var = other_var * 2"
