"""Layout a user can ask for: a cover sheet, a page break, per-block formatting.

A formatting instruction can only be obeyed if the schema can express it. These
cover the primitives that make open-ended requests answerable rather than
needing a new rule per phrasing.
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.paper.renderer import render_paper  # noqa: E402
from app.paper.schema import PaperSpec  # noqa: E402

_BODY = [{"type": "heading", "level": 1, "text": "Introduction"},
         {"type": "paragraph", "text": "Body text."}]


def _breaks(path) -> int:
    """Explicit page breaks in the document."""
    from docx.oxml.ns import qn
    doc = Document(str(path))
    return sum(1 for p in doc.paragraphs
               for br in p._p.iter(qn("w:br"))
               if br.get(qn("w:type")) == "page")


def _texts(path) -> list[str]:
    return [p.text for p in Document(str(path)).paragraphs if p.text.strip()]


# ── cover sheet ─────────────────────────────────────────────────────────────

def test_no_cover_sheet_unless_asked(tmp_path):
    """A paper must not suddenly grow a title page."""
    spec = PaperSpec.model_validate({
        "meta": {"title": "A Paper", "style": "ieee"}, "blocks": _BODY})
    out = render_paper(spec, tmp_path / "d.docx")
    assert _breaks(out) == 0


def test_cover_sheet_holds_only_what_was_named(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "Number Guessing Game", "style": "assignment",
                 "title_page": True,
                 "title_page_lines": ["Student ID: 503840", "Course: FOCP"],
                 "authors": [{"name": "zainab", "affiliation": "NUST"}]},
        "blocks": _BODY,
    })
    out = render_paper(spec, tmp_path / "d.docx")
    texts = _texts(out)
    cover = texts[:texts.index("1. Introduction")] if "1. Introduction" in texts else texts[:5]
    assert "Number Guessing Game" in cover
    assert "zainab" in cover and "NUST" in cover
    assert "Student ID: 503840" in cover and "Course: FOCP" in cover
    assert _breaks(out) == 1, "the document should start on the page after the cover"


def test_cover_sheet_keeps_an_abstract_off_the_cover(tmp_path):
    """An abstract the brief did ask for belongs after the cover, not on it."""
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment", "title_page": True,
                 "abstract": "A summary that was asked for."},
        "blocks": _BODY,
    })
    out = render_paper(spec, tmp_path / "d.docx")
    texts = _texts(out)
    assert texts.index("T") < texts.index("A summary that was asked for.")
    assert _breaks(out) == 1


def test_empty_title_page_lines_are_skipped(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment", "title_page": True,
                 "title_page_lines": ["Course: FOCP", "", "   "]},
        "blocks": _BODY,
    })
    out = render_paper(spec, tmp_path / "d.docx")
    assert "Course: FOCP" in _texts(out)


# ── page break ──────────────────────────────────────────────────────────────

def test_page_break_block_starts_a_new_page(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment"},
        "blocks": [*_BODY, {"type": "page_break"},
                   {"type": "heading", "level": 1, "text": "Conclusion"}],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    assert _breaks(out) == 1


def test_cover_sheet_and_page_breaks_coexist(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment", "title_page": True},
        "blocks": [*_BODY, {"type": "page_break"},
                   {"type": "heading", "level": 1, "text": "Appendix"}],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    assert _breaks(out) == 2


# ── per-block formatting overrides ──────────────────────────────────────────

@pytest.mark.parametrize("override,check", [
    ({"line_spacing": 2.0}, lambda p: p.paragraph_format.line_spacing == 2.0),
    ({"alignment": "center"}, lambda p: p.paragraph_format.alignment is not None),
    ({"bold": True}, lambda p: p.runs[0].bold is True),
    ({"font": "Arial"}, lambda p: p.runs[0].font.name == "Arial"),
    ({"size_pt": 14}, lambda p: p.runs[0].font.size.pt == 14),
])
def test_a_block_style_override_reaches_the_document(override, check, tmp_path):
    """"Double-space it", "use Arial", "centre that" must be expressible."""
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment"},
        "blocks": [{"type": "paragraph", "text": "Styled text.", "style": override}],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    para = next(p for p in Document(str(out)).paragraphs if p.text == "Styled text.")
    assert check(para)


# ── fitting a screenshot to the page ────────────────────────────────────────

def _page_images(path):
    """(width_pt, height_pt) of every image, page by page, via a real render."""
    import shutil
    import pymupdf
    from app.docos.parser.paginator import docx_to_pdf
    pdf = docx_to_pdf(Path(path).read_bytes())
    doc = pymupdf.open(str(pdf))
    return [[(i["bbox"][2] - i["bbox"][0], i["bbox"][3] - i["bbox"][1])
             for i in page.get_image_info()] for page in doc]


def _listing(lines: int) -> str:
    return "\n".join(f'    cout << "line {i}" << endl;' for i in range(lines))


def test_a_short_listing_stays_one_screenshot(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment"},
        "blocks": [{"type": "code", "language": "cpp", "text": _listing(8),
                    "render": "image", "filename": "a.cpp", "caption": "short"}],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    captions = [p.text for p in Document(str(out)).paragraphs if p.text.startswith("Listing")]
    assert captions == ["Listing 1. a.cpp — short"], "no part numbering when it fits"


def test_a_long_listing_is_split_rather_than_shrunk(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment"},
        "blocks": [{"type": "code", "language": "cpp", "text": _listing(90),
                    "render": "image", "filename": "a.cpp", "caption": "long"}],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    captions = [p.text for p in Document(str(out)).paragraphs if p.text.startswith("Listing")]
    assert len(captions) > 1, "a listing too tall for one page should be split"
    assert all("Listing 1." in c for c in captions), "the split is still one listing"
    assert "part 1 of" in captions[0]


def test_no_part_of_a_long_listing_overflows_its_page(tmp_path):
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment"},
        "blocks": [{"type": "code", "language": "cpp", "text": _listing(90),
                    "render": "image", "filename": "a.cpp", "caption": "long"}],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    text_w, text_h = (8.5 - 2.0) * 72, (11 - 2.0) * 72
    seen = 0
    for page in _page_images(out):
        for w, h in page:
            seen += 1
            assert w <= text_w + 1, f"image {w}pt wider than the {text_w}pt text area"
            assert h <= text_h + 1, f"image {h}pt taller than the {text_h}pt text area"
    assert seen >= 2


def test_a_split_listing_keeps_a_usable_width(tmp_path):
    """Shrinking one image to fit a page makes the code unreadable; splitting
    keeps every part near the full column width."""
    spec = PaperSpec.model_validate({
        "meta": {"title": "T", "style": "assignment"},
        "blocks": [{"type": "code", "language": "cpp", "text": _listing(90),
                    "render": "image", "filename": "a.cpp", "caption": "long"}],
    })
    out = render_paper(spec, tmp_path / "d.docx")
    widths = [w for page in _page_images(out) for w, _ in page]
    text_w = (8.5 - 2.0) * 72
    assert widths and min(widths) > text_w * 0.6, f"parts too narrow to read: {widths}"
