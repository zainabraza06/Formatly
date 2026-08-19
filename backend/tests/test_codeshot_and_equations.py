"""Code rendered as an editor screenshot, and equations typeset without LaTeX."""
from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.paper.codeshot import THEMES, render_code_image  # noqa: E402
from app.paper.equations import looks_like_math, render_equation_png  # noqa: E402
from app.paper.renderer import render_paper  # noqa: E402
from app.paper.schema import PaperSpec  # noqa: E402

CODE = 'import numpy as np\n\n\ndef solve(x):\n    return np.sqrt(x)\n'


def _render(tmp_path, blocks, style="assignment"):
    spec = PaperSpec.model_validate({"meta": {"title": "T"}, "blocks": blocks})
    return render_paper(spec, tmp_path / "out.docx", style=style)


# ── editor screenshots ──────────────────────────────────────────────────────

@pytest.mark.parametrize("theme", sorted(THEMES))
def test_a_screenshot_is_produced_in_each_theme(tmp_path, theme):
    out = render_code_image(CODE, tmp_path / f"{theme}.png", language="python",
                            filename="solve.py", theme=theme)
    with Image.open(str(out)) as img:
        assert img.width > 200 and img.height > 60


def test_the_window_chrome_carries_the_filename(tmp_path):
    """The title bar is taller than the code, so a long filename widens it."""
    short = render_code_image(CODE, tmp_path / "a.png", filename="a.py")
    long = render_code_image(CODE, tmp_path / "b.png",
                             filename="a_very_long_module_name_indeed.py")
    with Image.open(str(short)) as s, Image.open(str(long)) as ls:
        assert ls.width >= s.width


def test_an_unknown_language_still_renders(tmp_path):
    out = render_code_image("SELECT 1;", tmp_path / "x.png", language="not-a-language")
    assert out.exists()


def test_empty_code_is_refused(tmp_path):
    with pytest.raises(ValueError):
        render_code_image("   \n  ", tmp_path / "x.png")


def test_render_image_puts_a_picture_in_the_document(tmp_path):
    out = _render(tmp_path, [{"type": "code", "language": "python", "filename": "solve.py",
                              "caption": "the solver", "text": CODE, "render": "image"}])
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 1
    texts = [p.text for p in doc.paragraphs]
    assert "Listing 1. solve.py — the solver" in texts
    # the code itself is in the picture, not set as text
    assert "import numpy as np" not in texts


def test_text_render_stays_the_default(tmp_path):
    out = _render(tmp_path, [{"type": "code", "language": "python", "text": CODE}])
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 0
    assert "import numpy as np" in [p.text for p in doc.paragraphs]


# ── equations ───────────────────────────────────────────────────────────────

@pytest.mark.parametrize("expr", [
    r"d = \frac{\bar{x}_1 - \bar{x}_2}{\sqrt{(s_1^2 + s_2^2)/2}}",
    r"R(\tau) = \sum_{t=0}^{N-\tau} (x_t - \bar{x})^2",
    r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}",
    r"\int_0^\infty e^{-x^2}\,dx = \frac{\sqrt{\pi}}{2}",
])
def test_real_mathematics_typesets(tmp_path, expr):
    assert looks_like_math(expr)
    out, width_in = render_equation_png(expr, tmp_path / "eq.png", size_pt=12)
    assert out.exists() and 0.1 < width_in < 12


@pytest.mark.parametrize("wrapper", ["${}$", "$${}$$", r"\[{}\]", r"\({}\)", "{}"])
def test_the_delimiters_a_model_emits_are_all_accepted(tmp_path, wrapper):
    out, _ = render_equation_png(wrapper.format(r"\frac{a}{b}"), tmp_path / "eq.png")
    assert out.exists()


def test_plain_prose_equations_are_not_treated_as_markup():
    assert not looks_like_math("y = mx + c")
    assert not looks_like_math("the value is 12.4")


def test_unparseable_markup_raises_rather_than_writing_a_broken_image(tmp_path):
    with pytest.raises(ValueError):
        render_equation_png(r"\nonsense{q}", tmp_path / "eq.png")


def test_an_equation_needing_typesetting_becomes_a_picture(tmp_path):
    out = _render(tmp_path, [
        {"type": "equation", "text": r"d = \frac{a}{b}", "numbered": True}])
    assert len(Document(str(out)).inline_shapes) == 1


def test_a_plain_equation_stays_selectable_text(tmp_path):
    out = _render(tmp_path, [
        {"type": "equation", "text": "y = mx + c", "numbered": True}])
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 0
    assert any("y = mx + c" in p.text for p in doc.paragraphs)


def test_unparseable_markup_falls_back_to_text_instead_of_losing_the_equation(tmp_path):
    out = _render(tmp_path, [
        {"type": "equation", "text": r"\nonsense{q}", "render": "image", "numbered": False}])
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 0
    assert any(r"\nonsense{q}" in p.text for p in doc.paragraphs)


def test_a_numbered_equation_is_centred_with_the_number_at_the_margin(tmp_path):
    out = _render(tmp_path, [
        {"type": "equation", "text": r"d = \frac{a}{b}", "numbered": True}])
    para = next(p for p in Document(str(out)).paragraphs if p.text.endswith("(1)"))
    stops = [(s.position, s.alignment) for s in para.paragraph_format.tab_stops]
    assert len(stops) == 2
    assert stops[0][0] < stops[1][0]     # centre stop, then the right-hand stop
