"""Tests for the paper/report pipeline: schema → style resolver → DOCX renderer.

Runs fully offline (no LLM contacted); the generator is tested with a stub router.
"""
from __future__ import annotations

import json

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.paper.generator import PaperGenerationError, generate_paper
from app.paper.renderer import render_paper
from app.paper.schema import PaperSpec
from app.paper.styles import get_stylesheet, list_styles
from app.paper.stylesheet import alpha, resolve, roman

SPEC = {
    "meta": {
        "title": "An Evaluation of Models",
        "authors": [{"name": "A. Researcher", "affiliation": "Test Lab", "email": "a@lab.org"}],
        "abstract": "This paper evaluates several models.",
        "keywords": ["testing", "models"],
    },
    "blocks": [
        {"type": "heading", "level": 1, "text": "Introduction"},
        {"type": "paragraph", "text": "Background prose goes here."},
        {"type": "heading", "level": 2, "text": "Dataset"},
        {"type": "list", "ordered": False, "items": ["First point", "Second point"]},
        {"type": "equation", "text": "E = mc^2", "numbered": True},
        {"type": "heading", "level": 1, "text": "Results"},
        {"type": "table", "caption": "Model performance",
         "columns": ["Model", "Accuracy"], "rows": [["CNN", "0.94"], ["SVM", "0.85"]]},
        {"type": "figure", "caption": "Accuracy by model",
         "chart": {"kind": "bar", "title": "Accuracy", "labels": ["CNN", "SVM"],
                   "values": [0.94, 0.85], "rationale": "compares one metric across models"}},
    ],
    "references": ["A. Author, \"A title,\" Journal, 2023."],
    "visualization_plan": [
        {"data": "accuracy per model", "kind": "bar", "rationale": "categorical comparison"}
    ],
}


# ── numbering helpers ───────────────────────────────────────────────────────

def test_roman_and_alpha():
    assert roman(1) == "I" and roman(4) == "IV" and roman(9) == "IX" and roman(14) == "XIV"
    assert alpha(1) == "A" and alpha(2) == "B" and alpha(27) == "AA"


# ── resolver ────────────────────────────────────────────────────────────────

def test_resolver_stamps_explicit_formatting():
    spec = resolve(PaperSpec.model_validate(SPEC), "ieee")
    assert spec.resolved is True
    assert spec.meta.style == "ieee"
    assert spec.meta.page.columns == 2

    heading = spec.blocks[0]
    assert heading.style.font == "Times New Roman"
    assert heading.style.small_caps is True

    body = spec.blocks[1]
    assert body.style.alignment == "justify"
    assert body.style.size_pt == 10

    table = next(b for b in spec.blocks if b.type == "table")
    assert table.header_style.bold is True
    assert table.cell_style.size_pt == 8


def test_resolver_keeps_author_overrides():
    data = json.loads(json.dumps(SPEC))
    data["blocks"][1]["style"] = {"italic": True, "size_pt": 14}
    spec = resolve(PaperSpec.model_validate(data), "ieee")
    body = spec.blocks[1]
    assert body.style.italic is True      # override survived
    assert body.style.size_pt == 14
    assert body.style.font == "Times New Roman"  # base still applied


def test_style_registry_and_aliases():
    ids = {s["id"] for s in list_styles()}
    assert {"ieee", "ieee_1col", "assignment"} == ids
    assert get_stylesheet("IEEE").id == "ieee"
    assert get_stylesheet("Coursework").id == "assignment"
    assert get_stylesheet("something unknown").id == "ieee"  # safe default


# ── renderer ────────────────────────────────────────────────────────────────

@pytest.mark.parametrize("style,cols,font", [
    ("ieee", "2", "Times New Roman"),
    ("ieee_1col", "1", "Times New Roman"),
    ("assignment", "1", "Times New Roman"),
])
def test_render_each_style(tmp_path, style, cols, font):
    out = tmp_path / f"{style}.docx"
    render_paper(PaperSpec.model_validate(SPEC), out, style=style)
    assert out.exists()

    doc = Document(str(out))
    # last section carries the body column count
    last = doc.sections[-1]._sectPr.find(qn("w:cols"))
    assert last.get(qn("w:num")) == cols

    paras = [p for p in doc.paragraphs if p.text.strip()]
    assert paras[0].text == "An Evaluation of Models"
    assert paras[0].runs[0].font.name == font

    assert len(doc.tables) == 1
    # the bar chart, plus "E = mc^2" typeset as an image — a superscript cannot
    # be shown in a run of text, so an equation carrying one is rasterised
    assert len(doc.inline_shapes) == 2
    assert any("References" in p.text for p in paras)


def test_ieee_numbering_and_captions(tmp_path):
    out = tmp_path / "ieee.docx"
    render_paper(PaperSpec.model_validate(SPEC), out, style="ieee")
    texts = [p.text for p in Document(str(out)).paragraphs]

    assert any(t.startswith("I. Introduction") for t in texts)
    assert any(t.startswith("A. Dataset") for t in texts)
    assert any(t.startswith("II. Results") for t in texts)
    assert any(t.startswith("TABLE I") for t in texts)
    # IEEE figure captions run inline: "Fig. 1. Accuracy by model"
    assert any(t.startswith("Fig. 1.") and "Accuracy by model" in t and "\n" not in t
               for t in texts)


def test_decimal_numbering_for_the_assignment_style(tmp_path):
    out = tmp_path / "assignment.docx"
    render_paper(PaperSpec.model_validate(SPEC), out, style="assignment")
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any(t.startswith("1. Introduction") for t in texts)
    assert any(t.startswith("1.1 Dataset") for t in texts)
    assert any(t.startswith("2. Results") for t in texts)
    assert any(t.startswith("Table 1.") for t in texts)


def test_caption_separator_is_honoured(tmp_path):
    """IEEE puts the table title on its own line; the assignment style runs it on
    after the label. Both come from the sheet's caption separator."""
    ieee = tmp_path / "ieee.docx"
    render_paper(PaperSpec.model_validate(SPEC), ieee, style="ieee")
    cap = next(t for t in (p.text for p in Document(str(ieee)).paragraphs)
               if t.startswith("TABLE I"))
    assert cap == "TABLE I\nModel performance"   # label and title on separate lines

    formal = tmp_path / "assignment.docx"
    render_paper(PaperSpec.model_validate(SPEC), formal, style="assignment")
    cap = next(t for t in (p.text for p in Document(str(formal)).paragraphs)
               if t.startswith("Table 1"))
    assert cap == "Table 1. Model performance"


# ── generator (stubbed LLM) ─────────────────────────────────────────────────

class _StubRouter:
    def __init__(self, reply: str):
        self.reply = reply
        self.messages = None      # captured so tests can inspect the prompt

    def chat(self, messages, **kwargs):
        self.messages = messages
        return self.reply, "stub", 0.01

    @property
    def system(self) -> str:
        return next(m["content"] for m in self.messages if m["role"] == "system")

    @property
    def user(self) -> str:
        return next(m["content"] for m in self.messages if m["role"] == "user")


def test_generator_validates_and_resolves():
    spec, provider = generate_paper(
        raw_text="some material",
        style="ieee",
        router=_StubRouter(json.dumps(SPEC)),
    )
    assert provider == "stub"
    assert spec.resolved is True
    assert spec.meta.style == "ieee"
    assert spec.blocks[0].style.small_caps is True


def test_generator_tolerates_fenced_json():
    fenced = "```json\n" + json.dumps(SPEC) + "\n```"
    spec, _ = generate_paper(raw_text="x", router=_StubRouter(fenced))
    assert spec.meta.title == "An Evaluation of Models"


def test_generator_rejects_non_json():
    with pytest.raises(PaperGenerationError):
        generate_paper(raw_text="x", router=_StubRouter("I cannot help with that."))


def test_generator_requires_material():
    with pytest.raises(PaperGenerationError):
        generate_paper(raw_text="   ", router=_StubRouter("{}"))


def test_generator_keeps_visualization_plan():
    spec, _ = generate_paper(raw_text="x", router=_StubRouter(json.dumps(SPEC)))
    assert spec.visualization_plan[0].kind == "bar"
    assert "accuracy" in spec.visualization_plan[0].data


# ── arbitrary labelled material (no fixed "code"/"results" fields) ──────────

def test_attachments_reach_the_prompt_under_their_own_labels():
    router = _StubRouter(json.dumps(SPEC))
    generate_paper(
        raw_text="Write up the churn analysis.",
        doc_kind="report",
        attachments=[
            {"label": "Survey responses", "content": "412 replies; 63% cited price"},
            {"label": "Interview transcript", "content": "Q: why did you leave? A: too costly"},
        ],
        router=router,
    )
    user = router.user
    assert "Survey responses" in user and "63% cited price" in user
    assert "Interview transcript" in user and "too costly" in user
    assert "churn analysis" in user
    assert "report" in user


def test_blank_attachments_are_dropped():
    router = _StubRouter(json.dumps(SPEC))
    generate_paper(
        raw_text="material",
        attachments=[{"label": "Empty", "content": "   "}],
        router=router,
    )
    assert "additional_material" not in router.user


def test_no_attachments_is_fine():
    router = _StubRouter(json.dumps(SPEC))
    spec, _ = generate_paper(raw_text="Just prose, nothing else.", router=router)
    assert spec.resolved is True
    assert "additional_material" not in router.user


def test_prompt_is_domain_neutral():
    """The pipeline must not assume machine learning — that was only one example."""
    router = _StubRouter(json.dumps(SPEC))
    generate_paper(raw_text="A legal memorandum on lease obligations.",
                   doc_kind="memo", router=router)
    system = router.system.lower()
    for ml_ism in ("epoch", "model accuracy", "cnn", "svm", "dataset"):
        assert ml_ism not in system, f"prompt leaks a domain assumption: {ml_ism!r}"


def test_doc_kind_is_free_text():
    router = _StubRouter(json.dumps(SPEC))
    generate_paper(raw_text="x", doc_kind="grant proposal", router=router)
    assert "grant proposal" in router.user
