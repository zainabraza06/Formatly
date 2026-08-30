"""Lists: reading them, making them, and writing them back out.

"Make the contributions a bulleted list" used to select something and report
Done, because nothing in the system knew what a list was.
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from app.docos.actions import validate_batch
from app.docos.command.engine import CommandEngine
from app.docos.execution.engine import ExecutionEngine
from app.docos.execution.lists import split_items
from app.docos.export import graph_to_docx_bytes
from app.docos.graph import DocumentGraph, Node, NodeType
from app.docos.parser import parse_docx_bytes


def _graph(*bodies: str) -> DocumentGraph:
    root = Node(type=NodeType.DOCUMENT)
    root.children = [Node(type=NodeType.HEADING, content="Introduction")]
    root.children += [Node(type=NodeType.BODY, content=b) for b in bodies]
    return DocumentGraph(root=root, title="t")


def _run(graph: DocumentGraph, **params) -> DocumentGraph:
    batch = validate_batch({"reasoning": "t", "actions": [
        {"type": "list", "target": "body", "params": params or {"kind": "bullet"}}]})
    result = ExecutionEngine().execute(graph, batch)
    assert result.ok, result.error
    return result.graph


def _items(graph: DocumentGraph) -> list[tuple[str, str]]:
    return [((n.metadata.get("list") or {}).get("kind", ""), n.content)
            for n in graph.nodes() if n.type is NodeType.BODY]


# ── reading a list out of a Word file ────────────────────────────────────────

def test_word_list_styles_are_read_and_written_back():
    doc = Document()
    doc.add_paragraph("The contributions are:")
    doc.add_paragraph("A benchmark.", style="List Bullet")
    doc.add_paragraph("A feature set.", style="List Bullet")
    doc.add_paragraph("First step.", style="List Number")
    buf = io.BytesIO()
    doc.save(buf)

    graph = parse_docx_bytes(buf.getvalue(), title="t")
    kinds = [(n.metadata.get("list") or {}).get("kind") for n in graph.nodes()
             if n.content.strip()]
    assert kinds == [None, "bullet", "bullet", "number"]

    # and survives a round trip through Word's own styles
    again = parse_docx_bytes(graph_to_docx_bytes(graph), title="t")
    assert [(n.metadata.get("list") or {}).get("kind") for n in again.nodes()
            if n.content.strip()] == kinds


# ── cutting an enumerating paragraph into items ──────────────────────────────

@pytest.mark.parametrize("text, lead, items", [
    ("The contributions are: (i) a benchmark; (ii) a feature set; and (iii) a protocol.",
     "The contributions are:", ["A benchmark", "A feature set", "A protocol."]),
    ("We do three things. 1. Propose. 2. Release. 3. Evaluate.",
     "We do three things.", ["Propose.", "Release.", "Evaluate."]),
    ("We make three contributions: a benchmark; a feature set; a protocol.",
     "We make three contributions:", ["A benchmark", "A feature set", "A protocol."]),
])
def test_split_items_cuts_where_the_author_enumerated(text, lead, items):
    assert split_items(text) == (lead, items)


@pytest.mark.parametrize("text", [
    "This paper (see Section 2) presents a study of falls.",
    "Figure 1 shows the results.",
    "",
])
def test_split_items_leaves_ordinary_prose_alone(text):
    assert split_items(text) == ("", [])


# ── the action ───────────────────────────────────────────────────────────────

def test_bulleting_an_enumerating_paragraph_makes_one_item_per_thing():
    graph = _run(_graph("Falls are a risk.",
                        "The contributions are: (i) a benchmark; (ii) a feature set."))
    assert _items(graph) == [
        ("", "Falls are a risk."),          # untouched: it enumerates nothing
        ("", "The contributions are:"),     # the lead-in stays a paragraph
        ("bullet", "A benchmark"),
        ("bullet", "A feature set."),
    ]


def test_the_sentence_that_introduces_the_list_does_not_get_a_bullet():
    graph = _run(_graph("The main contributions of this work are:",
                        "A subject-independent benchmark.",
                        "An evidence-gated feature set."))
    assert _items(graph) == [
        ("", "The main contributions of this work are:"),
        ("bullet", "A subject-independent benchmark."),
        ("bullet", "An evidence-gated feature set."),
    ]


def test_a_colon_on_the_last_paragraph_is_still_an_item():
    """It announces nothing that is in scope, so it is one of the things."""
    graph = _run(_graph("Inputs:", "Outputs:"))
    assert _items(graph) == [("bullet", "Inputs:"), ("bullet", "Outputs:")]


def test_a_heading_in_scope_is_not_turned_into_a_bullet():
    graph = _run(_graph("First point.", "Second point."))
    assert [n.content for n in graph.nodes()
            if n.type is NodeType.HEADING and not n.metadata.get("list")] == ["Introduction"]


def test_word_list_paragraph_style_alone_is_not_a_list():
    """Word gives List Paragraph to any indented block, list or not."""
    doc = Document()
    doc.add_paragraph("The contributions are:", style="List Paragraph")
    buf = io.BytesIO()
    doc.save(buf)

    graph = parse_docx_bytes(buf.getvalue(), title="t")
    assert all(not n.metadata.get("list") for n in graph.nodes())


def test_plain_paragraphs_become_items_as_they_stand():
    graph = _run(_graph("First point.", "Second point."), kind="number")
    assert _items(graph) == [("number", "First point."), ("number", "Second point.")]


def test_a_list_can_be_taken_back_out():
    graph = _run(_graph("First point.", "Second point."), kind="number")
    assert _items(_run(graph, kind="none")) == [
        ("", "First point."), ("", "Second point.")]


def test_bulleting_twice_leaves_the_same_list():
    once = _run(_graph("First point.", "Second point."))
    assert _items(_run(once, kind="bullet")) == _items(once)


# ── the words people use for it ──────────────────────────────────────────────

@pytest.mark.parametrize("command, kind", [
    ("make the contributions listed in bullets", "bullet"),
    ("put the steps in a numbered list", "number"),
    ("itemise the limitations", "bullet"),
    ("turn these into bullet points", "bullet"),
    ("remove the bullets from the contributions", "none"),
])
def test_the_heuristic_understands_a_request_for_a_list(command, kind):
    batch = CommandEngine()._heuristic_actions(command)
    assert [a.type.value for a in batch.actions] == ["list"]
    assert batch.actions[0].params["kind"] == kind


def test_listing_the_headings_is_still_a_selection():
    batch = CommandEngine()._heuristic_actions("list the headings")
    assert [a.type.value for a in batch.actions] == ["select"]
