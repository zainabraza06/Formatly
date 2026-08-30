"""End-to-end tests for the DocOS engines.

Covers: DOCX parsing, action validation, execution + events, version engine
(snapshot/replay, undo/redo/rewind/diff) and the command engine heuristics.
Runs fully offline — no LLM provider is contacted (heuristic fallback is exercised).
"""
from __future__ import annotations

import io
import os
import tempfile

import pytest

from app.docos.actions import ActionValidationError, validate_batch
from app.docos.command import CommandEngine
from app.docos.execution import ExecutionEngine
from app.docos.graph import DocumentGraph, Node, NodeType, Run, Style
from app.docos.parser import parse_docx_bytes
from app.docos.versioning import VersionEngine
from app.docos.versioning.store import VersionStore


# ── fixtures ────────────────────────────────────────────────────────────────

def _sample_docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the first body paragraph.")
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph("We used a rigorous approach.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(1, 1).text = "D"
    doc.add_heading("References", level=1)
    doc.add_paragraph("Smith, J. (2020). A study.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _sample_graph() -> DocumentGraph:
    root = Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.HEADING, content="Intro"),
        Node(type=NodeType.BODY, content="Body one."),
        Node(type=NodeType.BODY, content="Body two."),
        Node(type=NodeType.FIGURE, children=[Node(type=NodeType.IMAGE, content="img")]),
        Node(type=NodeType.HORIZONTAL_RULE),
        Node(type=NodeType.REFERENCE, content="Ref 1"),
    ])
    return DocumentGraph(root=root, title="T")


@pytest.fixture()
def engine(tmp_path):
    store = VersionStore(db_path=tmp_path / "test.db")
    return VersionEngine(store=store)


# ── parser ──────────────────────────────────────────────────────────────────

def test_parser_classifies_nodes():
    graph = parse_docx_bytes(_sample_docx_bytes(), title="Sample")
    types = [n.type for n in graph.nodes()]
    assert NodeType.HEADING in types
    assert NodeType.SUBHEADING in types  # heading level 2
    assert NodeType.TABLE in types
    assert NodeType.REFERENCE in types   # paragraph after "References" heading
    tables = graph.find_by_types([NodeType.TABLE])
    assert tables and tables[0].metadata["cols"] == 2


# ── action validation ───────────────────────────────────────────────────────

def test_validate_rejects_unknown_target():
    with pytest.raises(ActionValidationError):
        validate_batch({"actions": [{"type": "select", "target": "banana"}]})


def test_validate_rejects_bulk_delete_without_confirm():
    with pytest.raises(ActionValidationError):
        validate_batch({"actions": [{"type": "delete", "target": "body"}]})


def test_validate_accepts_good_batch():
    batch = validate_batch({"actions": [
        {"type": "select", "target": "heading"},
        {"type": "format", "target": "heading", "style": {"bold": True, "font_size": 18}},
    ]})
    assert len(batch.actions) == 2


# ── execution ───────────────────────────────────────────────────────────────

def test_execution_formats_and_emits_events():
    graph = _sample_graph()
    batch = validate_batch({"actions": [
        {"type": "format", "target": "heading", "style": {"bold": True, "color": "#003366"}},
    ]})
    result = ExecutionEngine().execute(graph, batch)
    assert result.ok
    heading = result.graph.find_by_types([NodeType.HEADING])[0]
    assert heading.style.bold is True and heading.style.color == "#003366"
    names = {e.name.value for e in result.events}
    assert {"batch_started", "format_started", "format_finished", "batch_finished"} <= names
    # original graph untouched (execution works on a clone)
    assert graph.find_by_types([NodeType.HEADING])[0].style.bold is None


def test_execution_deletes_horizontal_rule():
    graph = _sample_graph()
    batch = validate_batch({"actions": [{"type": "delete", "target": "horizontal_rule"}]})
    result = ExecutionEngine().execute(graph, batch)
    assert result.ok
    assert not result.graph.find_by_types([NodeType.HORIZONTAL_RULE])


# ── version engine ──────────────────────────────────────────────────────────

def test_version_snapshot_replay_and_undo(engine):
    graph = _sample_graph()
    engine.init_document("doc1", "T", graph)

    # apply 12 edits so at least one non-checkpoint replay happens
    for i in range(12):
        g = engine.current_graph("doc1")
        batch = validate_batch({"actions": [
            {"type": "highlight", "target": "heading", "params": {"color": f"#00000{i%9}"}}
        ]})
        res = ExecutionEngine().execute(g, batch)
        engine.commit("doc1", batch, res.graph)

    hist = engine.history("doc1")
    assert len(hist) == 13                 # v0 import + 12 edits
    assert any(v.is_checkpoint and v.seq == 10 for v in hist)  # snapshot at seq 10

    # materialise a non-checkpoint version by replay
    mid = next(v for v in hist if v.seq == 7)
    g7 = engine.materialize(mid.id)
    assert g7.find_by_types([NodeType.HEADING])[0].style.highlight == "#000006"

    # undo moves pointer back one version
    info = engine.undo("doc1")
    assert info is not None and info.seq == 11
    redo = engine.redo("doc1")
    assert redo is not None and redo.seq == 12


def test_version_rewind_and_diff(engine):
    graph = _sample_graph()
    engine.init_document("doc2", "T", graph)
    g = engine.current_graph("doc2")
    batch = validate_batch({"actions": [
        {"type": "format", "target": "heading", "style": {"font_size": 22}}]})
    res = ExecutionEngine().execute(g, batch)
    v1 = engine.commit("doc2", batch, res.graph)

    diff = engine.diff(engine.history("doc2")[0].id, v1.id)
    assert diff.changed and diff.changed[0]["style"]["after"]["font_size"] == 22

    back = engine.rewind("doc2", engine.history("doc2")[0].id)
    assert back.seq == 0
    assert engine.current_graph("doc2").find_by_types([NodeType.HEADING])[0].style.font_size is None


# ── command engine (offline heuristics) ─────────────────────────────────────

def test_command_engine_heuristics():
    # force heuristic path by injecting a router that always fails
    class DeadRouter:
        def chat(self, *a, **k):
            raise RuntimeError("no providers")

    ce = CommandEngine(router=DeadRouter())
    graph = _sample_graph()

    r = ce.parse("Highlight all figures", graph)
    assert r.kind == "actions" and r.source == "heuristic"
    assert r.batch.actions[0].type.value == "highlight"

    r2 = ce.parse("Remove every horizontal line", graph)
    assert r2.batch.actions[0].type.value == "delete"
    assert r2.batch.actions[0].target == "horizontal_rule"

    r3 = ce.parse("Undo", graph)
    assert r3.kind == "control" and r3.control.kind == "undo"

    r4 = ce.parse("Rewind to version 8", graph)
    assert r4.control.kind == "rewind" and r4.control.params["seq"] == 8


# ── diff detail (what changed with what) ────────────────────────────────────

def test_word_segments_isolate_the_words_that_moved():
    from app.docos.versioning.diff import word_segments

    segs = word_segments("the quick brown fox", "the slow brown fox")
    assert [s["op"] for s in segs] == ["equal", "delete", "insert", "equal"]
    assert "".join(s["text"] for s in segs if s["op"] in ("equal", "delete")) == "the quick brown fox"
    assert "".join(s["text"] for s in segs if s["op"] in ("equal", "insert")) == "the slow brown fox"

    assert word_segments("same", "same") == [{"op": "equal", "text": "same"}]


def test_diff_carries_before_and_after_text_per_node():
    from app.docos.versioning.diff import diff_graphs

    before = DocumentGraph(title="T", root=Node(
        id="root", type=NodeType.DOCUMENT, children=[
            Node(id="p1", type=NodeType.BODY, content="We used a rigorous approach."),
            Node(id="p2", type=NodeType.BODY, content="Dropped later."),
        ]))
    after = DocumentGraph(title="T", root=Node(
        id="root", type=NodeType.DOCUMENT, children=[
            Node(id="p1", type=NodeType.BODY, content="We used a careful approach."),
            Node(id="p3", type=NodeType.BODY, content="Brand new."),
        ]))

    d = diff_graphs(before, after).to_dict()
    assert d["added"][0]["content"] == "Brand new."
    assert d["removed"][0]["content"] == "Dropped later."

    change = d["changed"][0]["content"]
    assert change["before"] == "We used a rigorous approach."
    assert change["after"] == "We used a careful approach."
    assert {"op": "insert", "text": "careful"} in change["segments"]
    assert {"op": "delete", "text": "rigorous"} in change["segments"]

    assert d["summary"] == {"added": 1, "removed": 1, "changed": 1, "text_changed": 1,
                            "style_changed": 0, "words_added": 1, "words_removed": 1}


def test_diff_names_the_style_fields_that_changed(engine):
    graph = _sample_graph()
    engine.init_document("doc3", "T", graph)
    g = engine.current_graph("doc3")
    batch = validate_batch({"actions": [
        {"type": "format", "target": "heading", "style": {"font_size": 22}}]})
    res = ExecutionEngine().execute(g, batch)
    v1 = engine.commit("doc3", batch, res.graph)

    d = engine.diff(engine.history("doc3")[0].id, v1.id).to_dict()
    fields = d["changed"][0]["style"]["fields"]
    assert {"field": "font_size", "before": None, "after": 22} in fields
    assert d["summary"]["style_changed"] == len(d["changed"])


# ── deleting an upload ──────────────────────────────────────────────────────

def test_delete_document_removes_it_and_its_versions(engine):
    graph = _sample_graph()
    engine.init_document("doc4", "T", graph)
    batch = validate_batch({"actions": [
        {"type": "format", "target": "heading", "style": {"bold": True}}]})
    res = ExecutionEngine().execute(engine.current_graph("doc4"), batch)
    engine.commit("doc4", batch, res.graph)
    assert len(engine.store.list_versions("doc4")) == 2

    assert engine.store.delete_document("doc4") is True
    assert engine.store.get_document("doc4") is None
    assert engine.store.list_versions("doc4") == []
    # A second delete is not an error, but it reports that nothing was there.
    assert engine.store.delete_document("doc4") is False


def test_delete_document_leaves_other_documents_alone(engine):
    engine.init_document("keep", "K", _sample_graph())
    engine.init_document("drop", "D", _sample_graph())

    engine.store.delete_document("drop")

    assert engine.store.get_document("keep") is not None
    assert len(engine.store.list_versions("keep")) == 1


def test_delete_document_refuses_someone_elses(tmp_path):
    from app.docos.service import DocOSService
    from app.docos.versioning import VersionEngine

    service = DocOSService(versions=VersionEngine(
        store=VersionStore(db_path=tmp_path / "own.db")))
    service.versions.init_document("mine", "M", _sample_graph(), owner_id="user-a")

    with pytest.raises(PermissionError):
        service.delete_document("mine", owner_id="user-b")
    with pytest.raises(KeyError):
        service.delete_document("nope", owner_id="user-a")

    assert service.delete_document("mine", owner_id="user-a") is True


# ── inline runs through the execution engine ────────────────────────────────

def _mixed_node():
    return Node(type=NodeType.BODY, content="Recent work outperforms the baseline.",
                runs=[Run(text="Recent work "),
                      Run(text="outperforms", style=Style(bold=True, italic=False)),
                      Run(text=" the baseline.")])


def test_formatting_a_paragraph_reaches_the_words_inside_it():
    """A run saying italic=False must not defeat "make the body italic"."""
    node = _mixed_node()
    node.apply_style(Style(italic=True))

    assert node.style.italic is True
    assert all(r.style.italic is None for r in node.runs), "the override should be cleared"
    # Only the attribute being set is cleared; the bold phrase stays bold.
    assert any(r.style.bold is True for r in node.runs)


def test_replace_keeps_inline_formatting_when_it_still_fits():
    node = _mixed_node()
    assert node.replace_text("baseline", "benchmark") is True
    assert node.content == "Recent work outperforms the benchmark."
    assert node.runs_describe_content()
    assert any(r.style.bold is True for r in node.runs)


def test_replace_across_a_run_boundary_drops_the_pieces_rather_than_guessing():
    node = _mixed_node()
    # "work outperforms" straddles two runs, so no run-wise replacement matches.
    assert node.replace_text("work outperforms", "work beats") is True
    assert node.content == "Recent work beats the baseline."
    assert node.runs == []
    assert [r.text for r in node.inline_runs()] == ["Recent work beats the baseline."]


def test_pasting_a_paragraph_copies_how_it_is_formatted():
    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[_mixed_node()]))
    original = graph.root.children[0]
    batch = validate_batch({"actions": [
        {"type": "copy", "node_ids": [original.id]},
        {"type": "paste", "node_ids": [original.id]},
    ]})
    result = ExecutionEngine().execute(graph, batch)

    copies = [n for n in result.graph.nodes() if n.content == original.content]
    assert len(copies) == 2, "the paragraph and its copy"
    assert all(any(r.style.bold for r in c.inline_runs()) for c in copies)


# ── what the document is ────────────────────────────────────────────────────

def _paper_graph() -> DocumentGraph:
    return DocumentGraph(title="A Study of Falls", root=Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.BODY, content="Abstract— This paper evaluates fall detection [1], [2]."),
        Node(type=NodeType.HEADING, content="I. INTRODUCTION"),
        Node(type=NodeType.BODY, content="Falls are a major cause of injury among older adults."),
        Node(type=NodeType.HEADING, content="II. METHOD"),
        Node(type=NodeType.SUBHEADING, content="A. Feature Extraction"),
        Node(type=NodeType.BODY, content=r"Each window is summarised by $\hat{E}_{j,c}$ per bin."),
        Node(type=NodeType.TABLE, children=[
            Node(type=NodeType.TABLE_ROW, children=[
                Node(type=NodeType.TABLE_CELL, content=r"$\mu_j = \frac{1}{n}\sum_i m_i$"),
                Node(type=NodeType.TABLE_CELL, content="(1)"),
            ]),
        ]),
        Node(type=NodeType.HEADING, content="REFERENCES"),
        Node(type=NodeType.REFERENCE, content="[1] Author, A. (2020). A study."),
    ]))


def test_the_brief_says_what_kind_of_document_it_is():
    from app.docos.command.brief import document_brief

    brief = document_brief(_paper_graph())
    assert brief["kind"] == "research paper"
    assert brief["title"] == "A Study of Falls"
    assert brief["conventions"]["citation_style"] == "numbered, [n]"
    assert "roman" in brief["conventions"]["heading_numbering"]


def test_the_brief_says_where_the_maths_lives():
    """The knowledge that was missing when "convert every equation" reached
    only the equations that happened to be in body paragraphs."""
    from app.docos.command.brief import document_brief

    where = document_brief(_paper_graph())["conventions"]["maths_appears_in"]
    assert where.get("body"), "inline maths in prose"
    assert where.get("table_cell"), "a display equation in its usual one-row table"


def test_the_brief_cuts_the_document_at_its_headings():
    from app.docos.command.brief import document_brief

    sections = document_brief(_paper_graph())["sections"]
    by_heading = {s["heading"]: s for s in sections}

    assert by_heading["I. INTRODUCTION"]["level"] == 1
    assert by_heading["A. Feature Extraction"]["level"] == 2
    assert by_heading["A. Feature Extraction"]["holds"]["tables"] == 1
    assert by_heading["REFERENCES"]["holds"]["references"] == 1
    assert by_heading["II. METHOD"]["holds"] == {}, "a heading followed by a subheading holds nothing itself"


def test_the_planner_is_told_what_the_document_is():
    from app.docos.command.prompt import build_user_message

    message = build_user_message("convert all latex equations", _paper_graph())
    assert "research paper" in message
    assert "maths_appears_in" in message
    assert "table_cell" in message, "so a plan can reach the equations in the table"


# ── placing an instruction in the document ──────────────────────────────────

_BRIEF = {"sections": [
    {"heading": "I. INTRODUCTION", "about": "motivates automated fall detection",
     "node_ids": ["a", "b"]},
    {"heading": "III. RESULTS", "about": "reports 89.01% accuracy on held-out subjects",
     "node_ids": ["c", "d"]},
    {"heading": "II. METHOD", "about": "describes feature extraction per bin",
     "node_ids": ["e", "f"]},
]}


def test_a_request_is_matched_to_the_section_it_describes():
    """The reading is what makes this work: "accuracy" appears nowhere in the
    heading "III. RESULTS", only in what the section was read to be about."""
    from app.docos.command.locate import locate_section

    assert locate_section("Make the section that reports the accuracy bold",
                          _BRIEF)["heading"] == "III. RESULTS"
    assert locate_section("Rewrite the part about the results to be more concise",
                          _BRIEF)["heading"] == "III. RESULTS"
    assert locate_section("tighten the section on feature extraction",
                          _BRIEF)["heading"] == "II. METHOD"


def test_a_request_about_the_whole_document_names_no_section():
    from app.docos.command.locate import locate_section

    assert locate_section("make every heading bold", _BRIEF) is None
    assert locate_section("convert all latex equations", _BRIEF) is None
    assert locate_section("justify the body paragraphs", _BRIEF) is None


def test_one_word_in_common_is_a_coincidence_not_a_match():
    from app.docos.command.locate import locate_section

    # "detection" alone should not carry a section: too thin to act on.
    assert locate_section("the section", _BRIEF) is None
    assert locate_section("the part with the thing", _BRIEF) is None


# ── a watcher that arrives a moment late ────────────────────────────────────

def test_the_hub_replays_what_a_late_watcher_missed():
    """A document is read as soon as it is imported, and the editor's socket
    opens only once the page has loaded — so the first pages are announced to
    an empty room."""
    import asyncio
    from app.docos.events.hub import EventHub

    class FakeSocket:
        def __init__(self):
            self.sent = []

        async def accept(self):
            pass

        async def send_json(self, message):
            self.sent.append(message)

    async def scenario():
        hub = EventHub()
        for page in (1, 2, 3):
            await hub.broadcast("doc", {"event": "reading_progress", "payload": {"page": page}})

        watcher = FakeSocket()
        await hub.connect("doc", watcher)
        return watcher.sent

    replayed = asyncio.run(scenario())
    assert [m["payload"]["page"] for m in replayed] == [1, 2, 3]
    assert all(m["replayed"] for m in replayed), "a replay is marked as one"


def test_the_hub_does_not_buffer_every_document_forever():
    import asyncio
    from app.docos.events.hub import EventHub, _ROOMS_BUFFERED

    async def scenario():
        hub = EventHub()
        for i in range(_ROOMS_BUFFERED + 10):
            await hub.broadcast(f"doc{i}", {"event": "x"})
        return len(hub._recent)

    assert asyncio.run(scenario()) <= _ROOMS_BUFFERED


# ── the planner is actually used ────────────────────────────────────────────

def test_the_llm_path_is_taken_when_the_model_answers():
    """A mistake inside `_llm_actions` — a name that did not exist — was caught
    by a bare `except` and turned every command into a heuristic guess, with
    nothing on the outside to show for it but the word "heuristic"."""
    from app.docos.command import CommandEngine

    class Router:
        def chat(self, messages, max_tokens=None, **_kw):
            return ('{"reasoning": "bold the headings", "actions": '
                    '[{"type": "format", "target": "heading", "style": {"bold": true}}]}'), "fake", 0.1

    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.HEADING, content="I. INTRODUCTION"),
    ]))
    result = CommandEngine(router=Router()).parse("make the headings bold", graph)

    assert result.source == "llm", "the planner answered, so its plan is the one used"
    assert not result.fell_back_because
    assert result.batch.actions[0].type.value == "format"


def test_a_reading_reaches_the_planner_without_breaking_it():
    """The signature bug: `reading` was threaded into the call and not into the
    function, so passing one raised NameError and fell back silently."""
    from app.docos.command import CommandEngine

    seen: dict[str, str] = {}

    class Router:
        def chat(self, messages, max_tokens=None, **_kw):
            seen["message"] = messages[-1]["content"]
            return '{"reasoning": "r", "actions": [{"type": "select", "target": "heading"}]}', "fake", 0.1

    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.HEADING, content="III. RESULTS"),
    ]))
    result = CommandEngine(router=Router()).parse(
        "select the headings", graph, {"III. RESULTS": "reports the accuracy"})

    assert result.source == "llm"
    assert "reports the accuracy" in seen["message"], "the reading travelled with the request"


def test_a_failed_planner_says_why_it_fell_back():
    from app.docos.command import CommandEngine

    class Router:
        def chat(self, *_a, **_kw):
            raise RuntimeError("no providers reachable")

    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.BODY, content="text"),
    ]))
    result = CommandEngine(router=Router()).parse("make the headings bold", graph)

    assert result.source == "heuristic"
    assert "no providers reachable" in result.fell_back_because


def test_the_heuristic_asks_for_a_rewrite_rather_than_selecting_something():
    """"Change all the latex equations" cannot be satisfied by selecting."""
    from app.docos.command import CommandEngine

    engine = CommandEngine()
    for command in ("change all latex equations into readable maths",
                    "convert the equations",
                    "simplify the abstract"):
        batch = engine._heuristic_actions(command)
        assert batch.actions[0].type.value == "rewrite", command
        assert batch.actions[0].params["instruction"] == command

    # and the formatting requests it does understand are untouched
    assert engine._heuristic_actions("make headings bold").actions[0].type.value == "format"
    assert engine._heuristic_actions("highlight figures").actions[0].type.value == "highlight"


# ── the headings in a table are not the document's headings ─────────────────

def _table_docx() -> bytes:
    import io
    from docx import Document

    d = Document()
    d.add_heading("III. RESULTS", 1)
    table = d.add_table(rows=3, cols=2)
    for col, text in enumerate(["Activity", "Precision"]):
        cell = table.cell(0, col)
        cell.text = ""
        cell.paragraphs[0].add_run(text).bold = True
    for row, values in enumerate([["Walking", "0.94"], ["Falling", "0.89"]], start=1):
        for col, value in enumerate(values):
            table.cell(row, col).text = value

    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


def test_a_tables_header_row_is_recognised():
    graph = parse_docx_bytes(_table_docx(), title="T")
    rows = [n for n in graph.nodes() if n.type is NodeType.TABLE_ROW]

    assert rows[0].metadata.get("header_row") is True
    assert not rows[1].metadata.get("header_row")


def test_table_headers_and_document_headings_are_different_targets():
    """"Make the headings in the table bold" used to bold every section title
    in the paper and no table at all."""
    graph = parse_docx_bytes(_table_docx(), title="T")

    assert [c.content for c in graph.resolve_target("table_header")] == ["Activity", "Precision"]
    assert [n.content for n in graph.resolve_target("heading")] == ["III. RESULTS"]


def test_a_request_about_headings_in_a_table_finds_the_table():
    from app.docos.command.engine import _guess_target

    assert _guess_target("make all headings inside the table bold") == "table_header"
    assert _guess_target("bold the table headers") == "table_header"
    # and a request about the document's headings still means those
    assert _guess_target("make every heading bold") == "heading"
    assert _guess_target("make the table borders thicker") == "table"


def test_formatting_the_table_headers_leaves_the_document_headings_alone():
    from app.docos.actions import validate_batch

    graph = parse_docx_bytes(_table_docx(), title="T")
    batch = validate_batch({"actions": [
        {"type": "format", "target": "table_header", "style": {"bold": True}}]})
    result = ExecutionEngine().execute(graph, batch)

    def bold_of(g):
        return {n.content: n.style.bold for n in g.nodes() if n.content}

    before, after = bold_of(graph), bold_of(result.graph)
    assert after["Activity"] is True and after["Precision"] is True
    assert before["Activity"] is not True, "the cells were not already bold in the model"
    # Word's Heading 1 style is bold, so the point is that the action did not
    # touch it — not that it is unbold.
    assert after["III. RESULTS"] == before["III. RESULTS"], "the paper's own heading was left alone"


# ── "nothing changed" is two different things ───────────────────────────────

def test_a_request_the_document_already_satisfies_says_so():
    """Every heading already bold is not the same as no headings at all, and
    reporting both as "nothing matched" made a carried-out instruction look
    like a failed one."""
    from app.docos.execution import ExecutionEngine

    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.HEADING, content="I. INTRODUCTION", style=Style(bold=True)),
        Node(type=NodeType.BODY, content="Some prose."),
    ]))
    batch = validate_batch({"actions": [
        {"type": "format", "target": "heading", "style": {"bold": True}}]})

    reached = ExecutionEngine().scope_of(graph, batch.actions[0])
    assert reached, "the heading was found; there was simply nothing to do to it"

    empty = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.BODY, content="No headings here."),
    ]))
    assert not ExecutionEngine().scope_of(empty, batch.actions[0]), "nothing matched"


def test_naming_a_heading_names_a_place():
    """"Make the contributions in introduction bullet points" names a place by
    its heading rather than by saying the word "section". Requiring the word
    sent the instruction to the whole document — thirteen rewrite passes to
    change one section, twelve of them for nothing."""
    from app.docos.command.locate import locate_section

    brief = {"sections": [
        {"heading": "I. INTRODUCTION", "about": "motivates automated fall detection",
         "node_ids": ["a", "b"]},
        {"heading": "IV. Evaluation Protocol", "about": "describes the validation",
         "node_ids": ["c", "d"]},
    ]}

    assert locate_section("make the contributions in introduction bullet points",
                          brief)["heading"] == "I. INTRODUCTION"
    assert locate_section("tighten the evaluation protocol",
                          brief)["heading"] == "IV. Evaluation Protocol"


def test_a_chance_word_in_a_summary_is_not_a_place():
    """Without a word like "section", only a heading will do: a word that
    happens to appear in what a section was read to be about is too thin a
    reason to send an instruction there and nowhere else."""
    from app.docos.command.locate import locate_section

    brief = {"sections": [
        {"heading": "I. INTRODUCTION", "about": "motivates automated fall detection",
         "node_ids": ["a"]},
    ]}
    # "detection" appears only in the reading, and the request names no place.
    assert locate_section("improve the detection wording", brief) is None
    # ...but saying "the part about" does name one.
    assert locate_section("improve the part about detection", brief) is not None


# ── showing the maths, rather than rewriting it ─────────────────────────────

def test_asking_to_see_the_equations_is_a_display_change():
    """"Convert the equations into readable mathematics" asks to see them set
    properly. Rewriting the words instead cost a model call a page, turned
    notation into prose, and could not be undone by looking at the result."""
    from app.docos.service import _wants_maths_drawn

    assert _wants_maths_drawn("convert all latex equations into readable formatted mathematics")
    assert _wants_maths_drawn("render the equations properly")
    assert _wants_maths_drawn("show the maths as maths")

    # Mentioning equations is not the same as asking to look at them.
    assert not _wants_maths_drawn("delete every equation")
    assert not _wants_maths_drawn("number all the equations")
    assert not _wants_maths_drawn("make every heading bold")


def test_drawing_the_maths_changes_no_words():
    from app.docos.actions import validate_batch

    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.BODY, content=r"Let $m_i = \|a_i\|_2$ denote the magnitude."),
    ]))
    original = graph.root.children[0].content

    batch = validate_batch({"actions": [{"type": "render_maths", "params": {"on": True}}]})
    result = ExecutionEngine().execute(graph, batch)

    assert result.graph.root.metadata["render_maths"] is True
    assert result.graph.root.children[0].content == original, "the LaTeX is still the LaTeX"


def test_an_equation_on_its_own_line_is_not_a_heading():
    """It is short, unpunctuated and set apart from the text — every signal a
    heading gives — so the outline gained entries like "m_i = ||a_i||_2"."""
    from app.docos.parser.docx_parser import _infer_structure

    emphatic = Style(bold=True, font_size=11)
    assert _infer_structure(r"m_i = ||a_i||_2 = sqrt(a_{i,x}^2)", emphatic, 10) is None
    assert _infer_structure(r"$m_i = \|\mathbf{a}_i\|_2$", emphatic, 10) is None

    # and the headings around it are still headings
    assert _infer_structure("B. Feature Extraction", emphatic, 10) is NodeType.SUBHEADING
    assert _infer_structure("I. INTRODUCTION", Style(bold=True, font_size=12), 10) is NodeType.HEADING


# ── the parts a paper names inside the paragraph ────────────────────────────

def _paper_with_an_abstract() -> DocumentGraph:
    return DocumentGraph(title="A Study", root=Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.BODY, content="A Lightweight Classifier for Fall Detection"),
        Node(type=NodeType.BODY, content="Abstract— This study evaluates subject-independent protocols.",
             metadata={"role": "abstract"}),
        Node(type=NodeType.BODY, content="Index Terms— fall detection, evaluation.",
             metadata={"role": "keywords"}),
        Node(type=NodeType.HEADING, content="I. INTRODUCTION"),
        Node(type=NodeType.BODY, content="Falls are a major cause of injury."),
    ]))


def test_an_abstract_is_a_part_of_the_document_even_without_a_heading():
    """An abstract announces itself with its first word rather than with a
    heading above it, so a document cut only at headings had no abstract in it
    and "rewrite the abstract" had nothing to aim at."""
    from app.docos.command.brief import document_brief

    headings = [s["heading"] for s in document_brief(_paper_with_an_abstract())["sections"]]
    assert "Abstract" in headings
    assert "Keywords" in headings
    assert "I. INTRODUCTION" in headings


def test_an_instruction_about_the_abstract_reaches_it():
    from app.docos.command.brief import document_brief
    from app.docos.command.locate import locate_section

    brief = document_brief(_paper_with_an_abstract())
    found = locate_section("rewrite the abstract to be more concise", brief)

    assert found and found["heading"] == "Abstract"
    assert found["node_ids"], "and it knows which paragraph that is"


def test_the_abstract_is_its_own_first_paragraph():
    """It is not a heading with content beneath it — the paragraph that names
    it is the content."""
    from app.docos.command.brief import document_brief

    brief = document_brief(_paper_with_an_abstract())
    abstract = next(s for s in brief["sections"] if s["heading"] == "Abstract")
    assert abstract["holds"].get("paragraphs") == 1


def test_a_first_row_is_a_table_header_even_before_anyone_bolds_it():
    """Requiring the row to be bold already made it undetectable in the one
    case someone asks about it: when they want it made bold."""
    import io
    from docx import Document

    d = Document()
    table = d.add_table(rows=2, cols=2)
    for col, text in enumerate(["Activity", "Precision"]):
        table.cell(0, col).text = text          # plain text, not bold
    table.cell(1, 0).text = "Walking"
    table.cell(1, 1).text = "0.94"
    buf = io.BytesIO(); d.save(buf)

    graph = parse_docx_bytes(buf.getvalue(), title="T")
    assert [c.content for c in graph.resolve_target("table_header")] == ["Activity", "Precision"]


def test_a_one_row_table_has_no_header():
    """A single-row table is a layout device — usually an equation with its
    number beside it — and its cells are not column headings."""
    import io
    from docx import Document

    d = Document()
    table = d.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "x = y + 1"
    table.cell(0, 1).text = "(3)"
    buf = io.BytesIO(); d.save(buf)

    graph = parse_docx_bytes(buf.getvalue(), title="T")
    assert graph.resolve_target("table_header") == []


# ── formatting words, not the paragraph around them ─────────────────────────

def test_bolding_a_word_bolds_the_word():
    """"Bold results in the abstract" means the word. Styling the node was all
    an action could do, so it bolded the whole abstract."""
    node = Node(type=NodeType.BODY,
                content="We report results for all subjects; the results are consistent.")

    assert node.style_span("results", Style(bold=True)) == 2
    assert node.style.bold is None, "the paragraph itself is not bold"
    assert node.content == "We report results for all subjects; the results are consistent."

    bolded = [r.text for r in node.inline_runs() if r.style.bold]
    assert bolded == ["results", "results"]


def test_formatting_a_phrase_keeps_the_formatting_around_it():
    node = Node(type=NodeType.BODY, content="A bold lead and results after it",
                runs=[Run(text="A bold lead", style=Style(bold=True)),
                      Run(text=" and results after it")])

    node.style_span("results", Style(italic=True))

    kept = {(r.text, r.style.bold, r.style.italic) for r in node.inline_runs()}
    assert ("A bold lead", True, None) in kept, "the bold lead is still bold"
    assert ("results", None, True) in kept


def test_a_phrase_that_is_not_there_changes_nothing():
    node = Node(type=NodeType.BODY, content="No mention of it here.")
    assert node.style_span("results", Style(bold=True)) == 0
    assert node.runs == []


def test_the_format_action_can_name_the_words_to_format():
    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(type=NodeType.BODY, content="Abstract— We report results for all subjects."),
    ]))
    batch = validate_batch({"actions": [
        {"type": "format", "target": "body", "style": {"bold": True},
         "params": {"find": "results"}}]})

    node = ExecutionEngine().execute(graph, batch).graph.root.children[0]

    assert node.style.bold is None, "the paragraph was not bolded"
    assert [r.text for r in node.inline_runs() if r.style.bold] == ["results"]


def test_bolding_one_word_counts_as_a_change():
    """It changes neither the paragraph's text nor the paragraph's own style,
    only how its pieces are formatted — and comparing those two alone called a
    real edit "nothing changed" and threw it away."""
    from app.docos.service import _changed_node_ids

    before = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(id="p1", type=NodeType.BODY, content="We report results here."),
    ]))
    after = before.model_copy(deep=True)
    after.root.children[0].style_span("results", Style(bold=True))

    assert _changed_node_ids(before, after) == {"p1"}


# ── restoring an old version actually restores it ───────────────────────────

def _versions_with_three_edits(tmp_path):
    from app.docos.versioning import VersionEngine
    from app.docos.versioning.store import VersionStore

    engine = VersionEngine(store=VersionStore(db_path=tmp_path / "v.db"))
    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(id="p", type=NodeType.BODY, content="v0 text"),
    ]))
    engine.init_document("d", "T", graph)

    for text in ("v1 text", "v2 text", "v3 text"):
        current = engine.current_graph("d")
        batch = validate_batch({"actions": [
            {"type": "replace", "target": "body",
             "params": {"find": current.get("p").content, "with": text}}]})
        engine.commit("d", batch, ExecutionEngine().execute(current, batch).graph)
    return engine


def _seq_id(engine, seq: int) -> str:
    return next(v.id for v in engine.store.list_versions("d") if v.seq == seq)


def test_restore_brings_back_the_version_it_names(tmp_path):
    """A version is normally stored as the actions that made it and replayed
    from the nearest snapshot. A restore has no actions — there is no edit that
    turns a document into an older one — so replaying it gave back the document
    it was meant to replace, and restore reported a new version having changed
    nothing."""
    engine = _versions_with_three_edits(tmp_path)
    assert engine.current_graph("d").get("p").content == "v3 text"

    engine.restore("d", _seq_id(engine, 1))
    assert engine.current_graph("d").get("p").content == "v1 text"


def test_restore_keeps_the_history_and_can_be_undone(tmp_path):
    engine = _versions_with_three_edits(tmp_path)
    before = [v.seq for v in engine.history("d")]

    engine.restore("d", _seq_id(engine, 1))

    assert [v.seq for v in engine.history("d")] == before + [4], "history grew, nothing was lost"
    engine.undo("d")
    assert engine.current_graph("d").get("p").content == "v3 text", "and the restore itself undoes"


def test_restore_works_from_a_rewound_pointer(tmp_path):
    """The case that hid the bug: after rewinding to the version being restored,
    committing the current graph happens to give the right answer."""
    engine = _versions_with_three_edits(tmp_path)
    engine.rewind("d", _seq_id(engine, 3))

    engine.restore("d", _seq_id(engine, 1))
    assert engine.current_graph("d").get("p").content == "v1 text"


def test_rewind_moves_the_pointer_without_writing(tmp_path):
    engine = _versions_with_three_edits(tmp_path)
    before = [v.seq for v in engine.history("d")]

    for seq, expected in ((1, "v1 text"), (2, "v2 text"), (0, "v0 text")):
        engine.rewind("d", _seq_id(engine, seq))
        assert engine.current_graph("d").get("p").content == expected

    assert [v.seq for v in engine.history("d")] == before, "rewinding writes nothing"


# ── formatting words a request describes rather than quotes ─────────────────

_ABSTRACT = ("Abstract— The Fusion model achieves 89.01% ± 0.47% accuracy on held-out "
             "subjects, about 1.1 percentage points below the 90.10% obtained when "
             "selection and evaluation share a subject pool.")


def test_a_span_the_model_invents_is_not_formatted():
    """Formatting text that is not there is worse than formatting nothing, and
    a model that retypes a figure instead of copying it produces exactly that."""
    import json
    from app.docos.command.spans import find_spans

    node = Node(type=NodeType.BODY, content=_ABSTRACT)

    class Router:
        def chat(self, messages, **_kw):
            return json.dumps({"spans": [
                {"id": node.id, "text": "89.01% ± 0.47% accuracy"},
                {"id": node.id, "text": "a figure that is not in the passage"},
            ]}), "fake", 0.1

    spans = find_spans([node], "the reported results", router=Router())
    assert [s["text"] for s in spans] == ["89.01% ± 0.47% accuracy"]


def test_described_spans_format_only_those_words():
    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(id="a", type=NodeType.BODY, content=_ABSTRACT),
    ]))
    batch = validate_batch({"actions": [{
        "type": "format", "target": "body", "style": {"bold": True},
        "params": {"describe": "the reported results",
                   "spans": [{"id": "a", "text": "89.01% ± 0.47% accuracy"},
                             {"id": "a", "text": "90.10%"}]},
    }]})

    node = ExecutionEngine().execute(graph, batch).graph.get("a")

    assert node.style.bold is None, "the abstract itself is not bold"
    assert [r.text for r in node.inline_runs() if r.style.bold] == [
        "89.01% ± 0.47% accuracy", "90.10%"]


def test_a_quoted_phrase_still_takes_the_literal_route():
    """Searching for text somebody named is faster than asking, and cannot be
    wrong about what they meant."""
    graph = DocumentGraph(root=Node(type=NodeType.DOCUMENT, children=[
        Node(id="a", type=NodeType.BODY, content="We report results here."),
    ]))
    batch = validate_batch({"actions": [{
        "type": "format", "target": "body", "style": {"bold": True},
        "params": {"find": "results"}}]})

    node = ExecutionEngine().execute(graph, batch).graph.get("a")
    assert [r.text for r in node.inline_runs() if r.style.bold] == ["results"]
