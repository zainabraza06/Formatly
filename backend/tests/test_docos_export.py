"""DocumentGraph → DOCX, which is how the exact view reaches a layout engine.

The editing canvas re-lays the document out in HTML and cannot reproduce Word's
line breaking or pagination. The only way to show the document exactly is to
hand a real file to something that can, built from the *current* graph so it
reflects edits rather than the file as it arrived.
"""
from __future__ import annotations

import base64
import io
import sys
from pathlib import Path

import pytest
from docx import Document as ReadDocument
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.docos.export import graph_to_docx_bytes  # noqa: E402
from app.docos.graph import DocumentGraph, Node, NodeType, Style  # noqa: E402


def _png_data_uri() -> str:
    buf = io.BytesIO()
    Image.new("RGB", (120, 60), "#3366cc").save(buf, format="PNG")
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode()


def _graph(children, page=None) -> DocumentGraph:
    root = Node(type=NodeType.DOCUMENT, metadata={"page": page or {
        "width_in": 8.5, "height_in": 11.0,
        "margin": {"top": 1.0, "right": 1.0, "bottom": 1.0, "left": 1.0},
    }}, children=children)
    return DocumentGraph(root=root, title="T")


def _read(graph) -> ReadDocument:
    data = graph_to_docx_bytes(graph)
    assert data[:2] == b"PK", "must be a real .docx"
    return ReadDocument(io.BytesIO(data))


def _texts(doc) -> list[str]:
    return [p.text for p in doc.paragraphs if p.text.strip()]


# ── content ─────────────────────────────────────────────────────────────────

def test_text_and_headings_survive():
    doc = _read(_graph([
        Node(type=NodeType.HEADING, content="Introduction", metadata={"level": 1}),
        Node(type=NodeType.BODY, content="Body text."),
    ]))
    assert "Introduction" in _texts(doc)
    assert "Body text." in _texts(doc)


def test_a_picture_is_embedded_not_described():
    doc = _read(_graph([
        Node(type=NodeType.FIGURE, children=[
            Node(type=NodeType.IMAGE, content="a figure",
                 metadata={"src": _png_data_uri()}),
        ]),
    ]))
    assert sum(1 for r in doc.part.rels.values() if "image" in r.reltype) == 1


def test_a_picture_with_no_data_leaves_its_caption_behind():
    """A linked image has no bytes to embed; the document should still say
    something stood there."""
    doc = _read(_graph([
        Node(type=NodeType.IMAGE, content="Figure 1. A diagram",
             metadata={"external": True, "linked_to": "file:///x.png"}),
    ]))
    assert "Figure 1. A diagram" in _texts(doc)


def test_a_table_keeps_its_cells():
    doc = _read(_graph([
        Node(type=NodeType.TABLE, children=[
            Node(type=NodeType.TABLE_ROW, children=[
                Node(type=NodeType.TABLE_CELL, content="Case"),
                Node(type=NodeType.TABLE_CELL, content="Result"),
            ]),
            Node(type=NodeType.TABLE_ROW, children=[
                Node(type=NodeType.TABLE_CELL, content="1"),
                Node(type=NodeType.TABLE_CELL, content="pass"),
            ]),
        ]),
    ]))
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["Case", "Result"]
    assert [c.text for c in table.rows[1].cells] == ["1", "pass"]


# ── the formatting that decides the layout ──────────────────────────────────

def test_the_documents_font_reaches_the_file():
    graph = _graph([Node(type=NodeType.BODY, content="text")], page={
        "width_in": 8.5, "height_in": 11.0,
        "margin": {"top": 1.0, "right": 1.0, "bottom": 1.0, "left": 1.0},
        "default_font": "Garamond", "default_size_pt": 13.0,
    })
    doc = _read(graph)
    assert doc.styles["Normal"].font.name == "Garamond"
    assert doc.styles["Normal"].font.size.pt == 13.0


def test_line_spacing_is_written_so_the_layout_matches():
    doc = _read(_graph([
        Node(type=NodeType.BODY, content="double", metadata={"line_spacing": 2.0}),
    ]))
    p = next(p for p in doc.paragraphs if p.text == "double")
    assert p.paragraph_format.line_spacing == 2.0


def test_page_size_and_margins_are_written():
    graph = _graph([Node(type=NodeType.BODY, content="x")], page={
        "width_in": 8.27, "height_in": 11.69,     # A4
        "margin": {"top": 0.8, "right": 0.7, "bottom": 0.8, "left": 0.7},
    })
    section = _read(graph).sections[0]
    assert round(section.page_width.inches, 2) == 8.27
    assert round(section.left_margin.inches, 2) == 0.7


def test_run_style_survives():
    doc = _read(_graph([
        Node(type=NodeType.BODY, content="bold italic",
             style=Style(bold=True, italic=True, font_size=14)),
    ]))
    run = next(p for p in doc.paragraphs if p.text == "bold italic").runs[0]
    assert run.bold and run.italic and run.font.size.pt == 14


# ── structure ───────────────────────────────────────────────────────────────

def test_a_page_break_is_a_real_page_break():
    from docx.oxml.ns import qn
    doc = _read(_graph([
        Node(type=NodeType.BODY, content="one"),
        Node(type=NodeType.PAGE_BREAK, metadata={"breaks": 1}),
        Node(type=NodeType.BODY, content="two"),
    ]))
    breaks = [br for p in doc.paragraphs for br in p._p.iter(qn("w:br"))
              if br.get(qn("w:type")) == "page"]
    assert len(breaks) == 1


def test_a_rule_is_written_as_a_border():
    from docx.oxml.ns import qn
    doc = _read(_graph([Node(type=NodeType.HORIZONTAL_RULE)]))
    assert any(p._p.find(qn("w:pPr")) is not None
               and p._p.find(qn("w:pPr")).find(qn("w:pBdr")) is not None
               for p in doc.paragraphs)


def test_an_empty_graph_still_produces_a_valid_file():
    assert graph_to_docx_bytes(_graph([]))[:2] == b"PK"


@pytest.mark.parametrize("bad", ["data:image/png;base64,!!notbase64!!", "http://x/y.png", ""])
def test_a_broken_image_does_not_cost_the_export(bad):
    doc = _read(_graph([
        Node(type=NodeType.IMAGE, content="caption", metadata={"src": bad}),
        Node(type=NodeType.BODY, content="text after"),
    ]))
    assert "text after" in _texts(doc)
