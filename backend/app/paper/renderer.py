"""PaperSpec JSON → DOCX, driven entirely by the selected stylesheet.

A deterministic executor: it applies exactly what the resolved JSON says, and adds
the structural furniture the stylesheet prescribes — heading numbering, table and
figure caption wording/position, table rules, column count, reference numbering.
Swap the stylesheet and the same JSON renders as IEEE, APA, ACM or a report.
"""
from __future__ import annotations

import re
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.paper import stylesheet as ss
from app.paper.figures import render_figure
from app.paper.references import format_reference
from app.paper.schema import (
    Code, Equation, Figure, Heading, ListBlock, Paragraph as PBlock, PaperSpec, Style, Table,
)
from app.paper.styles import StyleLike, resolve_style
from app.paper.styles.base import StyleSheet

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def render_paper(spec: PaperSpec, out_path: str | Path,
                 assets_dir: Optional[Path] = None,
                 style: StyleLike = None,
                 owner_id: Optional[str] = None) -> Path:
    """Render a PaperSpec to .docx.

    `style` overrides spec.meta.style and may be a built-in id, a user's custom
    style id/name, or a StyleSheet object. `owner_id` scopes custom-style lookup.
    """
    if style is not None or not spec.resolved:
        spec = ss.resolve(spec, style, owner_id)
    sheet = (style if isinstance(style, StyleSheet)
             else resolve_style(spec.meta.style, owner_id))

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    assets = Path(assets_dir) if assets_dir else Path(tempfile.mkdtemp(prefix="paper_figs_"))
    assets.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_page(doc, spec)
    _set_normal(doc, sheet)

    # Title block spans the full width; the body then flows in N columns.
    _title_block(doc, spec, sheet)

    if sheet.page.columns > 1:
        body = doc.add_section(WD_SECTION.CONTINUOUS)
        _apply_page_to_section(body, spec)
        _set_columns(body, sheet.page.columns, sheet.page.column_spacing_in)

    counters = {"h1": 0, "h2": 0, "h3": 0, "table": 0, "figure": 0, "equation": 0, "code": 0}
    col_w = _column_width_in(spec)

    for block in spec.blocks:
        if isinstance(block, Heading):
            _heading(doc, block, counters, sheet)
        elif isinstance(block, PBlock):
            _styled_paragraph(doc, block.text, block.style or sheet.body)
        elif isinstance(block, ListBlock):
            _list(doc, block, sheet)
        elif isinstance(block, Equation):
            counters["equation"] += 1
            _equation(doc, block, counters["equation"], col_w, sheet)
        elif isinstance(block, Table):
            counters["table"] += 1
            _table(doc, block, counters["table"], sheet)
        elif isinstance(block, Figure):
            # only claim the next figure number if there is actually a figure
            if _figure(doc, block, counters["figure"] + 1, assets, spec, sheet):
                counters["figure"] += 1
        elif isinstance(block, Code):
            counters["code"] += 1
            _code(doc, block, counters["code"], sheet)

    if spec.references:
        _references(doc, spec, sheet)

    doc.save(str(out_path))
    return out_path


# ── page / section setup ────────────────────────────────────────────────────

def _apply_page_to_section(section, spec: PaperSpec) -> None:
    p = spec.meta.page
    section.page_width = Inches(p.width_in)
    section.page_height = Inches(p.height_in)
    section.top_margin = Inches(p.margin_top_in)
    section.bottom_margin = Inches(p.margin_bottom_in)
    section.left_margin = Inches(p.margin_left_in)
    section.right_margin = Inches(p.margin_right_in)


def _setup_page(doc: Document, spec: PaperSpec) -> None:
    for section in doc.sections:
        _apply_page_to_section(section, spec)
        _set_columns(section, 1, spec.meta.page.column_spacing_in)


def _set_columns(section, num: int, spacing_in: float) -> None:
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(max(1, num)))
    cols.set(qn("w:space"), str(int(spacing_in * 1440)))  # twips
    cols.set(qn("w:equalWidth"), "1")


def _column_width_in(spec: PaperSpec) -> float:
    p = spec.meta.page
    usable = p.width_in - p.margin_left_in - p.margin_right_in
    if p.columns <= 1:
        return usable
    return (usable - p.column_spacing_in * (p.columns - 1)) / p.columns


def _set_normal(doc: Document, sheet: StyleSheet) -> None:
    normal = doc.styles["Normal"]
    font = sheet.body.font or "Times New Roman"
    normal.font.name = font
    normal.font.size = Pt(sheet.body.size_pt or 11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)


# ── style application ───────────────────────────────────────────────────────

def _apply_run(run, style: Style) -> None:
    font = style.font or "Times New Roman"
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)

    if style.size_pt:
        run.font.size = Pt(style.size_pt)
    if style.bold is not None:
        run.font.bold = style.bold
    if style.italic is not None:
        run.font.italic = style.italic
    if style.underline is not None:
        run.font.underline = style.underline
    if style.small_caps:
        run.font.small_caps = True
    if style.all_caps:
        run.font.all_caps = True
    if style.color:
        try:
            run.font.color.rgb = RGBColor.from_string(style.color.lstrip("#").upper())
        except Exception:
            pass


def _apply_para(p, style: Style) -> None:
    pf = p.paragraph_format
    if style.alignment:
        pf.alignment = _ALIGN.get(style.alignment)
    if style.space_before_pt is not None:
        pf.space_before = Pt(style.space_before_pt)
    if style.space_after_pt is not None:
        pf.space_after = Pt(style.space_after_pt)
    if style.line_spacing is not None:
        pf.line_spacing = style.line_spacing
    if style.left_indent_in is not None:
        pf.left_indent = Inches(style.left_indent_in)
    if style.hanging_indent_in:
        pf.first_line_indent = Inches(-style.hanging_indent_in)
    elif style.first_line_indent_in is not None:
        pf.first_line_indent = Inches(style.first_line_indent_in)
    if style.keep_with_next is not None:
        pf.keep_with_next = style.keep_with_next


# Real Word heading styles, so generated documents carry semantic structure:
# navigation pane, table of contents, accessibility, and correct re-parsing.
_WORD_HEADING = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}


# **bold** / __bold__ / *italic* / _italic_
_EMPHASIS = re.compile(r"\*\*(.+?)\*\*|__(.+?)__|\*(.+?)\*|_(.+?)_", re.DOTALL)


def _add_text(p, text: str, style: Style) -> None:
    """Write text, turning markdown emphasis into real runs.

    Models routinely italicise with asterisks — unavoidably so in styles whose
    reference format calls for an italic title. Written verbatim those land in
    the document as literal `*` characters, so they become formatting instead.
    """
    pos = 0
    for m in _EMPHASIS.finditer(text):
        if m.start() > pos:
            _apply_run(p.add_run(text[pos:m.start()]), style)
        bold = m.group(1) is not None or m.group(2) is not None
        inner = next(g for g in m.groups() if g is not None)
        _apply_run(p.add_run(inner),
                   style.merged(Style(bold=True) if bold else Style(italic=True)))
        pos = m.end()
    if pos < len(text):
        _apply_run(p.add_run(text[pos:]), style)


def _styled_paragraph(doc: Document, text: str, style: Style,
                      word_style: Optional[str] = None, rich: bool = True):
    p = doc.add_paragraph()
    if word_style:
        try:
            p.style = doc.styles[word_style]
        except KeyError:
            pass  # unusual template without the built-in styles
    # A justified line that ends in a manual break still gets stretched to the
    # full column width, and a line holding one unbreakable token has no gaps to
    # absorb the slack — it ends up flush right. Left-align multi-line text.
    if text and "\n" in text and style.alignment == "justify":
        style = style.merged(Style(alignment="left"))
    _apply_para(p, style)
    if text:
        # code is never markdown: underscores there are identifiers, not emphasis
        if rich:
            _add_text(p, text, style)
        else:
            _apply_run(p.add_run(text), style)
    return p


def _as_heading(style: Style) -> Style:
    """Word's built-in Heading styles are blue by default; our stylesheets decide
    colour, so pin it to black unless the sheet asks for something else."""
    return style if style.color else style.merged(Style(color="#000000"))


# ── title block ─────────────────────────────────────────────────────────────

def _title_block(doc: Document, spec: PaperSpec, sheet: StyleSheet) -> None:
    m = spec.meta
    _styled_paragraph(doc, m.title, sheet.title)

    for author in m.authors:
        _styled_paragraph(doc, author.name, sheet.author)
        detail = ", ".join(x for x in (author.affiliation, author.email) if x)
        if detail:
            _styled_paragraph(doc, detail, sheet.affiliation)

    if m.abstract:
        if sheet.abstract_as_heading:
            _styled_paragraph(doc, "Abstract", _as_heading(sheet.heading1),
                              word_style=_WORD_HEADING[1])
            _styled_paragraph(doc, m.abstract, sheet.abstract)
        else:
            p = doc.add_paragraph()
            _apply_para(p, sheet.abstract)
            if sheet.abstract_lead:
                _apply_run(p.add_run(sheet.abstract_lead),
                           sheet.abstract.merged(Style(bold=True, italic=True)))
            _apply_run(p.add_run(m.abstract), sheet.abstract)

    if m.keywords:
        p = doc.add_paragraph()
        _apply_para(p, sheet.keywords)
        if sheet.keywords_lead:
            _apply_run(p.add_run(sheet.keywords_lead),
                       sheet.keywords.merged(Style(bold=True, italic=True)))
        _apply_run(p.add_run(", ".join(m.keywords) + "."), sheet.keywords)


# ── blocks ──────────────────────────────────────────────────────────────────

def _heading(doc: Document, block: Heading, counters: dict[str, int], sheet: StyleSheet) -> None:
    level = max(1, min(3, block.level))
    if level == 1:
        counters["h1"] += 1
        counters["h2"] = counters["h3"] = 0
    elif level == 2:
        counters["h2"] += 1
        counters["h3"] = 0
    else:
        counters["h3"] += 1

    label = ss.heading_label(sheet, level, counters)
    style = _as_heading(block.style or sheet.heading_style(level))
    _styled_paragraph(doc, f"{label}{block.text}", style, word_style=_WORD_HEADING.get(level))


def _list(doc: Document, block: ListBlock, sheet: StyleSheet) -> None:
    style = block.style or sheet.list_item
    for i, item in enumerate(block.items, start=1):
        bullet = f"{i}) " if block.ordered else "• "
        _styled_paragraph(doc, f"{bullet}{item}", style)


def _equation(doc: Document, block: Equation, number: int, col_w: float, sheet: StyleSheet) -> None:
    style = block.style or sheet.equation
    p = doc.add_paragraph()
    _apply_para(p, style)
    if block.numbered:
        p.paragraph_format.tab_stops.add_tab_stop(Inches(col_w), WD_TAB_ALIGNMENT.RIGHT)
    _apply_run(p.add_run(block.text), style)
    if block.numbered:
        _apply_run(p.add_run(f"\t({number})"), style.merged(Style(italic=False)))


def _code_caption(doc: Document, block: Code, number: int, sheet: StyleSheet) -> None:
    """"Listing 3. regime_scalars.py — constrained autocorrelation." The filename
    matters to a reader who has to find the code again in their own project."""
    cap_style = block.caption_style or sheet.code_caption or sheet.figure_caption
    label = sheet.code_caption_prefix.format(num=number)
    p = doc.add_paragraph()
    _apply_para(p, cap_style)
    _apply_run(p.add_run(label), cap_style)

    title = " — ".join(x for x in (block.filename, block.caption) if x)
    if title:
        _apply_run(p.add_run(sheet.code_caption_separator + title),
                   _caption_title_style(cap_style, sheet))


def _code(doc: Document, block: Code, number: int, sheet: StyleSheet) -> None:
    """A listing is written one paragraph per line so it never reflows, shaded so
    it reads as a block, and held together so a column break cannot split it
    down the middle."""
    style = block.style or sheet.code
    lines = (block.text or "").splitlines() or [""]

    if sheet.code_caption_position == "above":
        _code_caption(doc, block, number, sheet)

    for i, line in enumerate(lines):
        # only the last line may be followed by a break, and only the first
        # carries the space above — the rest sit flush as one visual block
        edges = Style(
            space_before_pt=style.space_before_pt if i == 0 else 0,
            space_after_pt=style.space_after_pt if i == len(lines) - 1 else 0,
            keep_with_next=i < len(lines) - 1,
        )
        p = _styled_paragraph(doc, line, style.merged(edges), rich=False)
        if sheet.code_background:
            _shade_paragraph(p, sheet.code_background.lstrip("#"))

    if sheet.code_caption_position == "below":
        _code_caption(doc, block, number, sheet)


def _shade_paragraph(p, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    p.paragraph_format.element.get_or_add_pPr().append(shd)


# ── tables ──────────────────────────────────────────────────────────────────

def _caption_title_style(cap_style: Style, sheet: StyleSheet) -> Style:
    """The caption's title text: never small-caps, italic where the style says so,
    and un-bolded when the label itself carries the bold (APA)."""
    override = Style(small_caps=False)
    if sheet.caption_title_italic:
        override = Style(small_caps=False, italic=True, bold=False)
    return cap_style.merged(override)


def _table_caption(doc: Document, block: Table, number: int, sheet: StyleSheet) -> None:
    cap_style = block.caption_style or sheet.table_caption
    label = sheet.table_caption_prefix.format(num=ss.number_for(sheet, number, "table"))
    p = doc.add_paragraph()
    _apply_para(p, cap_style)
    _apply_run(p.add_run(label), cap_style)
    if block.caption:
        _apply_run(p.add_run(sheet.table_caption_separator + block.caption),
                   _caption_title_style(cap_style, sheet))


def _table(doc: Document, block: Table, number: int, sheet: StyleSheet) -> None:
    if sheet.table_caption_position == "above":
        _table_caption(doc, block, number, sheet)

    cols = block.columns or (block.rows[0] if block.rows else [])
    if cols:
        table = doc.add_table(rows=1, cols=len(cols))
        table.autofit = True
        _table_borders(table, sheet)

        header_style = block.header_style or sheet.table_header
        cell_style = block.cell_style or sheet.table_cell

        for i, name in enumerate(cols):
            cp = table.rows[0].cells[i].paragraphs[0]
            _apply_para(cp, header_style)
            _apply_run(cp.add_run(str(name)), header_style)
            if sheet.table_header_fill:
                _shade_cell(table.rows[0].cells[i], sheet.table_header_fill.lstrip("#"))

        for row in block.rows:
            cells = table.add_row().cells
            for i in range(len(cols)):
                value = str(row[i]) if i < len(row) else ""
                cp = cells[i].paragraphs[0]
                _apply_para(cp, cell_style)
                _apply_run(cp.add_run(value), cell_style)

    if sheet.table_caption_position == "below":
        _table_caption(doc, block, number, sheet)

    _styled_paragraph(doc, "", Style(size_pt=4, space_after_pt=6))


def _table_borders(table, sheet: StyleSheet) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")

    if sheet.table_borders == "none":
        edges = {e: None for e in ("top", "bottom", "left", "right", "insideH", "insideV")}
    elif sheet.table_borders == "grid":
        edges = {e: "4" for e in ("top", "bottom", "left", "right", "insideH", "insideV")}
    else:  # horizontal rules only (academic convention)
        edges = {"top": "8", "bottom": "8", "insideH": "4",
                 "left": None, "right": None, "insideV": None}

    for edge, size in edges.items():
        el = OxmlElement(f"w:{edge}")
        if size is None:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def _shade_cell(cell, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


# ── figures ─────────────────────────────────────────────────────────────────

def _figure_caption(doc: Document, block: Figure, number: int, sheet: StyleSheet) -> None:
    cap_style = block.caption_style or sheet.figure_caption
    label = sheet.figure_caption_prefix.format(num=number)
    p = doc.add_paragraph()
    _apply_para(p, cap_style)
    _apply_run(p.add_run(label), cap_style)
    if block.caption:
        _apply_run(p.add_run(sheet.figure_caption_separator + block.caption),
                   _caption_title_style(cap_style, sheet))


def _figure(doc: Document, block: Figure, number: int, assets: Path,
            spec: PaperSpec, sheet: StyleSheet) -> bool:
    """Write the figure and its caption. Returns False — writing nothing — when
    there is no image to show: a caption reading "Fig. 2." above a blank box (or
    above nothing at all) is worse for the reader than no figure, and it would
    also consume a figure number the surviving figures need."""
    image: Optional[Path] = None
    if block.chart:
        try:
            image = render_figure(block.chart, assets / f"fig_{number}.png")
        except Exception:
            image = None
    elif block.image_path and Path(block.image_path).exists():
        image = Path(block.image_path)

    if not (image and image.exists()):
        return False

    if sheet.figure_caption_position == "above":
        _figure_caption(doc, block, number, sheet)

    p = doc.add_paragraph()
    _apply_para(p, block.style or sheet.figure_body)
    pg = spec.meta.page
    width = _column_width_in(spec) if block.span == "column" else (
        pg.width_in - pg.margin_left_in - pg.margin_right_in
    )
    p.add_run().add_picture(str(image), width=Inches(max(1.0, width - 0.1)))

    if sheet.figure_caption_position == "below":
        _figure_caption(doc, block, number, sheet)
    return True


# ── references ──────────────────────────────────────────────────────────────

def _references(doc: Document, spec: PaperSpec, sheet: StyleSheet) -> None:
    _styled_paragraph(doc, sheet.references_title, _as_heading(sheet.references_heading),
                      word_style=_WORD_HEADING[1])
    entries = [r for r in (format_reference(ref) for ref in spec.references) if r]
    for i, ref in enumerate(entries, start=1):
        text = f"[{i}] {ref}" if sheet.number_references else ref
        _styled_paragraph(doc, text, sheet.reference)
