"""
DOCX generation engine.
Produces professionally formatted Word documents:
  - Centered title page with subtitle and date
  - Table of contents with dotted leaders
  - Justified body text
  - Proper Times New Roman / Calibri heading hierarchy
  - Correct line spacing via XML (python-docx's paragraph_format is unreliable)
  - Embedded chart images
"""
from __future__ import annotations

import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.schemas import ChartSpec, DocumentSection
from app.services.presets import StyleConfig, get_preset, pt


# ── helpers ───────────────────────────────────────────────────────────────────

def _set_margins(doc: Document, margin_in: float) -> None:
    for section in doc.sections:
        section.top_margin    = Inches(margin_in)
        section.bottom_margin = Inches(margin_in)
        section.left_margin   = Inches(margin_in)
        section.right_margin  = Inches(margin_in)


def _set_para_spacing(p, line_spacing: float) -> None:
    """Apply line spacing via raw XML (reliable across all viewers)."""
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    val = str(int(line_spacing * 240))   # 240 = single; 480 = double
    spacing.set(qn("w:line"), val)
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:after"), "120")    # 6 pt after paragraph
    # remove any previous spacing element
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    pPr.append(spacing)


def _set_normal_style(doc: Document, *, font_name: str, body_pt: int,
                      line_spacing: float) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = pt(body_pt)
    normal.font.color.rgb = RGBColor(0x11, 0x18, 0x27)


def _set_heading_styles(doc: Document, *, font_name: str,
                        heading1_pt: int, heading2_pt: int,
                        heading_bold: bool = True) -> None:
    for name, size, color in [
        ("Heading 1", heading1_pt, RGBColor(0x1e, 0x29, 0x3b)),
        ("Heading 2", heading2_pt, RGBColor(0x37, 0x41, 0x51)),
    ]:
        try:
            style = doc.styles[name]
            style.font.name  = font_name
            style.font.size  = pt(size)
            style.font.bold  = heading_bold
            style.font.color.rgb = color
            style.paragraph_format.space_before = pt(14)
            style.paragraph_format.space_after  = pt(6)
            style.paragraph_format.keep_with_next = True
        except Exception:
            continue


def _add_run(p, text: str, *, bold=False, italic=False, size_pt: int | None = None,
             font_name: str | None = None, color: RGBColor | None = None) -> Any:
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size_pt:
        run.font.size = pt(size_pt)
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run


# ── title page ────────────────────────────────────────────────────────────────

def _add_title_page(doc: Document, title: str, font_name: str,
                    subtitle: str = "") -> None:
    # Vertical padding — push content to ~1/3 down the page
    for _ in range(7):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = pt(0)

    # ── Main title ──────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.bold       = True
    run.font.name  = font_name
    run.font.size  = pt(28)
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    p_title.paragraph_format.space_after = pt(12)

    # ── Decorative rule ─────────────────────────────────
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_rule = p_rule.add_run("─" * 40)
    run_rule.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    run_rule.font.size = pt(14)
    p_rule.paragraph_format.space_after = pt(12)

    # ── Subtitle / document type ─────────────────────────
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(subtitle)
        run_sub.font.name  = font_name
        run_sub.font.size  = pt(14)
        run_sub.italic     = True
        run_sub.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        p_sub.paragraph_format.space_after = pt(24)

    # ── Date ─────────────────────────────────────────────
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(datetime.date.today().strftime("%B %Y"))
    run_date.font.name  = font_name
    run_date.font.size  = pt(11)
    run_date.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    # ── Bottom padding then page break ────────────────────
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = pt(0)

    doc.add_page_break()


# ── table of contents ─────────────────────────────────────────────────────────

def _add_toc(doc: Document, outline: list[str], font_name: str) -> None:
    p_h = doc.add_paragraph("Table of Contents")
    p_h.style = "Heading 1"
    p_h.paragraph_format.space_after = pt(10)

    for i, item in enumerate(outline, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.2)
        p.paragraph_format.space_before = pt(3)
        p.paragraph_format.space_after  = pt(3)

        # Number + title
        run_num = p.add_run(f"{i}. ")
        run_num.font.name = font_name
        run_num.font.size = pt(11)
        run_num.bold      = True

        run_text = p.add_run(item)
        run_text.font.name = font_name
        run_text.font.size = pt(11)

        # Tab + dotted leader + page placeholder
        run_dots = p.add_run(" " + "." * max(0, 55 - len(item)) + " ")
        run_dots.font.color.rgb = RGBColor(0xb0, 0xb8, 0xc8)
        run_dots.font.size = pt(10)

        run_pg = p.add_run(str(i + 1))
        run_pg.font.name = font_name
        run_pg.font.size = pt(11)
        run_pg.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_page_break()


# ── body text ─────────────────────────────────────────────────────────────────

def _add_body_paragraph(doc: Document, text: str, font_name: str,
                        body_pt: int, line_spacing: float) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = font_name
        run.font.size = pt(body_pt)
    _set_para_spacing(p, line_spacing)


# ── chart image ───────────────────────────────────────────────────────────────

def _add_chart_image(doc: Document, png_path: Path,
                     caption: str = "", font_name: str = "Calibri") -> None:
    doc.add_paragraph()
    try:
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        run.add_picture(str(png_path), width=Inches(5.8))
    except Exception:
        doc.add_paragraph(f"[Chart: {png_path.name}]")

    if caption:
        p_cap = doc.add_paragraph(caption)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_cap.runs:
            run.italic = True
            run.font.size = pt(9)
            run.font.name = font_name
            run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)


# ── main builder ──────────────────────────────────────────────────────────────

def build_docx(
    *,
    out_path: Path,
    title: str,
    outline: list[str],
    sections: list[DocumentSection],
    style_preset: str,
    extracted_rules: dict[str, Any],
    template_style: Optional[dict[str, Any]] = None,
    chart_specs: Optional[list[ChartSpec]] = None,
    chart_pngs: Optional[list[Path]] = None,
    include_title_page: bool = True,
    include_toc: bool = True,
) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    ts     = template_style or {}
    preset: StyleConfig = get_preset(style_preset)  # type: ignore[arg-type]

    font_name    = ts.get("font_name") or extracted_rules.get("font_name") or preset.font_name
    line_spacing = float(ts.get("line_spacing") or extracted_rules.get("line_spacing") or preset.line_spacing)
    margin_in    = float(ts.get("margin_in") or extracted_rules.get("margin_in") or preset.margin_in)
    heading1_pt  = int(ts.get("heading1_pt") or extracted_rules.get("heading_size_pt") or preset.heading1_pt)
    heading2_pt  = int(ts.get("heading2_pt") or max(heading1_pt - 2, 10))
    heading_bold = bool(ts.get("heading_bold", extracted_rules.get("heading_bold", True)))
    body_pt      = preset.body_pt

    doc = Document()
    _set_margins(doc, margin_in)
    _set_normal_style(doc, font_name=font_name, body_pt=body_pt, line_spacing=line_spacing)
    _set_heading_styles(doc, font_name=font_name, heading1_pt=heading1_pt,
                        heading2_pt=heading2_pt, heading_bold=heading_bold)

    # ── Title page ──────────────────────────────────────
    if include_title_page:
        _add_title_page(doc, title, font_name,
                        subtitle=style_preset.replace("_", " ").title())

    # ── Table of contents ────────────────────────────────
    if include_toc and outline:
        _add_toc(doc, outline, font_name)

    # ── Sections ─────────────────────────────────────────
    for sec in sections:
        p_head = doc.add_paragraph(sec.heading)
        p_head.style = "Heading 1"

        for para_text in (sec.content or "").split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            _add_body_paragraph(doc, para_text, font_name, body_pt, line_spacing)

        # gap after section
        doc.add_paragraph().paragraph_format.space_after = pt(6)

    # ── Charts ───────────────────────────────────────────
    if chart_pngs:
        doc.add_page_break()
        p_charts_h = doc.add_paragraph("Charts & Visualizations")
        p_charts_h.style = "Heading 1"

        for idx, png in enumerate(chart_pngs):
            if not Path(png).exists():
                continue
            caption = ""
            if chart_specs and idx < len(chart_specs):
                caption = f"Figure {idx + 1}: {chart_specs[idx].title}"
            _add_chart_image(doc, Path(png), caption=caption, font_name=font_name)

    doc.save(str(out_path))
    return out_path
