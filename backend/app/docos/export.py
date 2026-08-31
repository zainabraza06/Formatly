"""DocumentGraph → DOCX.

The editor lays the document out with HTML and CSS, which cannot reproduce
Word's line breaking, hyphenation or pagination — close, but not the document.
An exact view has to come from a real layout engine, and the only way to reach
one is to hand it a real file.

So this writes the *current* graph back to DOCX, edits and all. Rendering the
originally imported bytes would be easier and would be a lie the moment anything
changed.
"""
from __future__ import annotations

import base64
import binascii
import io
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.docos.graph import DocumentGraph, Node, NodeType

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_HEADING_LEVEL = {NodeType.HEADING: 1, NodeType.SUBHEADING: 2}

# Wider than any sensible column; the picture is clamped to the text width.
_MAX_IMAGE_IN = 6.5


def graph_to_docx_bytes(graph: DocumentGraph) -> bytes:
    doc = Document()
    page = (graph.root.metadata or {}).get("page") or {}
    _apply_page(doc, page)
    _apply_defaults(doc, page)

    for node in graph.root.children:
        _write(doc, node, page)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── page setup ──────────────────────────────────────────────────────────────

def _apply_page(doc: Document, page: dict[str, Any]) -> None:
    for section in doc.sections:
        section.page_width = Inches(float(page.get("width_in") or 8.5))
        section.page_height = Inches(float(page.get("height_in") or 11.0))
        margin = page.get("margin") or {}
        section.top_margin = Inches(float(margin.get("top", 1.0)))
        section.right_margin = Inches(float(margin.get("right", 1.0)))
        section.bottom_margin = Inches(float(margin.get("bottom", 1.0)))
        section.left_margin = Inches(float(margin.get("left", 1.0)))


def _apply_defaults(doc: Document, page: dict[str, Any]) -> None:
    """The document's own typeface, so anything without an explicit style still
    comes out in the right face rather than python-docx's."""
    normal = doc.styles["Normal"]
    font = page.get("default_font") or ""
    size = page.get("default_size_pt")
    if font:
        normal.font.name = font
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font)
    if size:
        normal.font.size = Pt(float(size))


def _text_width_in(page: dict[str, Any]) -> float:
    margin = page.get("margin") or {}
    width = float(page.get("width_in") or 8.5)
    return max(1.0, width - float(margin.get("left", 1.0)) - float(margin.get("right", 1.0)))


# ── nodes ───────────────────────────────────────────────────────────────────

def _write(doc: Document, node: Node, page: dict[str, Any]) -> None:
    kind = node.type
    meta = node.metadata or {}

    if kind == NodeType.PAGE_BREAK:
        for _ in range(max(1, int(meta.get("breaks", 1) or 1))):
            doc.add_page_break()
        return

    if kind == NodeType.HORIZONTAL_RULE:
        _rule(doc, meta)
        return

    if kind == NodeType.TABLE:
        _table(doc, node, page)
        return

    if kind in (NodeType.FIGURE, NodeType.PARAGRAPH) and node.children:
        # containers: a figure holds its pictures, a paragraph a run of lines
        for child in node.children:
            _write(doc, child, page)
        return

    if kind == NodeType.IMAGE:
        _picture(doc, node, page)
        return

    if kind in (NodeType.HEADER, NodeType.FOOTER):
        return    # written into the section, not the body

    _paragraph(doc, node, page)


def _paragraph(doc: Document, node: Node, page: dict[str, Any]):
    meta = node.metadata or {}
    if meta.get("page_break_before"):
        doc.add_page_break()

    level = _HEADING_LEVEL.get(node.type)
    paragraph = doc.add_paragraph()
    if level:
        try:
            paragraph.style = doc.styles[f"Heading {level}"]
        except KeyError:
            pass

    _apply_list(paragraph, doc, meta.get("list"))
    _apply_paragraph_format(paragraph, node)
    # One Word run per formatted piece, so a bold phrase or a superscript
    # citation comes back out of the file the way it went in.
    if _write_with_equations(paragraph, node, page, bool(level)):
        return paragraph

    for piece in node.inline_runs():
        _apply_run(paragraph.add_run(piece.text), node, page,
                   heading=bool(level), inline=piece.style)
    return paragraph


def _apply_list(paragraph, doc: Document, listing: Any) -> None:
    """Give a list item Word's own list style, so it exports as a real list.

    Word numbers its list styles from the second level up ("List Bullet 2"),
    and a document that nests deeper than the built-in styles go simply sits at
    the deepest one rather than losing its bullet.
    """
    if not isinstance(listing, dict):
        return
    base = "List Number" if listing.get("kind") == "number" else "List Bullet"
    try:
        level = max(0, min(int(listing.get("level") or 0), 2))
    except (TypeError, ValueError):
        level = 0
    for name in (f"{base} {level + 1}" if level else base, base):
        try:
            paragraph.style = doc.styles[name]
            return
        except KeyError:
            continue


def _write_with_equations(paragraph, node: Node, page: dict[str, Any],
                          heading: bool) -> bool:
    """Write a paragraph whose equations are still Word's own. True if it did.

    An equation read from the file kept its original XML, so an equation nobody
    edited goes back exactly as Word wrote it rather than as our reading of it.
    Once the text no longer contains the LaTeX we produced — the assistant was
    asked to turn the equations into something readable, and did — there is no
    equation left to restore and the words are written as words.

    Inline run formatting is not applied to these paragraphs: an equation and a
    run of bold in the same paragraph is rare, and losing the emphasis is a far
    smaller loss than dropping the equation.
    """
    equations = (node.metadata or {}).get("equations") or []
    if not equations:
        return False

    remaining = node.content
    segments: list[tuple[str, Any]] = []
    for equation in equations:
        marker = f"${equation.get('latex', '')}$"
        before, found, after = remaining.partition(marker)
        if not found:
            return False                     # rewritten: no equation to put back
        segments.append(("text", before))
        segments.append(("maths", equation.get("xml", "")))
        remaining = after
    segments.append(("text", remaining))

    from docx.oxml import parse_xml

    for kind, value in segments:
        if kind == "text":
            if value:
                _apply_run(paragraph.add_run(value), node, page, heading=heading)
        else:
            try:
                paragraph._p.append(parse_xml(value))
            except Exception:
                # Unparseable for any reason: the words are better than nothing.
                _apply_run(paragraph.add_run(node.content), node, page, heading=heading)
                return True
    return True


def _apply_paragraph_format(paragraph, node: Node) -> None:
    style = node.style
    meta = node.metadata or {}
    pf = paragraph.paragraph_format

    if style.alignment and style.alignment in _ALIGN:
        pf.alignment = _ALIGN[style.alignment]

    spacing = meta.get("line_spacing")
    if isinstance(spacing, (int, float)):
        pf.line_spacing = Pt(float(spacing)) if meta.get("line_spacing_exact") else float(spacing)
    for key, attr in (("space_before_pt", "space_before"), ("space_after_pt", "space_after"),
                      ("indent_left_pt", "left_indent"), ("indent_right_pt", "right_indent"),
                      ("indent_first_line_pt", "first_line_indent")):
        value = meta.get(key)
        if isinstance(value, (int, float)):
            setattr(pf, attr, Pt(float(value)))


def _apply_run(run, node: Node, page: dict[str, Any], heading: bool = False,
               inline: Optional[Style] = None) -> None:
    # The run's own formatting wins where it states any; everything it leaves
    # unstated it inherits from the paragraph, exactly as Word resolves it.
    style = node.style.merged(inline) if inline is not None else node.style
    font = style.font_family or page.get("default_font") or ""
    if font:
        run.font.name = font
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font)

    size = style.font_size or page.get("default_size_pt")
    if size:
        run.font.size = Pt(float(size))
    if heading and not style.color:
        # Word's built-in heading styles are blue; the document decides colour,
        # so an unstated one is black rather than the template's accent.
        run.font.color.rgb = RGBColor(0, 0, 0)
    # True is written plainly. False is written only when there is something to
    # cancel — a Heading style is bold in Word, and a paragraph can be italic —
    # because writing it everywhere else would state on every run what the file
    # never said, and the document would come back noisier than it went in.
    for attr in ("bold", "italic", "underline"):
        resolved = getattr(style, attr)
        if resolved:
            setattr(run.font, attr, True)
        elif resolved is False and (heading or getattr(node.style, attr)):
            setattr(run.font, attr, False)
    if style.color:
        try:
            run.font.color.rgb = RGBColor.from_string(str(style.color).lstrip("#").upper())
        except (ValueError, AttributeError):
            pass
    if style.vertical_align == "superscript":
        run.font.superscript = True
    elif style.vertical_align == "subscript":
        run.font.subscript = True


def _rule(doc: Document, meta: dict[str, Any]) -> None:
    if meta.get("page_break_before"):
        doc.add_page_break()
    paragraph = doc.add_paragraph()
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    borders.append(bottom)
    pPr.append(borders)


def _picture(doc: Document, node: Node, page: dict[str, Any]) -> None:
    """Place the picture the node carries. A node with no readable image keeps
    its caption, so the document still says something stood here."""
    data = _image_bytes(node)
    if data is None:
        if node.content:
            doc.add_paragraph(node.content)
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = min(_MAX_IMAGE_IN, _text_width_in(page))
    try:
        paragraph.add_run().add_picture(io.BytesIO(data), width=Inches(width))
    except Exception:
        # an unreadable or unsupported image must not cost the whole export
        if node.content:
            paragraph.add_run(node.content)


def _image_bytes(node: Node) -> Optional[bytes]:
    src = (node.metadata or {}).get("src")
    if not isinstance(src, str) or not src.startswith("data:"):
        return None    # linked images live elsewhere; there is nothing to embed
    try:
        return base64.b64decode(src.split(",", 1)[1])
    except (IndexError, binascii.Error, ValueError):
        return None


def _table(doc: Document, node: Node, page: dict[str, Any]) -> None:
    rows = [c for c in node.children if c.type == NodeType.TABLE_ROW]
    if not rows:
        return
    columns = max(len(r.children) for r in rows)
    if columns < 1:
        return

    table = doc.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    _apply_borders(table, (node.metadata or {}).get("borders"))
    for row in rows:
        cells = table.add_row().cells
        for i, cell in enumerate(row.children[:columns]):
            paragraph = cells[i].paragraphs[0]
            _apply_paragraph_format(paragraph, cell)
            # A cell holds a display equation as often as it holds a number,
            # and one nobody edited goes back as Word's own XML.
            if cell.content and not _write_with_equations(paragraph, cell, page, False):
                _apply_run(paragraph.add_run(cell.content), cell, page)
            for picture in cell.children:
                if picture.type == NodeType.IMAGE:
                    _picture_in_cell(cells[i], picture, page)

    for caption in (c for c in node.children if c.type == NodeType.CAPTION):
        doc.add_paragraph(caption.content)


_BORDER_TAGS = (("top", "w:top"), ("bottom", "w:bottom"), ("left", "w:left"),
                ("right", "w:right"), ("inside_h", "w:insideH"),
                ("inside_v", "w:insideV"))


def _apply_borders(table, borders: Any) -> None:
    """Write the table's own edges, so a table with two rules exports as one.

    Word measures a line in eighths of a point and spells "no line" as `none`.
    Without this the exporter always wrote Table Grid, and a table that had
    been given horizontal rules only came back out of the file fully boxed.
    """
    if not isinstance(borders, dict) or not borders:
        return

    properties = table._tbl.tblPr
    for existing in properties.findall(qn("w:tblBorders")):
        properties.remove(existing)

    element = OxmlElement("w:tblBorders")
    for side, tag in _BORDER_TAGS:
        try:
            width = float(borders.get(side, 0) or 0)
        except (TypeError, ValueError):
            width = 0.0
        edge = OxmlElement(tag)
        if width > 0:
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), str(max(2, int(round(width * 8)))))
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), "333333")
        else:
            edge.set(qn("w:val"), "none")
            edge.set(qn("w:sz"), "0")
        element.append(edge)
    properties.append(element)


def _picture_in_cell(cell, node: Node, page: dict[str, Any]) -> None:
    data = _image_bytes(node)
    if data is None:
        return
    try:
        cell.add_paragraph().add_run().add_picture(
            io.BytesIO(data), width=Inches(min(2.5, _text_width_in(page) / 2)))
    except Exception:
        pass
