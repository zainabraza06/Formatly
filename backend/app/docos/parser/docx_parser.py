"""DOCX → DocumentGraph.

Walks the document body in document order (paragraphs and tables interleaved),
classifies each block into a typed node, and aggregates run-level formatting into
a paragraph-level Style. Images, captions, horizontal rules, page breaks and
references are detected heuristically from styles and XML markers.
"""
from __future__ import annotations

import base64
import io
import re
from pathlib import Path
from typing import Any, NamedTuple, Optional

from docx import Document
from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.table import Table as _Table
from docx.text.paragraph import Paragraph as _Paragraph
from lxml import etree

from app.docos.graph import DocumentGraph, Node, NodeType, Run, Style, merge_runs
from app.docos.parser.omml import paragraph_parts

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_HEADING_STYLES = ("heading", "title")
_REFERENCE_HINTS = ("reference", "bibliography", "works cited")

# VML lives outside python-docx's namespace map, so these are spelled in full.
_V_NS = "urn:schemas-microsoft-com:vml"
_O_NS = "urn:schemas-microsoft-com:office:office"
_V_IMAGEDATA = f"{{{_V_NS}}}imagedata"
_V_RECT = f"{{{_V_NS}}}rect"
_O_HR = f"{{{_O_NS}}}hr"

# Images travel inline as data URIs so the editor can show the real picture
# without a second round trip. Past this size that stops being reasonable and
# the node keeps its place in the document without the bytes.
_MAX_INLINE_IMAGE_BYTES = 4 * 1024 * 1024


def parse_docx_bytes(data: bytes, *, title: str = "") -> DocumentGraph:
    return _parse(Document(io.BytesIO(data)), title=title)


def parse_docx(path: str | Path, *, title: str = "") -> DocumentGraph:
    p = Path(path)
    return _parse(Document(str(p)), title=title or p.stem)


# ── core ────────────────────────────────────────────────────────────────────

def _parse(doc: _Doc, *, title: str) -> DocumentGraph:
    root = Node(type=NodeType.DOCUMENT, metadata={"source": "docx"})
    graph = DocumentGraph(root=root, title=title)

    page = _page_geometry(doc)
    default_font, default_size = _document_default_font(doc)
    # Carried so the sheet is laid out in the document's own typeface. Rendering
    # it in the viewer's default is what made an imported file look shifted, and
    # the wrong metrics are what pushed text past the bottom of the page.
    page["default_font"] = default_font or ""
    page["default_size_pt"] = default_size or 11.0
    root.metadata["page"] = page

    # Every paragraph starts from these, so they are read once.
    spacing_defaults = _default_spacing(doc)

    in_references = False
    body = doc.element.body

    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            para = _Paragraph(child, doc)
            node = _paragraph_node(para, in_references, spacing_defaults, default_size)
            if node is None:
                continue
            if node.type == NodeType.HEADING and _looks_like_references(para.text):
                in_references = True
            root.children.append(node)
        elif tag == qn("w:tbl"):
            root.children.append(_table_node(_Table(child, doc), doc))

    _attach_headers_footers(doc, root)
    return graph


def _page_geometry(doc: _Doc) -> dict:
    """Real page size + margins (inches) from the first section, so the editor can
    render the sheet at the document's actual dimensions (A4, Letter, …)."""
    def _in(v, default: float) -> float:
        try:
            return round(float(v.inches), 3)
        except Exception:
            return default

    try:
        sec = doc.sections[0]
    except (IndexError, Exception):  # noqa: B014 - defensive
        return {"width_in": 8.5, "height_in": 11.0,
                "margin": {"top": 1.0, "right": 1.0, "bottom": 1.0, "left": 1.0}}

    w = _in(sec.page_width, 8.5)
    h = _in(sec.page_height, 11.0)
    landscape = str(getattr(sec, "orientation", "")).endswith("LANDSCAPE")
    if landscape and w < h:
        w, h = h, w
    return {
        "width_in": w,
        "height_in": h,
        "landscape": landscape,
        "margin": {
            "top": _in(sec.top_margin, 1.0),
            "right": _in(sec.right_margin, 1.0),
            "bottom": _in(sec.bottom_margin, 1.0),
            "left": _in(sec.left_margin, 1.0),
        },
    }


def _looks_like_references(text: str) -> bool:
    t = (text or "").strip().lower()
    return any(h in t for h in _REFERENCE_HINTS)


def _paragraph_node(para: _Paragraph, in_references: bool,
                    spacing_defaults: Optional[dict] = None,
                    default_size: Optional[float] = None) -> Optional[Node]:
    style_name = (para.style.name if para.style else "") or ""
    lname = style_name.lower()

    # `para.text` returns the w:t runs and nothing else, so a Word equation —
    # stored as OMML beside them — was invisible. A paragraph holding only one
    # looked empty and was dropped, taking the equation out of the document.
    # The parts are read in order and the maths written as LaTeX between
    # dollars, which is how the rest of the document writes maths and what an
    # instruction about "the equations" already means.
    parts = paragraph_parts(para)
    maths = [(latex, element) for kind, latex, element in parts if kind == "maths"]
    text = ("".join(latex if kind == "text" else f"${latex}$"
                    for kind, latex, _element in parts)
            if maths else (para.text or ""))

    # Page-boundary signals. `lastRenderedPageBreak` markers are written by Word
    # and record where each page actually broke on the last render — we use them
    # so the editor's page count matches the real document (not a node count).
    explicit, rendered = _page_break_signals(para)
    breaks_before = (1 if explicit else 0) + rendered

    images = _images(para)
    is_rule = not text.strip() and (_has_bottom_border(para) or _is_vml_rule(para))

    # A rule drawn as a picture, or decoration whose image will not resolve: an
    # empty paragraph with nothing readable in it is not a figure, and showing
    # it as a broken one puts a false alarm in the middle of the text.
    #
    # This is decided before the page-break case on purpose. A rule frequently
    # sits exactly where Word broke the page, and treating that paragraph as a
    # bare break threw the rule away — so the line vanished from precisely the
    # pages a reader would notice. The rule keeps the marker instead.
    if is_rule:
        meta: dict = {"style_name": style_name}
        _apply_breaks(meta, breaks_before)
        return Node(type=NodeType.HORIZONTAL_RULE, metadata=meta)

    # dedicated page-break paragraph (a break with no real content)
    if breaks_before and not text.strip() and not images:
        return Node(type=NodeType.PAGE_BREAK,
                    metadata={"style_name": style_name, "breaks": breaks_before})

    # figure: a paragraph carrying one or more actual pictures
    if images:
        style = _paragraph_style(para)
        children = [
            Node(
                type=NodeType.IMAGE,
                content=text.strip() if len(images) == 1 else "",
                style=style,
                metadata={"style_name": style_name, "is_figure": True, **img},
            )
            for img in images
        ]
        fig = Node(type=NodeType.FIGURE, children=children,
                   metadata={"style_name": style_name})
        _apply_breaks(fig.metadata, breaks_before)
        return fig

    # skip truly empty spacer paragraphs (keep them out of the graph)
    if not text.strip():
        return None

    style = _paragraph_style(para)
    node_type = _classify_paragraph(lname, in_references)
    # The style name is believed when it says something. When it does not — a
    # plain paragraph — the line's own shape is read, so a paper that never
    # applied Word's Heading styles still has headings.
    if node_type is NodeType.BODY and not in_references:
        if _looks_like_caption(text):
            node_type = NodeType.CAPTION
        else:
            inferred = _infer_structure(text, style, default_size)
            if inferred is not None:
                node_type = inferred
    meta: dict = {"style_name": style_name, "level": _heading_level(lname)}
    role = _named_part(text)
    if role:
        meta["role"] = role
    bullet = _list_of(para, style_name)
    if bullet:
        meta["list"] = bullet
    if maths:
        from lxml import etree
        meta["equations"] = [
            {"latex": latex, "xml": etree.tostring(element, encoding="unicode")}
            for latex, element in maths
        ]
    meta.update(_spacing(para, spacing_defaults))
    _apply_breaks(meta, breaks_before)
    return Node(
        type=node_type,
        content=text,
        style=style,
        metadata=meta,
        runs=_paragraph_runs(para),
    )


def _format_spacing(pf) -> dict:
    """What one paragraph format states about spacing, in the units CSS wants."""
    out: dict = {}
    if pf is None:
        return out
    try:
        if pf.line_spacing is not None:
            # a float is a multiple; a Length is an exact height
            out["line_spacing"] = (float(pf.line_spacing)
                                   if isinstance(pf.line_spacing, float)
                                   else round(pf.line_spacing.pt, 1))
            out["line_spacing_exact"] = not isinstance(pf.line_spacing, float)
        if pf.space_before is not None:
            out["space_before_pt"] = round(pf.space_before.pt, 1)
        if pf.space_after is not None:
            out["space_after_pt"] = round(pf.space_after.pt, 1)
        # Indents. A display equation is often just an indented paragraph, and
        # a hanging indent is what makes a reference list look like one.
        for attr, key in (("left_indent", "indent_left_pt"),
                          ("right_indent", "indent_right_pt"),
                          ("first_line_indent", "indent_first_line_pt")):
            value = getattr(pf, attr, None)
            if value is not None:
                out[key] = round(value.pt, 1)
    except Exception:
        pass
    return out


def _default_spacing(doc: _Doc) -> dict:
    """The spacing every paragraph starts from: `docDefaults`.

    Word writes the document's real defaults here — `w:after="200"` is 10pt
    between paragraphs, `w:line="276"` is 1.15 line spacing — and most files
    never state either again. Reading only the paragraph found nothing, so
    pages were filled with tighter, gapless text than the document describes
    and more of it fit than really does.
    """
    try:
        spacing = doc.styles.element.find(
            f"{{{_W}}}docDefaults/{{{_W}}}pPrDefault/{{{_W}}}pPr/{{{_W}}}spacing")
    except Exception:
        return {}
    if spacing is None:
        return {}

    out: dict = {}
    after = spacing.get(qn("w:after"))
    before = spacing.get(qn("w:before"))
    line = spacing.get(qn("w:line"))
    rule = spacing.get(qn("w:lineRule")) or "auto"
    try:
        if after is not None:
            out["space_after_pt"] = round(int(after) / 20, 1)      # twips → pt
        if before is not None:
            out["space_before_pt"] = round(int(before) / 20, 1)
        if line is not None:
            # "auto" is a multiple of single, in 240ths; the others are twips.
            out["line_spacing"] = (round(int(line) / 240, 3) if rule == "auto"
                                   else round(int(line) / 20, 1))
            out["line_spacing_exact"] = rule != "auto"
    except (TypeError, ValueError):
        return {}
    return out


def _spacing(para: _Paragraph, defaults: Optional[dict] = None) -> dict:
    """Line spacing and the gaps around a paragraph, with what it inherits.

    Nearest wins: the paragraph, then the styles it is based on, then the
    document's defaults — the order Word resolves them in.
    """
    resolved: dict = dict(defaults or {})

    chain: list[dict] = []
    style = getattr(para, "style", None)
    seen = 0
    while style is not None and seen < 10:      # guard a cyclic base_style chain
        chain.append(_format_spacing(getattr(style, "paragraph_format", None)))
        style = getattr(style, "base_style", None)
        seen += 1

    for stated in reversed(chain):              # furthest ancestor first
        resolved.update(stated)
    resolved.update(_format_spacing(getattr(para, "paragraph_format", None)))
    return resolved


def _apply_breaks(meta: dict, breaks_before: int) -> None:
    """Encode page-boundary metadata: `page_break_before` starts a new page before
    this node; `extra_pages` is how many *additional* pages the node spans (a long
    paragraph or a table that Word split across pages)."""
    if breaks_before >= 1:
        meta["page_break_before"] = True
    if breaks_before > 1:
        meta["extra_pages"] = breaks_before - 1


# How a numbered heading announces itself. Sub-level patterns are tried first,
# because "1.1" also starts like "1." and "IV-C" like "IV".
# A single letter is a sub-heading — "A. Dataset" — except where it is also a
# section numeral. "I." opens a paper; a section numbered C would be the
# hundredth, so the letters that read both ways are given to the numerals.
_SUB_NUMBER = re.compile(
    r"^(?:\d+\.\d+|[ABCDEFGHJKLMNOPQRSTUWYZ][.)]\s|[IVXLCDM]+-[A-Z0-9]\b)", re.ASCII)
_TOP_NUMBER = re.compile(r"^(?:[IVXLCDM]+|\d+)[.)]\s", re.ASCII)

# Parts a paper names inside the paragraph rather than above it. An abstract
# rarely gets a heading of its own — it announces itself with its first word —
# so a document cut only at its headings has no abstract in it anywhere, and an
# instruction about "the abstract" has nothing to reach.
_NAMED_PARTS = (
    ("abstract", re.compile(r"^abstract\s*[-—–:.]", re.IGNORECASE)),
    ("keywords", re.compile(r"^(index\s+terms|keywords?)\s*[-—–:.]", re.IGNORECASE)),
)


def _named_part(text: str) -> Optional[str]:
    """The part of a paper this paragraph announces itself as, if any."""
    line = (text or "").strip()
    for name, pattern in _NAMED_PARTS:
        if pattern.match(line):
            return name
    return None


def _list_of(para: _Paragraph, style_name: str) -> Optional[dict]:
    """Is this paragraph a list item, and of what kind?

    Word says so two ways and documents use both: a List Bullet or List Number
    style, or an inline numbering reference. Reading only one of them misses
    half the lists in the world.
    """
    lname = (style_name or "").lower()
    kind: Optional[str] = None
    if "list number" in lname:
        kind = "number"
    elif "list bullet" in lname:
        kind = "bullet"

    level = 0
    properties = para._p.find(qn("w:pPr"))
    numbering = properties.find(qn("w:numPr")) if properties is not None else None
    if numbering is not None:
        ilvl = numbering.find(qn("w:ilvl"))
        if ilvl is not None:
            try:
                level = int(ilvl.get(qn("w:val")) or 0)
            except ValueError:
                level = 0
        # A numbering reference without a style to name it: bullet is the
        # commoner of the two and the safer guess, since a wrong bullet reads
        # as a bullet and a wrong number implies an order that is not there.
        kind = kind or "bullet"
    # Deliberately not "List Paragraph": Word gives that style to any indented
    # block, list or not, and treating it as one put a bullet on the sentence
    # that introduces the list. A List Paragraph with real numbering is caught
    # above, by its numbering.

    return {"kind": kind, "level": level} if kind else None


# How a caption announces itself: the thing it captions, its number, and a
# separator. Most authors type a caption as an ordinary paragraph rather than
# applying Word's Caption style, so a document full of captions could contain
# none as far as an instruction about them was concerned.
#
# The separator is what keeps "Figure 1 shows a clear trend" out: that is a
# sentence about a figure, not the caption of one. A caption says "Figure 1."
# or "Fig. 2:" and then describes it.
_CAPTION_LINE = re.compile(
    # An optional bracketed lead-in, because a draft often carries one:
    # "[PLACEHOLDER — Fig. 1: …]".
    r"^\s*(?:[\[(][^\])]{0,40}?[—–:-]\s*)?"
    r"(figure|fig|table|tbl|chart|algorithm|listing|scheme|plate)\.?\s*"
    r"(\d+(?:\.\d+)*|[IVXLCDM]+)\s*[.:—–)-]",
    re.IGNORECASE)


def _looks_like_caption(text: str) -> bool:
    """Is this paragraph a caption, whatever style it was given?"""
    line = (text or "").strip()
    # A caption is a label, not a section of prose. The cap is generous: some
    # journals write a paragraph-long one.
    return bool(line) and len(line) <= 400 and bool(_CAPTION_LINE.match(line))


# The shapes mathematics takes: LaTeX, or the symbols it is written with.
_LOOKS_LIKE_MATHS = re.compile(
    r"\$|\[a-zA-Z]{2,}|[=≈≤≥±×÷∑∏∫√∞]|\|\||_\{|\^\{|[a-zA-Z]_[a-z0-9]\b")

# Things that begin like a heading and are not one.
_NOT_A_HEADING = re.compile(r"^(fig\.|figure|table|eq\.|equation|algorithm|appendix\s+\w+\s*[:.]?\s*\S)",
                            re.IGNORECASE)


def _infer_structure(text: str, style: Style, default_size: Optional[float]) -> Optional[NodeType]:
    """Is this hand-formatted paragraph really a heading?

    Most papers never apply Word's Heading styles. The author types "I.
    INTRODUCTION" in bold at 12pt with the Normal style, and reading the style
    name alone then finds a document with no structure at all — nothing to
    target with "standardise the headings", and nothing to show in an outline.

    So the shape of the line is read instead: short, unpunctuated, set apart by
    weight or size, and often numbered. All of that has to agree, because a
    bold sentence is not a heading and calling it one is worse than missing it.
    """
    line = (text or "").strip()
    if not line or len(line) > 90 or len(line.split()) > 14:
        return None
    if _NOT_A_HEADING.match(line):
        return None
    # An equation on its own line is short, unpunctuated and set apart from the
    # text — every signal a heading gives. It is not a heading, and calling it
    # one puts "m_i = ||a_i||_2" in the document's outline.
    if _LOOKS_LIKE_MATHS.search(line):
        return None

    sub = bool(_SUB_NUMBER.match(line))
    top = bool(_TOP_NUMBER.match(line)) and not sub
    # A heading does not end a sentence. A numbered one may end with a period
    # only because of its own numbering, which the patterns above have matched.
    if line.endswith((".", "?", "!", ";", ",")) and not (sub or top):
        return None

    bigger = bool(style.font_size and default_size and style.font_size >= default_size + 0.5)
    much_bigger = bool(style.font_size and default_size and style.font_size >= default_size + 2)
    shouted = line.isupper() and len(line) > 3

    # Numbering alone is not enough — a numbered list item is numbered too —
    # so the line must also be set apart from the text around it.
    set_apart = bool(style.bold) or bigger or shouted
    if not set_apart:
        return None
    if not (sub or top or bigger or shouted):
        return None

    if sub:
        return NodeType.SUBHEADING
    if top:
        return NodeType.HEADING
    return NodeType.HEADING if (much_bigger or shouted) else NodeType.SUBHEADING


def _classify_paragraph(lname: str, in_references: bool) -> NodeType:
    if lname.startswith("title"):
        return NodeType.HEADING
    if lname.startswith("heading"):
        lvl = _heading_level(lname)
        return NodeType.HEADING if lvl <= 1 else NodeType.SUBHEADING
    if "caption" in lname:
        return NodeType.CAPTION
    if in_references:
        return NodeType.REFERENCE
    if "footnote" in lname:
        return NodeType.FOOTNOTE
    return NodeType.BODY


def _heading_level(lname: str) -> int:
    if lname.startswith("title"):
        return 0
    if lname.startswith("heading"):
        tail = lname.replace("heading", "").strip()
        return int(tail) if tail.isdigit() else 1
    return 0


class StyleFont(NamedTuple):
    """What a paragraph style, and the styles it inherits from, say about type."""
    name: Optional[str] = None
    size: Optional[float] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    color: Optional[str] = None


def _style_font(style) -> StyleFont:
    """Walk a paragraph style and the styles it inherits from for its type.

    Most documents set their typeface once, on a style, and never touch a run.
    Reading only run formatting therefore finds nothing and the document renders
    in whatever the viewer defaults to — which is why an imported file came out
    in the wrong face at the wrong size.

    Emphasis works the same way and was read only off the runs, so a paragraph
    that is italic because its style is italic — an IEEE abstract, a block
    quotation — arrived with no italic at all.
    """
    found: dict[str, Any] = {}
    seen = 0
    while style is not None and seen < 10:      # guard a cyclic base_style chain
        font = getattr(style, "font", None)
        if font is not None:
            if found.get("name") is None and font.name:
                found["name"] = font.name
            if found.get("size") is None and font.size is not None:
                try:
                    found["size"] = float(font.size.pt)
                except (AttributeError, TypeError):
                    pass
            # None means "inherit from the style I am based on", so only a real
            # True/False here settles the question.
            for attr in ("bold", "italic", "underline"):
                if found.get(attr) is None and getattr(font, attr, None) is not None:
                    found[attr] = bool(getattr(font, attr))
            if found.get("color") is None:
                rgb = getattr(getattr(font, "color", None), "rgb", None)
                if rgb is not None:
                    found["color"] = f"#{rgb}"
        if len(found) == 6:
            break
        style = getattr(style, "base_style", None)
        seen += 1
    return StyleFont(**found)


def _document_default_font(doc: _Doc) -> tuple[Optional[str], Optional[float]]:
    """The typeface a document falls back to: the Normal style, then docDefaults."""
    try:
        name, size = _style_font(doc.styles["Normal"])[:2]
    except (KeyError, AttributeError):
        name = size = None

    if name and size:
        return name, size

    try:
        rpr = doc.styles.element.find(qn("w:docDefaults"))
        if rpr is not None:
            fonts = rpr.find(".//" + qn("w:rFonts"))
            if name is None and fonts is not None:
                # A literal face if there is one; otherwise the document names a
                # *theme* font, and the actual typeface lives in theme1.xml.
                name = (fonts.get(qn("w:ascii")) or fonts.get(qn("w:hAnsi"))
                        or _theme_font(doc, fonts.get(qn("w:asciiTheme"))
                                       or fonts.get(qn("w:hAnsiTheme"))))
            sz = rpr.find(".//" + qn("w:sz"))
            if size is None and sz is not None and sz.get(qn("w:val")):
                size = float(sz.get(qn("w:val"))) / 2   # half-points
    except Exception:
        pass
    return name, size


_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def _theme_font(doc: _Doc, theme_ref: Optional[str]) -> Optional[str]:
    """Resolve "minorHAnsi" / "majorHAnsi" to the typeface the theme names."""
    if not theme_ref:
        return None
    scheme = "majorFont" if theme_ref.startswith("major") else "minorFont"
    try:
        for rel in doc.part.rels.values():
            if not rel.reltype.endswith("/theme") or rel.is_external:
                continue
            part = rel.target_part
            # The theme is an opaque part in python-docx, so it usually has to
            # be parsed from its bytes rather than read as an element tree.
            root = getattr(part, "element", None)
            if root is None:
                from lxml import etree
                root = etree.fromstring(part.blob)
            element = root.find(f".//{{{_A_NS}}}{scheme}/{{{_A_NS}}}latin")
            if element is not None:
                return element.get("typeface") or None
    except Exception:
        pass
    return None


def _run_style(run) -> Style:
    """What one run says about itself — and only that.

    Anything it leaves unsaid stays None so it inherits from the paragraph, the
    same way it does in Word. Resolving it here instead would freeze the
    paragraph's formatting into every run, and a later "make the body 12pt"
    would appear to do nothing.
    """
    f = run.font
    color = None
    if f.color is not None and f.color.rgb is not None:
        color = f"#{str(f.color.rgb)}"

    vertical = None
    if getattr(f, "superscript", None):
        vertical = "superscript"
    elif getattr(f, "subscript", None):
        vertical = "subscript"

    return Style(
        font_family=f.name or None,
        font_size=float(f.size.pt) if f.size is not None else None,
        bold=None if run.bold is None else bool(run.bold),
        italic=None if run.italic is None else bool(run.italic),
        underline=None if run.underline is None else bool(run.underline),
        color=color,
        vertical_align=vertical,
    )


def _paragraph_runs(para: _Paragraph) -> list[Run]:
    """The paragraph's inline formatting, or nothing when it has none.

    A paragraph formatted alike throughout needs no runs — its `style` already
    says everything — so the common case stays as small as it was.
    """
    runs = merge_runs(Run(text=r.text or "", style=_run_style(r)) for r in para.runs)
    if len(runs) <= 1:
        return []
    return runs


def _paragraph_style(para: _Paragraph) -> Style:
    """Aggregate run formatting to a paragraph-level style (majority wins)."""
    # Each entry is (value, weight), the weight being how many characters carry
    # it. Counting runs instead let a three-word italic label outvote the
    # paragraph it introduces.
    sizes: list[tuple[float, int]] = []
    bolds: list[tuple[bool, int]] = []
    italics: list[tuple[bool, int]] = []
    underlines: list[tuple[bool, int]] = []
    colors: list[tuple[str, int]] = []
    families: list[tuple[str, int]] = []

    # What the paragraph's style says, so a run that states nothing can still
    # vote with the value it actually inherits. Without this only the runs that
    # spoke up counted, and one bold word carried a whole paragraph.
    from_style = _style_font(getattr(para, "style", None))
    emphasis_stated = False

    for run in para.runs:
        f = run.font
        weight = len(run.text or "")
        if not weight:
            continue                      # an empty run describes no text
        if f.size is not None:
            sizes.append((f.size.pt, weight))
        if f.color is not None and f.color.rgb is not None:
            colors.append((f"#{str(f.color.rgb)}", weight))
        if f.name:
            families.append((f.name, weight))

        for value, inherited, votes in (
            (run.bold, from_style.bold, bolds),
            (run.italic, from_style.italic, italics),
            (run.underline, from_style.underline, underlines),
        ):
            if value is not None or inherited is not None:
                emphasis_stated = True
            # Unstated on both the run and the style means Word's own default,
            # which is upright, unbolded and unlined — a real vote, not silence.
            votes.append((bool(value if value is not None else inherited), weight))

    # A paragraph with no runs at all (or none carrying text) still has a style.
    length = len(para.text or "") or 1
    if not families and from_style.name:
        families.append((from_style.name, length))
    if not sizes and from_style.size:
        sizes.append((from_style.size, length))
    if not colors and from_style.color:
        colors.append((from_style.color, length))
    if not bolds and from_style.bold is not None:
        bolds.append((from_style.bold, length)); emphasis_stated = True
    if not italics and from_style.italic is not None:
        italics.append((from_style.italic, length)); emphasis_stated = True
    if not underlines and from_style.underline is not None:
        underlines.append((from_style.underline, length)); emphasis_stated = True

    # Nobody ever mentioned emphasis, so the node says nothing about it either
    # rather than asserting a bold=False that would override a future style.
    if not emphasis_stated:
        bolds = italics = underlines = []

    align = None
    if para.alignment is not None:
        align = {0: "left", 1: "center", 2: "right", 3: "justify"}.get(int(para.alignment), None)

    def _majority(weighted: list[tuple[Any, int]]) -> Optional[Any]:
        """The value carrying the most text, and on a tie the one that came first.

        `max(set(...))` used to leave a tie to set ordering, so a paragraph split
        evenly between italic and upright runs came out one way on one import and
        the other way on the next.
        """
        if not weighted:
            return None
        totals: dict[Any, int] = {}
        for value, weight in weighted:
            totals[value] = totals.get(value, 0) + weight
        order = [value for value, _ in weighted]
        return max(totals, key=lambda v: (totals[v], -order.index(v)))

    return Style(
        font_family=_majority(families),
        font_size=_majority(sizes),
        bold=_majority(bolds),
        italic=_majority(italics),
        underline=_majority(underlines),
        color=_majority(colors),
        alignment=align,
    )


def _is_header_row(row, index: int) -> bool:
    """Is this the table's header row?

    Word says so outright when the row is set to repeat across pages
    (`w:tblHeader`). Most tables never set it and simply bold the first row, so
    that counts too. Either way it is a distinct thing from a document heading,
    and "the headings in the table" means these.
    """
    try:
        properties = row._tr.find(qn("w:trPr"))
        if properties is not None and properties.find(qn("w:tblHeader")) is not None:
            return True
    except Exception:
        pass

    if index != 0:
        return False

    # The first row of a table is its header, which is the convention papers
    # follow and the reason "the headings in the table" means anything at all.
    # Requiring it to be bold already was exactly backwards: it made the row
    # undetectable in the one case someone asks about it — when they want it
    # made bold — so the request found nothing and reported success.
    #
    # Two things disqualify it. A single-row table is a layout device, usually
    # an equation with its number beside it, and has no header. A first row
    # with an empty cell is a table that starts with data.
    if len(row.table.rows) < 2:
        return False
    return all((cell.text or "").strip() for cell in row.cells)


def _table_node(table: _Table, doc: _Doc) -> Node:
    rows: list[Node] = []
    for row_index, r in enumerate(table.rows):
        cells: list[Node] = []
        for c in r.cells:
            # A cell's pictures come with it. Reading only c.text dropped any
            # screenshot placed in a table — a common way to lay out figures.
            pictures: list[Node] = []
            for para in c.paragraphs:
                for img in _images(para):
                    pictures.append(Node(type=NodeType.IMAGE,
                                         metadata={"is_figure": True, **img}))

            # `c.text` has the same blind spot as `para.text`: it reads the
            # w:t runs and skips Word's equations entirely. A display equation
            # is very often laid out as a one-row table with the number beside
            # it, so that blind spot emptied exactly the cells a paper puts its
            # mathematics in.
            lines: list[str] = []
            cell_maths: list[dict] = []
            for para in c.paragraphs:
                parts = paragraph_parts(para)
                if any(kind == "maths" for kind, _text, _el in parts):
                    lines.append("".join(text if kind == "text" else f"${text}$"
                                         for kind, text, _el in parts))
                    cell_maths.extend(
                        {"latex": text, "xml": etree.tostring(element, encoding="unicode")}
                        for kind, text, element in parts if kind == "maths")
                else:
                    lines.append(para.text)

            cell_meta: dict = {"equations": cell_maths} if cell_maths else {}
            cells.append(
                Node(
                    type=NodeType.TABLE_CELL,
                    content="\n".join(lines),
                    children=pictures,
                    metadata=cell_meta,
                )
            )
        row_meta: dict = {}
        if _is_header_row(r, row_index):
            row_meta["header_row"] = True
        rows.append(Node(type=NodeType.TABLE_ROW, children=cells, metadata=row_meta))
    meta: dict = {"rows": len(table.rows), "cols": len(table.columns)}
    # A long table Word split across pages carries lastRenderedPageBreak markers
    # inside its cells; count them so the total page count stays accurate.
    internal_breaks = len(table._tbl.findall(".//" + qn("w:lastRenderedPageBreak")))
    if internal_breaks:
        meta["extra_pages"] = internal_breaks
    return Node(type=NodeType.TABLE, children=rows, metadata=meta)


def _attach_headers_footers(doc: _Doc, root: Node) -> None:
    for section in doc.sections:
        htext = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
        ftext = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
        if htext:
            root.children.insert(0, Node(type=NodeType.HEADER, content=htext))
        if ftext:
            root.children.append(Node(type=NodeType.FOOTER, content=ftext))
        break  # first section only — keeps the graph clean


# ── low-level XML probes ────────────────────────────────────────────────────

def _has_drawing(para: _Paragraph) -> bool:
    """Any drawing markup at all — a picture, but also a shape or a rule."""
    return bool(para._p.findall(".//" + qn("w:drawing"))) or bool(
        para._p.findall(".//" + qn("w:pict"))
    )


def _image_rel_ids(para: _Paragraph) -> list[str]:
    """Relationship ids of the *pictures* in this paragraph.

    A picture is markup that points at an image part: a DrawingML blip, or a VML
    shape with imagedata. Drawing markup on its own proves nothing — Word writes
    a horizontal rule as a VML rect inside w:pict, and treating that as a picture
    is what turned every rule in an imported document into a figure.
    """
    ids: list[str] = []
    for blip in para._p.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rid:
            ids.append(rid)
    for imagedata in para._p.findall(".//" + _V_IMAGEDATA):
        rid = imagedata.get(qn("r:id")) or imagedata.get(qn("r:href"))
        if rid:
            ids.append(rid)
    return ids


def _images(para: _Paragraph) -> list[dict[str, Any]]:
    """The pictures in this paragraph, each with its bytes inlined as a data URI.

    A picture whose part cannot be read, or that is too large to inline, still
    yields a node — the document keeps its shape, and only the pixels are
    missing.
    """
    out: list[dict[str, Any]] = []
    for rid in _image_rel_ids(para):
        item: dict[str, Any] = {"rel_id": rid}
        try:
            item.update(_resolve_image(para.part, rid))
        except Exception:
            item["unreadable"] = True   # the node still holds the picture's place
        out.append(item)
    return out


def _resolve_image(part: Any, rid: str) -> dict[str, Any]:
    """Read one image relationship.

    Two kinds arrive. An *embedded* image has its bytes inside the file and is
    inlined as a data URI. A *linked* image has none — the document only points
    at where the picture lives, which is what LibreOffice writes when it converts
    HTML and what Word writes for "Link to File". Those are not corrupt, and
    reporting them as unreadable was wrong: an http(s) target is passed through
    for the browser to load, exactly as Word would, and anything else is marked
    as linked so the reader is told the picture is not in the document.
    """
    rel = part.rels[rid]

    if getattr(rel, "is_external", False):
        target = str(getattr(rel, "target_ref", "") or "")
        if target.lower().startswith(("http://", "https://")):
            return {"src": target, "external": True}
        return {"linked_to": target, "external": True}

    blob = rel.target_part.blob
    if len(blob) > _MAX_INLINE_IMAGE_BYTES:
        return {"too_large": len(blob)}

    content_type = getattr(rel.target_part, "content_type", "") or "image/png"
    return {
        "src": f"data:{content_type};base64," + base64.b64encode(blob).decode("ascii"),
        "bytes": len(blob),
    }


def _is_vml_rule(para: _Paragraph) -> bool:
    """Is this paragraph one of Word's horizontal rules?

    Word writes a rule two ways, and both are drawing markup. The plain one is a
    VML rect. The one from the Horizontal Line gallery is a *picture of a line* —
    a VML shape carrying imagedata — so "it references an image" cannot be the
    test. What both share is the o:hr flag, which says outright that the shape is
    a rule, wherever it sits in the markup.
    """
    for pict in para._p.findall(".//" + qn("w:pict")):
        for element in pict.iter():
            if element.get(_O_HR) is not None:
                return True
        # a rect with no image behind it is decoration, not a picture
        if not pict.findall(".//" + _V_IMAGEDATA) and pict.findall(".//" + _V_RECT):
            return True
    return False


def _page_break_signals(para: _Paragraph) -> tuple[bool, int]:
    """Return (has_explicit_break, rendered_break_count).

    `explicit`  — a manual page break the author inserted (w:br type=page).
    `rendered`  — how many `w:lastRenderedPageBreak` markers Word left in this
                  paragraph, i.e. how many page boundaries fall on it.
    """
    explicit = any(
        br.get(qn("w:type")) == "page"
        for br in para._p.findall(".//" + qn("w:br"))
    )
    rendered = len(para._p.findall(".//" + qn("w:lastRenderedPageBreak")))
    return explicit, rendered


def _has_bottom_border(para: _Paragraph) -> bool:
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        return False
    return pBdr.find(qn("w:bottom")) is not None or pBdr.find(qn("w:top")) is not None
